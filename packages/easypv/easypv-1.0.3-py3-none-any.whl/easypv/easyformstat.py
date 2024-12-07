#!/usr/bin/env python
# coding: utf-8

# 第一部分：程序说明###################################################################################
# coding=utf-8
# 开发人：蔡权周

# 第二部分：导入基本模块及初始化 ########################################################################
import threading
import tkinter as tk  
from tkinter import ttk, scrolledtext,simpledialog, messagebox, filedialog, simpledialog,StringVar
from tkinter.scrolledtext import ScrolledText 
from tkinter.messagebox import showinfo
from datetime import datetime, timedelta
import datetime
import time
import pandas as pd
import re 
import matplotlib.pyplot as plt  
import numpy as np 
from matplotlib.backends.backend_tkagg import FigureCanvasTkAgg 
from matplotlib.backends.backend_tkagg import NavigationToolbar2Tk
import seaborn as sns 
from collections import defaultdict 
import math
from datetime import datetime 
from collections import Counter  
import ast  
import os
import hashlib
import random
import itertools
import sqlite3 
import scipy.stats as stats
from pandas.core.groupby.groupby import GroupBy  
from pandas.core.base import NoNewAttributesMixin  
from scipy.stats import ranksums
import scipy.stats as st
from matplotlib import cm  
import zipfile
from openpyxl import load_workbook
from sqlalchemy import create_engine, text

import argparse


#import jieba as jb
#from wordcloud import WordCloud
from matplotlib.text import Text  
from matplotlib.transforms import Affine2D  
from scipy.stats import iqr  
from scipy.stats import chi2_contingency, fisher_exact  
from scipy.special import factorial  
from docx import Document

global version_now
global usergroup
global setting_cfg
global csdir
global peizhidir
global biaozhun
global global_dfs
global psur
global mytitle

global_dfs = {}
biaozhun={}
version_now="1.0.0" 

csdir =str (os .path .abspath (__file__ )).replace (str (__file__ ),"")
if csdir=="":
    csdir =str (os .path .dirname (__file__ ))#
    csdir =csdir +csdir.split ("easyformstat")[0 ][-1 ]#

parser = argparse.ArgumentParser(description='传入启动参数')
# 使用选项参数，并且设置为布尔类型
parser.add_argument('--psur', action='store_true', help='是否启动定制版本（默认不启动）')
# 如果没有提供 --psurs，则 args.psurs 将为 False
args = parser.parse_args()
 
# 使用 args.psurs
if args.psur:
    psur=1
else:
    psur=0 #超级账户 psur==1 为精简版


#
usergroup="用户组="+str(psur)


if psur==1:
    mytitle='药品品种监测数据辅助分析工具 EFS_PINZHONG_'+version_now
else:
    mytitle='易析数据分析工具 EFS_'+version_now
    

#############################################
#备份的
def AAAAA_bak():
    pass


#########################################
#用户组和配置表文件编码
def AAAAA_psur():
    pass
########################################

def INI_extract_zip_file(zip_file_path, extract_path):
    #import shutil
    import zipfile
    if extract_path=="":
        return 0
    with zipfile.ZipFile(zip_file_path, 'r') as zip_ref:
        for file_info in zip_ref.infolist():

            file_info.filename = file_info.filename.encode('cp437').decode('gbk')
            zip_ref.extract(file_info, extract_path)
            #source_path = os.path.join(extract_path, file_info.filename)
            #print(file_info)
            #target_path = os.path.join(extract_path, file_info.filename.encode('gbk').decode('utf-8'))
            #shutil.move(source_path, extract_path)
            
def INI_get_directory_path(directory_path):
    global csdir  # 假设 csdir 是之前定义好的包含 zip 文件的目录
    
    # 检查目录是否存在指定的文件
    file_path = os.path.join(directory_path, '！！！配置表版本（请勿删除本文件）.txt')
    if not os.path.isfile(file_path):
        # 创建一个 Tkinter 根窗口（隐藏主窗口）
        root = Toplevel()
        root.withdraw()  # 隐藏主窗口
        from tkinter import messagebox
        # 弹出确认框
        message = "程序将在该目录内生成相关配置文件。这个目录内的同名文件将会被取代，建议做好备份，请问是否继续？"
        user_response = messagebox.askyesno("确认解压", message)
        
        # 根据用户响应决定是否解压
        if user_response:
            # 假设 csdir + "def.zip" 是正确的 zip 文件路径
            zip_file_path = os.path.join(csdir, "def.py")  # 修改为正确的 zip 文件名
            extract_zip_file(zip_file_path, directory_path)
        else:
            # 用户选择否，退出程序
            root.destroy()  # 销毁隐藏的 Tkinter 窗口
            quit()
    
    # 检查目录路径是否为空，如果为空则退出程序
    if directory_path == "":
        quit()
    
    # 返回目录路径
    return directory_path

    


def INI_convert_and_compare_dates(date_str):
    import datetime
    current_date = datetime.datetime.now()

    try:
       date_obj = datetime.datetime.strptime(str(int(int(date_str)/4)), "%Y%m%d") 
    except:
        print("fail")
        return  "已过期"
  
    if date_obj > current_date:
        
        return "未过期"
    else:
        return "已过期"
    
def INI_read_setting_cfg():
    global csdir
    # 读取 setting.cfg 文件
    if os.path.exists(csdir+'setting.cfg'):
        print("已完成初始化\n")
        with open(csdir+'setting.cfg', 'r') as f:
            setting_cfg = eval(f.read())
    else:
        # 创建 setting.cfg 文件，如果文件已存在则覆盖
        setting_cfg_path =csdir+ 'setting.cfg'
        with open(setting_cfg_path, 'w') as f:
            f.write('{"settingdir": 0, "sidori": 0, "sidfinal": "11111180000808"}')
        print("未初始化，正在初始化...\n")
        setting_cfg = INI_read_setting_cfg()
    return setting_cfg
    

def INI_open_setting_cfg():
    global csdir
    # 打开 setting.cfg 文件
    with open(csdir+"setting.cfg", "r") as f:
        # 将文件内容转换为字典
        setting_cfg = eval(f.read())
    return setting_cfg

def INI_update_setting_cfg(keys,values):
    global csdir
    # 打开 setting.cfg 文件
    with open(csdir+"setting.cfg", "r") as f:
        # 将文件内容转换为字典
        setting_cfg = eval(f.read())
    
    if setting_cfg[keys]==0 or setting_cfg[keys]=="11111180000808" :
        setting_cfg[keys]=values
        # 保存字典覆盖源文件
        with open(csdir+"setting.cfg", "w") as f:
            f.write(str(setting_cfg))


def INI_generate_random_file():
    # 生成一个六位数的随机数
    random_number = random.randint(200000, 299999)
    # 将随机数保存到文本文件中
    INI_update_setting_cfg("sidori",random_number)

def INI_display_random_number():
    global csdir
    mroot = Toplevel()
    mroot.title("ID")
    
    sw = mroot.winfo_screenwidth()
    sh = mroot.winfo_screenheight()
    # 得到屏幕高度
    ww = 80
    wh = 70
    # 窗口宽高为100
    x = (sw - ww) / 2
    y = (sh - wh) / 2
    mroot.geometry("%dx%d+%d+%d" % (ww, wh, x, y))
    
    # 打开 setting.cfg 文件
    with open(csdir+"setting.cfg", "r") as f:
        # 将文件内容转换为字典
        setting_cfg = eval(f.read())
    random_number=int(setting_cfg["sidori"])
    sid=random_number*2+183576

    print(sid)
    # 创建标签和输入框
    label = ttk.Label(mroot, text=f"机器码: {random_number}")
    entry = ttk.Entry(mroot)

    # 将标签和输入框添加到窗口中
    label.pack()
    entry.pack()

    # 监听输入框的回车键事件
    #entry.bind("<Return>", INI_check_input)
    ttk.Button(mroot, text="验证", command=lambda:INI_check_input(entry.get(),sid)).pack()
    
def INI_check_input(input_numbers,sid):

    # 将输入的数字转换为整数'

    try:
        input_number = int(str(input_numbers)[0:6])
        day_end=INI_convert_and_compare_dates(str(input_numbers)[6:14])
    except:
        showinfo(title="提示", message="不匹配，注册失败。")
        return 0
    # 核对输入的数字是否等于随机数字
    if input_number == sid and day_end=="未过期":
        INI_update_setting_cfg("sidfinal",input_numbers)
        showinfo(title="提示", message="注册成功,请重新启动程序。")
        quit()
    else:
        showinfo(title="提示", message="不匹配，注册失败。")

###############################


def INI_update_software(package_name):
    # 检查当前安装的版本
    global version_now   
    print("当前版本为："+version_now+",正在检查更新...(您可以同时执行分析任务)") 
    try: 
        latest_version = requests.get(f"https://pypi.org/pypi/{package_name}/json",timeout=2).json()["info"]["version"]
    except:
        return "...更新失败。"
    if latest_version > version_now:
        text.insert(END,"\n最新版本为："+latest_version+",正在尝试自动更新....")        
        # 如果 PyPI 中有更高版本的软件，则使用 `pip install --upgrade` 进行更新
        pip.main(['install', package_name, '--upgrade'])
        text.insert(END,"\n您可以开展工作。")
        return "...更新成功。"












#########################################
#兼容PSUR项目
def AAAAA_psur():
    pass
########################################
#

def PSUR_open(data):
    if isinstance(data, pd.DataFrame):
        psur_data=data	
    else:
        psur_data=SMALL_read_excel_files()
    
    if '是否已清洗' not in psur_data.columns:
        print('●正在执行基础清洗...')
        psur_data=CLEAN_replay_operations(psur_data,biaozhun["药品清洗"]) 
        print('●正在执行添加药品分类...') 
        try:       
            CLEAN_easystat(psur_data,biaozhun["药品分类库"],'赋关键词')
        except:
            print('药品分类添加失败，原因可能是药品分类库没有正确配置。')
        
        print('●正在执行扩行...')
        psur_data=CLEAN_expand_rows(psur_data, '；', ["不良反应名称（规整）"])
        print('●正在执行PT规整...')
        psur_data=CLEAN_replay_operations(psur_data,biaozhun["药品PT清洗"])
        print('●再次执行扩行...')
        psur_data=CLEAN_expand_rows(psur_data, '；', ["不良反应名称（规整）"])
        print('●正在执行SOC映射...')
        psur_data=PSUR_merge_dataframes(psur_data, "不良反应名称（规整）",biaozhun["meddra"], df2_col='PT') 
        print('●正在执行用量单位规整...')   
        psur_data['用量（规整后）']= psur_data['用量'].copy()
        psur_data['用量单位（规整后）']= psur_data['用量单位'].copy()
        psur_data.loc[psur_data['用量单位（规整后）'].isin(['g', '克']), '用量（规整后）'] *= 1000  
        psur_data.loc[psur_data['用量单位（规整后）'].isin(['g', '克','毫克']), '用量单位（规整后）'] = 'mg'
        psur_data['用法用量']=psur_data['用量（规整后）'].astype(str)+psur_data['用量单位（规整后）'].astype(str)+' '+psur_data['用法-日'].astype(str)+'日'+psur_data['用法-次'].astype(str)+'次'
        print('●正在执行重点关注标注...')
        psur_data=CLEAN_replay_operations(psur_data,pd.read_excel(peizhidir+'easy_药品规整-基础清洗.xlsx',sheet_name="扩行后").reset_index(drop=True) ) 

        psur_data['疑似新的ADR']='未载入说明书'  
        psur_data['SOC']=psur_data['SOC'].fillna('其他（未规整）')     
        psur_data.loc[(psur_data['重点关注']=='重点关注'),'重点关注ADR'] = psur_data.loc[psur_data['重点关注']=='重点关注','不良反应名称（规整）']
        psur_data['是否已清洗']='是'
        print('●数据清洗完成。')          
    else:
        print('●检测到导入的数据前期已清洗过，不再执行数据清洗。')        
    #PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(psur_data,[['怀疑/并用','批准文号','通用名称'], ['-伤害'], ['报告表编码'], ['nunique'], {'不良反应名称（规整）': 'count_ALL','重点关注ADR': 'count_ALL'},  ['报告表编码合计']]),0,psur_data)     

    PROGRAM_display_df_in_treeview(psur_data,'psur',psur_data) 

def PSUR_ini(data):  
    """快速启动PSUR子项目"""  
    global psur
    
    print('程序正在启动...')
    print('正在载入，请稍候...')
    try:
        biaozhun["药品清洗"]=pd.read_excel(peizhidir+'easy_药品规整-基础清洗.xlsx').reset_index(drop=True) 
        biaozhun["药品PT清洗"]=pd.read_excel(peizhidir+'easy_药品规整-PT标准化.xlsx').reset_index(drop=True)  
    except:
        for i in range(1000000000000):   
            a=input("未找到标准库和品种分析库，请联系开发者索取。")     
    if isinstance(biaozhun["meddra"], int):
        try:
            biaozhun["meddra"]=pd.read_excel(peizhidir+'share（easy_adrmdr）药品规整-SOC-Meddra库.xlsx').reset_index(drop=True)  
            print("已载入品种分析专用库。")
        except:
            for i in range(1000000000000):   
                a=input("未在配置表文件夹找到品种分析专用库，请联系开发者索取。") 
    if psur==0:
        PSUR_open(data)	




def PSUR_check_adr_in_word(df): 
    '''载入说明书进行比对 ''' 
    # 弹出文件选择对话框，让用户选择一个Word文档  
    root = tk.Tk()  
    root.withdraw()  # 隐藏tkinter的主窗口  
    file_path = filedialog.askopenfilename(  
        title="选择Word文档",  
        filetypes=[("Word files", "*.docx")]  
    )  
    root.destroy()  # 销毁tkinter窗口  
      
    if not file_path:  
        return df  # 如果用户取消选择文件，则直接返回原df  
      
    # 读取Word文档内容  
    doc = Document(file_path)  
    word_content = ''  
      
    # 提取段落文本  
    for paragraph in doc.paragraphs:  
        word_content += paragraph.text + ' '  
      
    # 提取表格中的文本  
    for table in doc.tables:  
        for row in table.rows:  
            for cell in row.cells:  
                word_content += cell.text + ' '  
                
    df['疑似新的ADR']=""  
    # 遍历df的每一行  
    for index, row in df.iterrows():  
        adr = row['不良反应名称（规整）']  
        if adr not in word_content:  
            df.at[index, '疑似新的ADR'] = adr  
      
    PROGRAM_display_df_in_treeview(df,'psur',0)  

   
def PSUR_merge_dataframes(df1, df1_col, df2, df2_col='PT'):  
    
    merged_df = pd.merge(df1, df2, left_on=df1_col, right_on=df2_col, how='left')  
      
    return merged_df 


def PSRU_ori_owercount_easyread(df):
    if '报表类型' in df.columns:
        m=df.loc[0,'报表类型']
        df=SMALL_expand_dict_like_columns(df)
        df['报表类型']=m
        new_columns = [col.replace('报告表编码', '') if '报告表编码' in col else col for col in df.columns]  
        df.columns = new_columns
        for col in df.columns:  
            try:  
                df[col] = df[col].fillna(0).astype(int)  
            except ValueError:  
                pass
        
    return df
#########################################
##统计函数 (暂时没有发挥作用） 
def not_finish_SMALL_stats(numbers):  
      
    # 计算均值  
    mean = np.mean(numbers)  
      
    # 计算标准差  
    std = np.std(numbers)  
      
    # 计算均值加减标准差  
    mean_plus_1std = mean + std  
    mean_plus_2std = mean + 2 * std  
    mean_plus_3std = mean + 3 * std  
      
    # 计算分位数  
    q25 = np.percentile(numbers, 25)  
    q50 = np.percentile(numbers, 50)  
    q75 = np.percentile(numbers, 75)  
      
    # 计算IQR  
    iqr = q75 - q25  
      
    # 计算1.5倍IQR  
    iqr_multiplier = 1.5 * iqr  
      
    # 初始化信号提示  
    signal = 0  
      
    # 检查是否有值超过阈值，并更新信号提示  
    for num in numbers:  
        if num > mean_plus_2std:  
            signal += 1  
        if num > mean_plus_3std:  
            signal += 1  
        if num > iqr_multiplier:  
            signal += 1  
      
    # 将结果存储在字典中  
    statistics = {  
        'Mean': mean,  
        'Std': std,  
        'Mean+1Std': mean_plus_1std,  
        'Mean+2Std': mean_plus_2std,  
        'Mean+3Std': mean_plus_3std,  
        '25th Percentile': q25,  
        '50th Percentile': q50,  
        '75th Percentile': q75,  
        'IQR': iqr,  
        '1.5xIQR': iqr_multiplier,  
        'Signal': signal  
    }  
      
    return statistics  



#PSUR
def TEST_psur(df1,separator0, df2=None):    
    if df2 is not None:    
        # 根据df2对df1增加一列“事件分类”    
        # 初始化'事件分类'列为空字符串  
        df1['事件分类'] = ''  
          
        def update_classification(row):  
            mask = ((df1['不良反应-术语'].str.contains(str(row['值']), na=False))&(~df1['不良反应-术语'].str.contains(str(row['排除值']), na=False)))
            df1.loc[mask, '事件分类'] += separator0 + str(row['分类'])  
  
        # 使用apply函数代替循环  
        df2.apply(update_classification, axis=1)  
  
        # 删除以";"开头的值（如果需要的话）  
        df1['事件分类'] = df1['事件分类'].str.replace("^;", "")
        print(df1)
    # 创建GUI让用户选择参数    
    root = tk.Tk()    
    root.title("数据处理工具")    
        
    # 变量用于存储用户的选择    
    selected_event_column = tk.StringVar(root)    
    selected_pivot_column = tk.StringVar(root)    
    selected_name_column = tk.StringVar(root)    
    selected_code_column = tk.StringVar(root)    
    selected_separator = tk.StringVar(root)    
        
    def on_ok(df1):     
        # 获取用户的选择    
        event_column = selected_event_column.get()    
        pivot_column = selected_pivot_column.get()    
        name_column = selected_name_column.get()    
        code_column = selected_code_column.get()    
        separator = selected_separator.get()    
            
        # 根据选择的事件列和分隔符扩行  
        if df2 is not None:    
            df1 = df1.assign(**{event_column: df1[event_column].str.split(separator0)}).explode(event_column)    
        else:
            df1 = df1.assign(**{event_column: df1[event_column].str.split(separator)}).explode(event_column) 			    
        # 生成List1：事件列组按透视列进行数据透视（计算：编码列nunique）    
        list1 = df1.groupby(event_column)[pivot_column].unique().apply(pd.Series).stack().reset_index(level=1, drop=True).to_frame('编码列').groupby(level=0)['编码列'].nunique().reset_index().rename(columns={0: 'nunique'})    
            
        # 生成List2：按事件列分组（计算：名称列sum） 
        aggregations = {name_column: lambda x: SMALL_count_mode(x, separator)}  
        list2 = df1.groupby(event_column).agg(aggregations).reset_index()
   
        #list2 = df1.groupby(event_column)[name_column].sum().reset_index()#.rename(columns={0: lambda x: SMALL_count_mode(x, separator) })    
            
        # list1拼接list2(按事件列,左)    
        result = pd.merge(list1, list2, on=event_column, how='left')    
            
        # 打印最后的结果    
        print(result)    
        messagebox.showinfo("完成", "处理完成，请查看控制台输出")    
        #root.destroy()    
        
    def on_cancel():    
        root.destroy()    
        
    # 创建下拉菜单和标签等GUI元素    
    msdos=df1.columns.to_list()  
    ttk.Label(root, text="请选择事件列:").grid(row=0, column=0, padx=5, pady=5)    
    event_column_dropdown = ttk.Combobox(root, textvariable=selected_event_column, values=msdos)    
    event_column_dropdown.grid(row=0, column=1, padx=5, pady=5)    
    event_column_dropdown.current(0)  # 默认选择第一个列名    
        
    ttk.Label(root, text="请选择透视列:").grid(row=1, column=0, padx=5, pady=5)    
    pivot_column_dropdown = ttk.Combobox(root, textvariable=selected_pivot_column, values=msdos)    
    pivot_column_dropdown.grid(row=1, column=1, padx=5, pady=5)    
    pivot_column_dropdown.current(0)  # 默认选择第一个列名    
        
    ttk.Label(root, text="请选择字典列:").grid(row=2, column=0, padx=5, pady=5)    
    name_column_dropdown = ttk.Combobox(root, textvariable=selected_name_column, values=msdos)    
    name_column_dropdown.grid(row=2, column=1, padx=5, pady=5)    
    name_column_dropdown.current(0)  # 默认选择第一个列名    
        
    ttk.Label(root, text="请选择编码列:").grid(row=3, column=0, padx=5, pady=5)    
    code_column_dropdown = ttk.Combobox(root, textvariable=selected_code_column, values=msdos)    
    code_column_dropdown.grid(row=3, column=1, padx=5, pady=5)    
    code_column_dropdown.current(0)  # 默认选择第一个列名    
        
    ttk.Label(root, text="请选择分隔符:").grid(row=4, column=0, padx=5, pady=5)    
    separator_dropdown = ttk.Combobox(root, textvariable=selected_separator, values=[',', ';', ' ', '\t'])    
    separator_dropdown.grid(row=4, column=1, padx=5, pady=5)    
    separator_dropdown.current(0)  # 默认选择第一个分隔符    
        
    # 创建确定和取消按钮    
    ttk.Button(root, text="确定", command=lambda:on_ok(df1)).grid(row=5, column=0, padx=5, pady=(10, 0))    
    ttk.Button(root, text="取消", command=on_cancel).grid(row=5, column=1, padx=5, pady=(10, 0))    
        
    root.mainloop()  # 运行GUI主循环，等待用户操作






  








#######################################
#定制的和个性化设置的自定义函数  
def AAAA_Setting():
    pass
######################################



def SETTING_create_menu(data):  
    """创建菜单栏，包含File、Edit和Help三个菜单项。"""  

    win=data["windows"]
    win_progressbar=data["win_progressbar"]
    ori_owercount_easyread=data["ori_owercount_easyread"]
    ori=data["ori"]   
    datacols=ori_owercount_easyread.columns.to_list() 
    menu_bar = tk.Menu(win)  
    win.config(menu=menu_bar)
    
    file_menu = tk.Menu(menu_bar, tearoff=0)  
    menu_bar.add_cascade(label="文件", menu=file_menu)  
    file_menu.add_command(label="载入文件(xls)", command=lambda:PROGRAM_display_df_in_treeview(SMALL_read_excel_files(),0,0)) 
    file_menu.add_command(label="载入文件(csv)", command=lambda:PROGRAM_display_df_in_treeview(SMALL_read_csv_files(),0,0))     
    
    
    file_menu.add_command(label="导出文件", command=lambda:SMALL_save_dict(ori_owercount_easyread))     
    file_menu.add_separator()    
    file_menu.add_command(label="从数据库载入", command=lambda:SQL_create_query_gui(ori_owercount_easyread))   
    file_menu.add_command(label="另存为数据库", command=lambda:SQL_df_to_sqlite_db_with_gui(ori_owercount_easyread))      
    file_menu.add_command(label="添加到数据库", command=lambda:SQL_update_sqlite_db_with_df(ori_owercount_easyread))           

    file_menu.add_separator()           
    file_menu.add_command(label="文件转字典", command=lambda:PROGRAM_display_content_in_textbox(str(ori_owercount_easyread.to_dict(orient='list')))) 
    file_menu.add_command(label="比较excel文件", command=SMALL_compare_excel_files)   

    file_menu.add_separator()    
    file_menu.add_command(label="切换视图", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_easyreadT(ori_owercount_easyread),0,0))  
    file_menu.add_separator()   
    file_menu.add_command(label="问题和建议", command=lambda:showinfo(title="联系我们", message="如有任何问题或建议，请联系蔡老师，411703730（微信或QQ）。"))  




    nomal_menu = tk.Menu(menu_bar, tearoff=0)  
    menu_bar.add_cascade(label="预制清洗", menu=nomal_menu)
    

    
    #加载器械清洗菜单和相关清洗任务
    nomal_menu.add_command(label="自定义标准清洗(字典任务)", command=lambda:PROGRAM_display_df_in_treeview(CLEAN_replay_operations(ori_owercount_easyread.copy(),SMALL_read_excel_files()),0,0)) 
    nomal_menu.add_command(label="自定义标准清洗（赋关键词）", command=lambda:PROGRAM_display_df_in_treeview(CLEAN_easystat(ori_owercount_easyread.copy(),SMALL_read_excel_files(),'赋关键词'),0,0)) 
    nomal_menu.add_command(label="自定义标准清洗（加关键词）", command=lambda:PROGRAM_display_df_in_treeview(CLEAN_easystat(ori_owercount_easyread.copy(),SMALL_read_excel_files(),'加关键词'),0,0)) 

    nomal_menu.add_separator()
    if biaozhun!={} and '器械故障表现' in datacols:
        nomal_menu.add_command(label="○器械预制-基础清洗-标准库", command=lambda:PROGRAM_display_df_in_treeview(CLEAN_replay_operations(ori_owercount_easyread.copy(),biaozhun["器械清洗"]),0,0)) 
        nomal_menu.add_command(label="○ 器械预制-加SOC（标准库-关键词法-之后SOC不扩行）", command=lambda:PROGRAM_display_df_in_treeview(CLEAN_easystat(ori_owercount_easyread.copy(),biaozhun["器械关键词"],'加关键词'),0,0)) 
        nomal_menu.add_command(label="○ 器械预制-加SOC（标准库-关键词法-之后SOC扩行）", command=lambda:PROGRAM_display_df_in_treeview(CLEAN_expand_rows(CLEAN_easystat(ori_owercount_easyread.copy(),biaozhun["器械关键词"],'加关键词'),';',  ["-表现归类(关键词法)"]),0,0)) 
        nomal_menu.add_command(label="○ 器械预制-加SOC（标准库-关键词法-之后SOC不扩行-（仅故障表现））", command=lambda:PROGRAM_display_df_in_treeview(CLEAN_easystat(ori_owercount_easyread.copy(),biaozhun["器械关键词（仅故障表现）"],'加关键词'),0,0)) 
        nomal_menu.add_command(label="○ 器械预制-加SOC（标准库-关键词法-之后SOC扩行-（仅故障表现））", command=lambda:PROGRAM_display_df_in_treeview(CLEAN_expand_rows(CLEAN_easystat(ori_owercount_easyread.copy(),biaozhun["器械关键词（仅故障表现）"],'加关键词'),';',  ["-表现归类(关键词法)"]),0,0)) 

    if biaozhun!={} and '不良反应名称' in datacols:
        nomal_menu.add_separator()
        nomal_menu.add_command(label="●药品预制-基础清洗-标准库", command=lambda:PROGRAM_display_df_in_treeview(CLEAN_replay_operations(ori_owercount_easyread.copy(),biaozhun["药品清洗"]),0,0))   
        nomal_menu.add_command(label="● 药品预制-PT标准化清洗-标准库", command=lambda:PROGRAM_display_df_in_treeview(CLEAN_replay_operations(ori_owercount_easyread.copy(),biaozhun["药品PT清洗"]),0,0))     
        nomal_menu.add_command(label="● 药品预制-加SOC（PT扩行-关键词法-标准库）", command=lambda:PROGRAM_display_df_in_treeview(CLEAN_easystat(CLEAN_expand_rows(ori_owercount_easyread.copy(), '；', ["不良反应名称（规整）"]),biaozhun["药品关键词"],'加关键词'),0,0)) 
        nomal_menu.add_command(label="● 药品预制-加SOC（PT扩行-MedDRA法）", command=lambda:PROGRAM_display_df_in_treeview(SMALL_merge_dataframes(CLEAN_expand_rows(ori_owercount_easyread.copy(), '；', ["不良反应名称（规整）"]), "不良反应名称（规整）", peizhidir+'share（easy_adrmdr）药品规整-SOC-Meddra库.xlsx', df2_col='PT'),0,0)) 

        nomal_menu.add_separator()

        
        nomal_menu.add_command(label="以药品定制模式打开（执行规整）", command=lambda:PSUR_ini(ori_owercount_easyread.copy()))  
        nomal_menu.add_command(label="以药品定制模式打开（不执行规整）", command=lambda:PROGRAM_display_df_in_treeview(ori_owercount_easyread.copy(),'psur',0)) 


     

    if '---器械数据规整---' in datacols or '---药品数据规整---' in datacols: 
        nomal_menu2 = tk.Menu(menu_bar, tearoff=0)  
        menu_bar.add_cascade(label="预制统计", menu=nomal_menu2)

        if 1==1:
            nomal_menu2.add_command(label="报告年份", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['报告年份'], [], ['报告编码'], ['nunique'], '', ['报告编码']]),0,ori)) 
            nomal_menu2.add_command(label="事件发生年份", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['事件发生年份'], [], ['报告编码'], ['nunique'], '', ['报告编码']]),0,ori)) 
            nomal_menu2.add_separator()
            
            nomal_menu2.add_command(label="-注册人备案人", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['-注册人备案人'], ['-伤害'], ['报告编码'], ['nunique'], '', ['报告编码合计']]),0,ori)) 
            nomal_menu2.add_command(label="-产品类别", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['产品类别'], ['-伤害'], ['报告编码'], ['nunique'], '', ['报告编码合计']]),0,ori)) 
            nomal_menu2.add_command(label="-产品名称", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[["产品类别",'-产品名称'], ['-伤害'], ['报告编码'], ['nunique'], '', ['报告编码合计']]),0,ori)) 


            nomal_menu2.add_command(label="-证号", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['-注册人备案人', '产品类别', '-注册证备案证'], ['-伤害'], ['报告编码'], ['nunique'], '', ['报告编码合计']]),0,ori)) 
            nomal_menu2.add_command(label="-批号", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['-注册人备案人', '产品类别', '-注册证备案证', '-批号'], ['-伤害'], ['报告编码'], ['nunique'],'', ['报告编码合计']]),0,ori)) 
            nomal_menu2.add_command(label="-型号", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['-注册人备案人', '产品类别', '-注册证备案证', '-型号'], ['-伤害'], ['报告编码'], ['nunique'],'', ['报告编码合计']]),0,ori)) 
            nomal_menu2.add_command(label="-规格", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['-注册人备案人', '产品类别', '-注册证备案证', '-规格'], ['-伤害'], ['报告编码'], ['nunique'], '', ['报告编码合计']]),0,ori)) 

            nomal_menu2.add_separator()
            nomal_menu2.add_command(label="-性别", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['-性别'], ['-伤害'], ['报告编码'], ['nunique'], '', ['报告编码合计']]),0,ori)) 
            nomal_menu2.add_command(label="-年龄段", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['年龄段'], ['-伤害'], ['报告编码'], ['nunique'],'', ['报告编码合计']]),0,ori)) 
   


        if '---器械数据规整---' in datacols:        
            nomal_menu2.add_separator()
            nomal_menu2.add_command(label="-监测机构（去重）", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread.drop_duplicates("报告编码"),[['-监测机构'], ['-伤害'], ['报告编码'], ['nunique'], {'报告季度': 'SMALL_count_mode', '-单位名称': 'SMALL_count_mode', '严重报告超时': sum}, ['报告编码合计',["报告编码严重伤害",'严重报告超时'],['报告编码合计','报告编码严重伤害']]]),0,ori)) 
            nomal_menu2.add_command(label="-报告单位（去重）", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread.drop_duplicates("报告编码"),[['-监测机构', '-单位名称'], ['-伤害'], ['报告编码'], ['nunique'], {'报告季度': 'SMALL_count_mode',  '严重报告超时': sum} ,['报告编码合计',["报告编码严重伤害",'严重报告超时'],['报告编码合计','报告编码严重伤害']]]),0,ori)) 

        if '---药品数据规整---' in datacols:
            nomal_menu2.add_command(label="-时隔", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['不良反应发生时间减用药开始时间'], [], ['报告编码'], ['nunique'], '', ['报告编码']]),0,ori)) 
            nomal_menu2.add_separator()
            nomal_menu2.add_command(label="原患疾病（扩行）", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(CLEAN_expand_rows(ori_owercount_easyread.copy(),";",["原患疾病"]),[['原患疾病'], [], ['报告编码'], ['nunique'], '', ['报告编码']]),0,ori)) 
            nomal_menu2.add_command(label="用药原因（扩行）", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(CLEAN_expand_rows(ori_owercount_easyread.copy(),";",["用药原因"]),[['用药原因'], [], ['报告编码'], ['nunique'], '', ['报告编码']]),0,ori)) 
            nomal_menu2.add_command(label="不良反应（扩行）", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(CLEAN_expand_rows(ori_owercount_easyread.copy(),"；",["不良反应名称（规整）"]),[['不良反应名称（规整）'], [], ['报告编码'], ['nunique'], '', ['报告编码']]),0,ori)) 
            nomal_menu2.add_separator()
        if '-监测机构'in datacols and '-单位名称' in datacols:
            nomal_menu2.add_command(label="-监测机构（去重）", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread.drop_duplicates("报告编码"),[['-监测机构'], ['-伤害'], ['报告编码'], ['nunique'], {'报告季度': 'SMALL_count_mode', '-单位名称': 'SMALL_count_mode', '报告超时': sum, '新的加严重的': sum}, ['报告编码合计',["报告编码严重",'新的加严重的','报告超时'],['报告编码合计','报告编码合计','报告编码合计']]]),0,ori)) 
            nomal_menu2.add_command(label="-报告单位（去重）", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread.drop_duplicates("报告编码"),[['-监测机构', '-单位名称'], ['-伤害'], ['报告编码'], ['nunique'], {'报告季度': 'SMALL_count_mode',  '报告超时': sum, '新的加严重的': sum} ,['报告编码合计',["报告编码严重",'新的加严重的','报告超时'],['报告编码合计','报告编码合计','报告编码合计']]]),0,ori)) 



          
    if '---器械数据规整---' in datacols or '---药品数据规整---' in datacols: 
        nomal_menu3 = tk.Menu(menu_bar, tearoff=0)  
        menu_bar.add_cascade(label="风险监测", menu=nomal_menu3)
        if '器械故障表现' in datacols and ('报告编码' in datacols or '报告表编码'):
            nomal_menu3.add_separator()        
            nomal_menu3.add_command(label="风险预警", command=lambda:TOOLS_keti(ori_owercount_easyread)) 

      
        if '---药品数据规整---' in datacols:
            nomal_menu3.add_separator()   
            nomal_menu3.add_command(label="新的不良反应（集合法）", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_get_new_GUI(ori_owercount_easyread),0,0))  
 
            nomal_menu3.add_command(label="新的不良反应（匹配库法）-（历史数据-证号）", command=lambda:TOOLS_analyze_products(ori_owercount_easyread,biaozhun['药品不良反应库'],biaozhun['药品重点关注库'],'证号'))       
            nomal_menu3.add_command(label="新的不良反应（匹配库法）-AI证号", command=lambda:TOOLS_analyze_products(ori_owercount_easyread,biaozhun['药品不良反应库-AI'],biaozhun['药品重点关注库'],'证号'))       
        
            
    
    stat_menu = tk.Menu(menu_bar, tearoff=0)  
    menu_bar.add_cascade(label="常用工具", menu=stat_menu) 
    stat_menu.add_command(label="ROR和PRR计算(df)", command=lambda:TOOLS_ROR_from_df_with_gui(ori_owercount_easyread.copy())) 
    stat_menu.add_command(label="ROR和PRR计算(db)", command=lambda:TOOLS_ROR_from_DB_GUI(ori_owercount_easyread))  

    stat_menu.add_separator()
    stat_menu.add_command(label="批量透视(df)", command=lambda:TOOLS_stat_all_gui(ori_owercount_easyread.copy())) 
    stat_menu.add_command(label="趋势分析(df)", command=lambda:TOOLS_trend_analysis_GUI(ori_owercount_easyread.copy()))        
    stat_menu.add_separator()   
    stat_menu.add_command(label="秩和检验(依据统计表)", command=lambda:TOOLS_rank_sum_test(ori_owercount_easyread.copy())) 
    stat_menu.add_command(label="卡方检验(依据统计表)", command=lambda:TOOLS_drug_reaction_CH2_create_gui(ori_owercount_easyread.copy()))  
    stat_menu.add_separator()  
    stat_menu.add_command(label="数据脱敏", command=lambda:PROGRAM_display_df_in_treeview(TOOLS_data_masking(ori_owercount_easyread),0,0))  
    stat_menu.add_command(label="转化字典列", command=lambda:PROGRAM_display_df_in_treeview(SMALL_expand_dict_like_columns(ori_owercount_easyread),0,0))   
      
    stat_menu.add_separator()  
    stat_menu.add_command(label="追加合并表格", command=lambda:TOOLS_merge_dataframes(ori_owercount_easyread))     
    stat_menu.add_command(label="删除尾行", command=lambda:PROGRAM_display_df_in_treeview(ori_owercount_easyread[:-1],0,0))      
    stat_menu.add_command(label="删除首行", command=lambda:PROGRAM_display_df_in_treeview(ori_owercount_easyread[1:],0,0))      


def SETTING_get_width():  
    column_widths = {  
        "评分说明": 800,  
        "该单位喜好上报的品种统计": 200,  
        "报告编码": 200,  
        "产品名称": 200,  
        "上报机构描述": 200,  
        "持有人处理描述": 200,  
        "该注册证编号/曾用注册证编号报告数量": 200,  
        "通用名称": 200,  
        "该批准文号报告数量": 200,  
        "上市许可持有人名称": 200,  
        "注册证编号/曾用注册证编号": 140,  
        "监测机构": 140,  
        "报告月份": 140,  
        "报告季度": 140,  
        "单位列表": 140,  
        "单位名称": 140,  
        "管理类别": 40,  
        "报告日期": 100,          
        "报告人": 50,         
        "报告表编码一般": 100,  
        "报告表编码严重": 100,      
        "报告表编码新的一般": 100,  
        "报告表编码新的严重": 100,  
        "报告表编码其他": 100,  
        "报告表编码合计": 100,  
        "一般": 40,  
        "严重": 40,      
        "新的一般": 55,  
        "新的严重": 55,    
        "合计": 50,          
        "信息":1000,                          
        "#0":40,
        "发生地":40,      
        "联系人":40, 
        "产品类别":40,
        "产品批号":60,      
        "型号":60, 
        "规格":60,         
        "-批号":60,      
        "-型号":60, 
        "-规格":60,         
        "曾用注册证编号上报":40, 
        "联系电话":40  
        
    }  
    return column_widths
    
def SETTING_exe(s_dict,methon_treeview,ori):
    #先处理报表类型。
    
    report_type=eval(str(s_dict["报表类型"]))


    
    if "grouped" in s_dict["报表类型"]:
        data_s=ori.copy()
        for i in report_type["grouped"]:
            mask=data_s[i].astype(str)==str(s_dict[i])
            data_s =data_s[mask].copy() # bao括的
        PROGRAM_display_df_in_treeview(data_s,methon_treeview,data_s)
        return 0	
        
    elif "group_sep" in s_dict["报表类型"]:
        data_s=ori.copy()        
        for i in report_type["group_sep"]:
            #print(i)
            #escaped_value = re.escape(str(s_dict[i]))
             
            mask=ori[i].str.contains(str(s_dict[i]),na=False)
            
            data_s =ori[mask].copy() # bao括的
        PROGRAM_display_df_in_treeview(data_s,methon_treeview,data_s)
        return 0	     

    
#######################################
#SQL 
def AAA_SQL():
    pass
######################################

  

#df转DB
def SQL_df_to_sqlite_db_with_gui(df, table_name='table1'):  
    # 创建Tkinter窗口实例  
    root = tk.Tk()  
    root.withdraw()  # 隐藏主窗口  
  
    # 弹出文件选择对话框让用户选择数据库文件  
    db_path = filedialog.asksaveasfilename(defaultextension=".db",  
                                           filetypes=[("SQLite Database Files", "*.db"), ("All Files", "*.*")])  
  
    # 如果用户选择了文件  
    if db_path:  
        try:  
            # 连接到SQLite数据库（如果文件不存在，它会被创建）  
            conn = sqlite3.connect(db_path)   
            # 显示正在处理的消息  
            print(f"Saving DataFrame to {db_path} as table {table_name}...")   
            df.to_sql(table_name, conn, if_exists='replace', index=False)  
            # 关闭数据库连接  
            conn.close()  
  
            # 显示完成消息  
            messagebox.showinfo(title="提示", message=f"\nDataFrame has been saved to {db_path} as table {table_name}.")  
  
        except sqlite3.Error as e:  
            messagebox.showinfo(title="提示", message=f"An error occurred: {e}")  
    else:  
        messagebox.showinfo(title="提示", message="No file selected.") 


#DB更新
def SQL_update_sqlite_db_with_df(df):  
    # 创建Tkinter窗口实例  
    root = tk.Tk()  
    root.withdraw()  # 隐藏主窗口  
  
    # 弹出文件选择对话框让用户选择数据库文件  
    db_path = filedialog.askopenfilename(title="Select SQLite Database", filetypes=[("SQLite Database Files", "*.db"), ("All Files", "*.*")])  
  
    # 如果用户选择了文件  
    if db_path:  
        try:  
            # 连接到SQLite数据库  
            conn = sqlite3.connect(db_path)  
  
            # 获取数据库第一个表的名称  
            cursor = conn.cursor()  
            cursor.execute("SELECT name FROM sqlite_master WHERE type='table' ORDER BY name LIMIT 1;")  
            table_name = cursor.fetchone()[0]  
  
            # 从数据库中读取现有数据的列名  
            db_columns_query = cursor.execute(f"PRAGMA table_info({table_name})")  
            db_columns = [column[1] for column in db_columns_query.fetchall()]  
  
            # 获取DataFrame的列名  
            df_columns = list(df.columns)  
  
            # 检查列差异  
            columns_in_df_not_in_db = [col for col in df_columns if col not in db_columns]  
            columns_in_db_not_in_df = [col for col in db_columns if col not in df_columns]  
  
            # 打印列差异  
            if columns_in_df_not_in_db:  
                print(f"The following columns are in the DataFrame but not in the database table: {columns_in_df_not_in_db}")  
            if columns_in_db_not_in_df:  
                print(f"The following columns are in the database table but not in the DataFrame: {columns_in_db_not_in_df}")  
  
            # 如果列名完全相同，则追加数据并进行去重  
            if not (columns_in_df_not_in_db or columns_in_db_not_in_df):  
                # 将DataFrame数据追加到数据库表  
                df.to_sql(table_name, conn, if_exists='append', index=False)  
  
 
                # 提交更改  
                conn.commit()  
  
                messagebox.showinfo(title="提示", message=f"Data has been appendedin table {table_name} in {db_path}.")  
  
            # 关闭数据库连接  
            conn.close()  
  
        except sqlite3.Error as e:  
            messagebox.showinfo(title="提示", message=f"An error occurred: {e}")  
            if conn:  
                conn.close()  
    else:  
        messagebox.showinfo(title="提示", message="No file selected.")  


#SQL查询
def SQL_create_query_gui(df,methon=None,conn=None):  
    #methon用于输入一个文本框
    def add_query_condition(event=None):  
        # 获取选择的列名和输入的查询内容  
        column_name = column_dropdown.get()  
        query_value = query_entry.get()  
      
      
        # 如果两者都不为空，则构建查询条件并添加到文本框  
        if column_name and query_value:  
            
            
            # 检查是否输入了算数运算符  
            match = re.match(r"([<>!=]=?)\s*(\d+)", query_value)  
            if match:  
                operator, value = match.groups()  
          
                # 将用户输入的运算符转换为SQL运算符  
                sql_operator = {  
                    "=": "=",  
                    "!=": "<>",  
                    ">": ">",  
                    "<": "<",  
                    ">=": ">=",  
                    "<=": "<="  
                }.get(operator)  
          
                if sql_operator:  
                    condition = f"{column_name} {sql_operator} {value}"  
                else:  
                    messagebox.showerror("Error", "Invalid operator. Please use >, <, >=, <=, !=, or =.")  
                    return 
            
            
            # 检查是否输入了日期范围  
            elif query_value.startswith('time(') and query_value.endswith(')'):  
                try:  
                    # 提取日期范围  
                    start, end = query_value[5:-1].split('-')  
                    # 格式化日期为YYYY-MM-DD  
                    start_date = f"{start[:4]}-{start[4:6]}-{start[6:]}"  
                    end_date = f"{end[:4]}-{end[4:6]}-{end[6:]}"  
                    # 构建日期范围查询条件  
                    condition = f"{column_name} BETWEEN '{start_date}' AND '{end_date}'"  
                except ValueError:  
                    # 如果日期格式不正确，显示错误消息  
                    messagebox.showerror("Error", "Invalid date range format. Please use 'time(YYYYMMDD-YYYYMMDD)'.")  
                    return  
            elif query_value.startswith("df['") and query_value.endswith("']"):
                
                        # 检查是否提供了 DataFrame 和列名  
                if df is None or column_name not in df.columns:  
                    messagebox.showerror("Error", "DataFrame or column name not provided.")  
                    return #CDR2023220233983732874313728 
                
                inner_column_name = query_value[4:-2]
                values_to_match= df[inner_column_name].unique().tolist()
                if len(values_to_match) >900:  
                    print("元素过多，超过了900个，可能会导致性能问题.") 
                    # 构建 SQL IN 子句  
                    in_clause = ', '.join([f"'{value}'" for value in values_to_match])  
                    condition = f"{column_name} IN ({in_clause})"   
                else:    
                    conditions = [f"{column_name} LIKE '%{value}%'" for value in values_to_match]  
                    condition = " OR ".join(conditions)  
                      
            else:  
                # 构建查询条件  
                values_to_match = query_value.split("|")      
                conditions = [f"{column_name} LIKE '%{value}%'" for value in values_to_match]  
                condition = " OR ".join(conditions)  
      
        # 检查文本框是否为空，如果不为空，则添加AND连接符  
        if sql_text.get("1.0", tk.END).strip():  
            sql_text.insert(tk.END, f" AND ({condition})")  
        else:  
            sql_text.insert(tk.END, f"WHERE ({condition})")  
      
        # 清空输入框以便输入下一个条件  
        query_entry.delete(0, tk.END)  
      
    def execute_combined_query():  
        
        # 获取完整的查询语句  
        query = f"SELECT * FROM {first_table} {sql_text.get('1.0', tk.END)}"  
      
        try:  
            data = pd.read_sql_query(query, conn)  
            PROGRAM_display_df_in_treeview(data,0,0)  
        except sqlite3.Error as e:  
            messagebox.showerror("Error", f"An error occurred: {e.args[0]}")  
    def return_query():  
        methon.delete("1.0",  tk.END)   
        s = sql_text.get('1.0', tk.END)  
        # 查找第一个 "WHERE" 的位置  
        index = s.find("WHERE")  
          
        # 如果找到了 "WHERE"  
        if index != -1:  
            # 使用切片来删除第一个 "WHERE"  
            s = s[:index] + s[index + len("WHERE"):] 
        methon.insert(tk.END, s.strip()) 
        root.destroy()  

              
    
    def populate_columns():  
        # 获取第一个表的列名并填充到下拉菜单  
        cursor = conn.cursor()  
        cursor.execute(f"PRAGMA table_info({first_table})")  
        columns = [row[1] for row in cursor.fetchall()]  
        column_dropdown['values'] = columns  

    def on_closing():  
        conn.close()  
        root.destroy()   

    # 创建主窗口  
    root = tk.Tk()  
    root.title("SQL查询工具")  
    root.protocol("WM_DELETE_WINDOW", on_closing)  # 设置关闭窗口时的回调函数  
    sw = root.winfo_screenwidth()  
    sh = root.winfo_screenheight()  
    ww = 470  # 窗口宽度  
    wh = 250  # 窗口高度  
    x = (sw - ww) // 2  
    y = (sh - wh) // 2  
    root.geometry(f"{ww}x{wh}+{x}+{y}")  
    if conn==None:
        file_path = filedialog.askopenfilename(filetypes=[("SQLite Database Files", "*.db"), ("All Files", "*.*")])  
        conn = sqlite3.connect(file_path) 
    root.lift()
    root.attributes("-topmost", True)
    root.attributes("-topmost", False)
    # 获取数据库的第一个表名  
    cursor = conn.cursor()  
    cursor.execute("SELECT name FROM sqlite_master WHERE type='table' ORDER BY name LIMIT 1;")  
    global first_table  # 使用全局变量以便在多个函数间共享  
    first_table = cursor.fetchone()[0]  
  
    # 创建带滚动条的文本框用于显示构建的查询  
    sql_text = scrolledtext.ScrolledText(root, width=40, height=10, wrap=tk.WORD)  
    sql_text.grid(row=0, column=0, columnspan=3, padx=2, pady=2, sticky=tk.W+tk.E)  
  
    # 创建列名下拉菜单  
    column_label = tk.Label(root, text="请选择列:")  
    column_label.grid(row=2, column=0, padx=10, sticky=tk.W)  
    column_dropdown = ttk.Combobox(root)  
    column_dropdown.grid(row=2, column=1, padx=10)  
    populate_columns()  
  
    # 创建输入框用于输入查询内容  
    query_entry = tk.Entry(root)  
    query_entry.grid(row=2, column=2, padx=10)  
  
    # 创建“增加”按钮  
    add_button = tk.Button(root, text="添加条件", command=add_query_condition)  
    add_button.grid(row=3, column=2, padx=10)  
  
    # 创建执行按钮  
    execute_button = tk.Button(root, text="执行", command=execute_combined_query)  
    execute_button.grid(row=3, column=0, columnspan=3, padx=10, pady=10)  
    
    if methon!=None:
        return_button = tk.Button(root, text="返回", command=return_query)  
        return_button.grid(row=3, column=1, columnspan=3, padx=10, pady=10)  
            

    # 运行主循环  
    root.mainloop()  

#######################################
#改写和注册部分函数和其他小函数 
def AAA_Small():
    pass
######################################

def SMALL_merge_dataframes(df1, df1_col, df2_path=None, df2_col='PT'):  
	#药品预制-加SOC（PT扩行-MedDra法） 
    # 如果df2_path为None，弹出文件选择对话框让用户选择文件  
    if df2_path is None:  
        root = tk.Tk()  
        root.withdraw()  # 隐藏主窗口  
        df2_path = filedialog.askopenfilename()  # 弹出文件选择对话框  
      
    # 读取df2文件（如果提供了路径） 
     
    if df2_path: 
        script_dir = os.path.dirname(os.path.abspath(__file__))  
        config_file = os.path.join(script_dir, df2_path) 
        df2 = pd.read_excel(config_file)  # 假设df2是CSV文件，如果是其他格式请相应修改  
    
    else:  
        raise ValueError("No path provided for df2 and no file selected.")  
        
    try:  
        df1 = df1.drop(['code', 'Chinese', '级别', 'PT','HLT', 'HLGT', 'SOC', '主SOC'], axis=1)  
    except:  
        pass
          
    # 使用merge函数将df2拼接到df1上（基于指定的列名）  
    merged_df = pd.merge(df1, df2, left_on=df1_col, right_on=df2_col, how='left')  
    try:  
        merged_df['SOC']= merged_df['SOC'].fillna('其他（未规整）')  
    except:  
        pass 
    return merged_df  



def SMALL_get_dataframe_hash(df):  
    """处理DataFrame并返回其哈希值"""  
    columns_sorted = sorted(df.columns)  
    df_sorted = df[columns_sorted]  
    return hashlib.md5(df_sorted.to_csv(index=False).encode('utf-8')).hexdigest()  
  
def SMALL_compare_excel_files():  
    """让用户选择两个Excel文件并比较它们的内容"""   
    # 创建一个隐藏的Tkinter窗口以便使用filedialog模块  
    root = tk.Tk()  
    root.withdraw()  
  
    # 弹出对话框让用户选择第一个Excel文件  
    file_path1 = filedialog.askopenfilename(title="选择第一个Excel文件", filetypes=[("Excel files", "*.xlsx;*.xls")])  
    if not file_path1:  # 如果用户取消了选择，则退出程序  
        return  
  
    # 弹出对话框让用户选择第二个Excel文件  
    file_path2 = filedialog.askopenfilename(title="选择第二个Excel文件", filetypes=[("Excel files", "*.xlsx;*.xls")])  
    if not file_path2:  # 如果用户取消了选择，则退出程序  
        return  
  
    # 读取Excel文件并转换为DataFrame  
    df1 = pd.read_excel(file_path1).reset_index(drop=True) 
    print(df1)
    df2 = pd.read_excel(file_path2).reset_index(drop=True)   
    print(df2)
    # 计算DataFrame的哈希值并比较  
    hash1 = SMALL_get_dataframe_hash(df1)  
    hash2 = SMALL_get_dataframe_hash(df2)  
  
    # 打印比较结果  
    if hash1 == hash2:  
        print("两个Excel文件内容相同。")  
    else:  
        print("两个Excel文件内容不同。")  
  



def SMALL_apply_operation(df, mycols,col_name ,operation):  
    """  
    对DataFrame中的两列应用任意运算。  
  
    参数:  
        df (pd.DataFrame): 输入的DataFrame。  
        col1 (str): 第一列的名称。  
        col2 (str): 第二列的名称。  
        operation (str): 运算式子，例如 'a+b' 或 'a-b'。  
  
    返回:  
        pd.Series: 运算结果。  
    """ 
    def check_time_format(df, column_name):  
        # 检查列中的数据是否可以转换为时间格式  
        try:  
            datetime.strptime(df[column_name][0], '%Y-%m-%d %H:%M:%S')  
            return True  
        except:  
            try:  
                datetime.strptime(df[column_name][0], '%Y-%m-%d')  
                return True  
            except:  
                return False       
    col1, col2 = mycols  
      
    # 检测输入列是否为时间格式  
    if check_time_format(df, col1) and  check_time_format(df, col2) :  
        time1 = pd.to_datetime(df[col1])    
        time2 = pd.to_datetime(df[col2])  
        is_time=True
    else:
        is_time=False  
    # 根据运算式子执行相应的计算  

    if is_time:  
        # 如果是时间运算，计算时间差（以天为单位）
        if operation=="a-b":  
            time_diff = (time1 - time2).dt.days  
        elif operation=="b-a":  
            time_diff = (time2 - time1).dt.days 
        else:
            raise ValueError(f"时间只支持相减。") 
        df[col_name]=time_diff                  
    else:  
        # 如果是非时间运算，使用eval()函数执行运算  
        try:  
            df[col_name] = eval(operation, {"__builtins__": None}, {'a': df[col1], 'b': df[col2]})  
        except Exception as e:  
            raise ValueError(f"无法执行运算：{e}")  
      
    return df



    
def SMALL_find_based_on_expression(df,expression):  
    #条件查找，范例："(df['a']>5) | ((3<=df['a']) & (df['a']<10)) | (df['d']!=6) | (df['e']==7) | (df['f']=='r')"  
    try:  
        mask = pd.eval(expression, engine='python', parser='pandas', local_dict={'df': df})  
    except Exception as e:  
        raise ValueError(f"无法解析表达式: {e}")  
          
    # 返回符合条件的行  
    return df[mask]  
  

  
def SMALL_assign_value_based_on_expression(df, target_column, expression, value):  
    """  
    根据给定的表达式给DataFrame的target_column列赋值value。  
    value 可以是数字、文本或者一个关于df的表达式（如 "df['a']" 或 "df['a']/12"）。  
      
    参数:  
        df (pd.DataFrame): 输入的DataFrame。  
        target_column (str): 需要赋值的列名。  
        expression (str): 用于生成布尔掩码的表达式字符串。  
        value: 需要赋给满足表达式条件的行的值，可以是数字、字符串或者关于df的表达式。  
  
    返回:  
        pd.DataFrame: 更新后的DataFrame。  
    """  
    try:  
        # 使用pd.eval来评估表达式并生成布尔掩码  
        mask = pd.eval(expression, engine='python', parser='pandas', local_dict={'df': df})  
        # 确保mask是布尔类型的Series  
        if not isinstance(mask, pd.Series) or not mask.dtype == bool:  
            raise ValueError("表达式返回的不是有效的布尔序列")  
  
        # 如果value是一个字符串并且包含"df", 则尝试将其解析为一个关于df的表达式  
        if isinstance(value, str) and "df" in value:  
            try:  
                value_series = pd.eval(value, engine='python', parser='pandas', local_dict={'df': df})  
                if not isinstance(value_series, pd.Series):  
                    raise ValueError("value 表达式返回的不是有效的pandas Series")  
            except Exception as e:  
                raise ValueError(f"无法解析value表达式: {e}")  
        # 如果value不是字符串或者不包含"df", 则直接将其视为一个常量值  
        else:  
            value_series = pd.Series(value, index=df.index)  # 创建一个常量Series，以便可以广播到目标列  
  
        # 根据掩码给目标列赋值  
        df.loc[mask, target_column] = value_series  
        return df  
  
    except Exception as e:  
        raise ValueError(f"无法解析表达式或赋值: {e}")  

   


  

    
def SMALL_last_non_null_value(df, columns_list, new_column_name):  
    """  
    在DataFrame df中创建一个新列，其值为columns_list中最后一个非空值。  
      
    参数:  
        df (pd.DataFrame): 输入的DataFrame。  
        columns_list (list of str): 要检查的列名列表，按顺序。  
        new_column_name (str): 新列的名称。  
          
    返回:  
        pd.DataFrame: 包含新列的原始DataFrame。  
    """  
      
    # 遍历每一行，找到最后一个非空值  
    def get_last_non_null(row):  
        for col in reversed(columns_list):  # 从后向前遍历列  
            if pd.notna(row[col]):  
                return row[col]  
        return np.nan  # 如果没有非空值，则返回NaN  
  
    # 应用函数并创建新列  
    df[new_column_name] = df.apply(get_last_non_null, axis=1)  
      
    return df  


def SMALL_calculate_ratios(df, list1, list2):  
    # 确保输入列表长度相等  
    if len(list1) != len(list2):  
        raise ValueError("两个列表的长度必须相等")  
  
    # 遍历列表中的列名，并计算比值  
    for i in range(len(list1)):  
        col_name1 = list1[i]  
        col_name2 = list2[i]  
  
        # 确保列名在DataFrame中存在  
        if col_name1 not in df.columns or col_name2 not in df.columns:  
            raise ValueError(f"列名 {col_name1} 或 {col_name2} 不在DataFrame中")  
  
        # 计算比值并保留两位小数  
        df[f"{col_name1}_{col_name2}_ratio"] = df[col_name1] / df[col_name2]*100  
        df[f"{col_name1}_{col_name2}_ratio"] = df[f"{col_name1}_{col_name2}_ratio"].round(2)  
  
    return df 

def SMALL_expand_dict_like_columns(df):    
    def format_dict_like_values(value):    
        if isinstance(value, dict):    
            # 使用列表推导式生成格式化后的键值对字符串，然后用' '.join连接  
            formatted_items = [f"{k}（{v}）" for k, v in value.items()]  
            # 使用'、'.join连接字符串，并使用rstrip移除尾部的'、'  
            return '、'.join(formatted_items).rstrip('、')  
        elif isinstance(value, str) and value.startswith('{') and value.endswith('}'):    
            try:    
                data_dict = eval(value)  # 注意：使用eval有风险，应确保数据安全  
                if isinstance(data_dict, dict):    
                    formatted_items = [f"{k}（{v}）" for k, v in data_dict.items()]  
                    return '、'.join(formatted_items).rstrip('、')  
            except:    
                pass    
        return value  # 如果不是字典或类似字典的字符串，则返回原始值    
        
    df_transformed = df.copy()    
    for column in df_transformed.columns:    
        df_transformed[column] = df_transformed[column].apply(format_dict_like_values)    
    return df_transformed  
      
def SMALL_add_composition_ratio(df, column_name, new_column_name='构成比'):  
    """  
    在原始DataFrame（除最后一行合计外）中添加指定列的构成比。  
    这里假设column_name列已经包含了频数，且最后一行是合计。  
      
    参数:  
        df (pd.DataFrame): 输入的DataFrame。  
        column_name (str): 需要计算构成比的列名，该列应包含频数（除最后一行）。  
        new_column_name (str): 新添加的构成比列名，默认为'构成比'。  
          
    返回:  
        pd.DataFrame: 原始DataFrame的拷贝（除最后一行外），其中包含一个新列用于存储构成比（保留两位小数）。  
    """  
    # 检查列名是否在DataFrame中  
    if column_name not in df.columns:  
        raise ValueError(f"Column '{column_name}' not found in DataFrame.")  
      
    # 获取最后一行的合计值，并创建一个不包含最后一行的DataFrame拷贝  
    total = df.iloc[-1][column_name]  

      
    # 计算构成比（以百分比表示）  
    composition_ratio = (df[column_name] / total) * 100  
      
    # 在DataFrame拷贝中添加新列，存储构成比  
    df[new_column_name] = composition_ratio.round(2)  
      
    return df.sort_values('构成比', ascending=False)
          
def SMALL_count_mode(series, mode):  
    """  
    对给定的pandas Series进行计数，并根据mode参数指定的分隔符拆分复合症状并计算每个症状的计数。  
      
    参数:  
        series (pandas.Series): 输入的pandas Series对象，包含要计数的数据。  
        mode (str): 指定用于拆分复合症状的分隔符，可以是单个字符或多个字符。  
      
    返回:  
        dict: 拆分后并计算每个症状计数的新字典。  
    """  
    result = series.value_counts().to_dict()  
    filtered_result = {k: v for k, v in result.items() if v > 0}  
    if mode=="":
        return filtered_result
      
    new_dict = {}  
    for key, value in filtered_result.items():  
        symptoms = re.split(mode, key)  # 使用正则表达式根据指定的分隔符进行拆分  
        for symptom in symptoms:  
            symptom = symptom.strip()  
            if symptom:  # 检查症状是否为非空字符串  
                if symptom in new_dict:  
                    new_dict[symptom] += value  
                else:  
                    new_dict[symptom] = value  
      
    return new_dict
    
def SMALL_get_list(input_str, df, *args):  
    input_str = str(input_str)  
      
    if pd.isnull(input_str):  
        return []  
  
    result_list = []  
      
    if ("use(" in input_str) and (").file"  in input_str):  
        pattern = r"[(](.*?)[)]"  
        matches = re.findall(pattern, input_str)  
        result_list = df[matches[0]].astype(str).tolist()  
    else:  
        result_list = [input_str]  
      
    # 合并重复的字符串操作  
    result_list = ",".join(result_list)  
    result_list = result_list.replace("（严重）", "").replace("（一般）", "")  
    for separator in ["┋", ";", "；", "、", "，", ","]:  
        result_list = result_list.replace(separator, ",")  
    #print(result_list)  
    result_list = result_list.split(",") 
    #result_list.sort()  
      
    return result_list    

def SMALL_save_dict(data):  
    """保存文件"""  
    file_path_flhz = filedialog.asksaveasfilename(  
        title="保存文件",  
        initialfile="排序后的原始数据",  
        defaultextension=".xlsx",  
        filetypes=[("Excel 工作簿", "*.xlsx"), ("Excel 97-2003 工作簿", "*.xls")],  
    )  
    if not file_path_flhz:  
        return  # 如果用户点击取消，则退出函数
        
    if "详细描述T" in data.columns:
        data["详细描述T"]=data["详细描述T"].astype(str)

    if "报告编码" in data.columns:
        data["报告编码"]=data["报告编码"].astype(str)     
        
        
    try:  
        with pd.ExcelWriter(file_path_flhz, engine="xlsxwriter") as writer:  
            data.to_excel(writer, sheet_name="导出的数据", index=False)  
        messagebox.showinfo(title="提示", message="文件写入成功。")  
    except Exception as e:  
        messagebox.showerror("错误", f"保存文件时出错: {e}")


def SMALL_read_csv_files():  
    """打开文件选择对话框，读取CSV文件，并返回合并后的数据。"""  
    file_types = [("CSV files", "*.csv"), ("All files", "*.*")]  
    root = filedialog.Tk()  # 需要一个Tk root实例来运行filedialog  
    root.withdraw()  # 隐藏Tk窗口  
    file_paths = filedialog.askopenfilenames(title="请选择CSV文件", filetypes=file_types)  
      
    if not file_paths:  
        return None  # 如果用户取消了对话框，直接返回None  
  
    # 读取并合并CSV文件  
    all_data = []  
    for file_path in file_paths:  
        df = pd.read_csv(file_path)  
        all_data.append(df)  
      
    combined_data = pd.concat(all_data, ignore_index=True) if len(all_data) > 1 else all_data[0]  
      
    # 尝试删除名为"Unnamed"的列（如果存在）  
    try:  
        combined_data = combined_data.loc[:, ~combined_data.columns.str.contains("^Unnamed")]  
    except:  
        pass  
      
    return combined_data.reset_index(drop=True)  


def SMALL_read_excel_files():  
    """打开文件选择对话框，读取Excel文件，并返回合并后的数据。"""  
    file_types = [("Excel files", "*.xls;*.xlsx"), ("All files", "*.*")]  
    file_paths = filedialog.askopenfilenames(title="请选择Excel文件", filetypes=file_types)  
    if not file_paths:  
        return None  # 如果用户取消了对话框，直接返回None  
        
  
    # 读取并合并Excel文件  
    all_data = []  
    for file_path in file_paths:  
        df = pd.read_excel(file_path)  
        all_data.append(df)  
    combined_data = pd.concat(all_data, ignore_index=True) if len(all_data) > 1 else all_data[0]  
    try:
        combined_data=combined_data.loc[ : , ~combined_data.columns.str.contains("^Unnamed")]
    except:
        pass
    return combined_data.reset_index(drop=True)  



def SMALL_pre_process_dataframe(df):  
    """  
    处理DataFrame中的非数字列和缺失值。  
      
    参数:  
        df (pd.DataFrame): 输入的DataFrame。  
          
    返回:  
        pd.DataFrame: 处理后的DataFrame。  
    """  
    # 复制原始DataFrame以防修改原始数据  
    df_processed = df.copy()  
      
    # 遍历每一列  
    for col in df_processed.columns:  
        # 检查列的数据类型  
        if pd.api.types.is_string_dtype(df_processed[col]):  
            # 如果是字符串类型，尝试将其转换为类别型并进行编码  
            df_processed[col] = df_processed[col].astype('category').cat.codes  
        elif pd.api.types.is_numeric_dtype(df_processed[col]):  
            # 如果是数字类型但有缺失值，使用均值填充  
            df_processed[col].fillna(df_processed[col].mean(), inplace=True)  
        else:  
            # 对于其他类型，可以考虑删除列或转换为适当的类型（根据具体情况）  
            # 这里我们选择删除列作为示例（根据实际情况调整）  
            df_processed.drop(columns=col, inplace=True)  
            print(f"Column '{col}' has been dropped because its data type is not supported.")  
      
    # 额外处理：检查是否有任何非数字值（可能是由于之前的步骤中未处理的特殊情况）  
    # 如果有，将这些值替换为NaN，并使用该列的均值填充  
    for col in df_processed.columns:  
        df_processed[col] = pd.to_numeric(df_processed[col], errors='coerce')  
        if df_processed[col].isnull().sum() > 0:  
            df_processed[col].fillna(df_processed[col].mean(), inplace=True)  
      
    return df_processed  
    
def SMALL_align_values_to_most_frequent(df, group_col, cols_to_align):  
    """  
    对于DataFrame中的每个group_col的唯一值，找到cols_to_align中对应列的最频繁出现的值，  
    并将该值赋给所有具有相同group_col值的行。不在cols_to_align中的列保持不变。  
      
    参数:  
        df (pd.DataFrame): 输入的DataFrame。  
        group_col (str): 用于分组的列名。  
        cols_to_align (list of str): 需要对齐的列名列表。  
          
    返回:  
        pd.DataFrame: 更新后的DataFrame，包含所有原始列，且列顺序不变。  
    """  
    # 复制原始DataFrame以保持其不变  
    df_aligned = df.copy()  
      
    # 遍历需要对齐的列  
    for col in cols_to_align:  
        # 使用groupby找到每个group_col值的最频繁出现的值  
        mode_series = df.groupby(group_col)[col].transform(lambda x: x.mode()[0])  
        # 将结果对齐到原始DataFrame的相应列  
        df_aligned[col] = mode_series  
      
    return df_aligned  
    
def SMALL_add_count_and_ratio(df, a, agg_col=None, methon=None):  
    """  
    函数说明：  
    该函数接受一个pandas DataFrame（df）、列名（a）、聚合列名（agg_col）以及聚合方法（methon）作为输入。  
    使用groupby方法对列a进行分组，并根据agg_col和methon计算每个组的聚合值。  
    接下来，计算每个组的构成比，并将其转换为百分数形式，保留两位小数。  
    最后，按构成比的大小进行排序，并返回一个新的DataFrame，  
    其中包含列a的唯一值、每个唯一值的计数或聚合值以及构成比。  
  
    参数：  
    df：pandas DataFrame，包含要进行分组和聚合的数据。  
    a：字符串或字符串列表，指定要进行分组的列名。  
    agg_col：字符串，可选，指定要进行聚合的列名。如果为None，则对整个组进行聚合。  
    methon：字符串，指定聚合方法（例如，'count'、'sum'、'nunique'等）。  
  
    返回值：  
    pandas DataFrame，包含列a的唯一值、每个唯一值的计数或聚合值以及构成比（百分数形式，保留两位小数），并按构成比大小排序。  
    """  
  
  
    # 对列a进行分组  
    if agg_col is not None:  
        grouped = df.groupby(a)[agg_col]  
    else:  
        grouped = df.groupby(a).size().reset_index(name='计数')  
      
    # 根据methon计算聚合值  
    if methon == 'count':  
        grouped = grouped.count().reset_index(name='计数')  
    elif methon == 'sum':  
        grouped = grouped.sum().reset_index(name='计数')  
    elif methon == 'nunique':  
        grouped = grouped.nunique().reset_index(name='计数')  
  
    # 计算构成比并转换为百分数形式，保留两位小数  
    total_count = grouped['计数'].sum()  
    grouped['构成比'] = (grouped['计数'] / total_count) * 100  
    grouped['构成比'] = grouped['构成比'].round(2)  
  
    # 按构成比大小进行排序  
    grouped = grouped.sort_values('构成比', ascending=False)  
  
    # 添加报表类型列（注意：此部分逻辑可能需要根据实际需求进行调整）  
    lst = []  
    if isinstance(a, str):  
        lst.append(a)  
    elif isinstance(a, list):  
        lst.extend(a)  
    grouped["报表类型"] = "{'grouped':" + str(lst) + "}"  
  
    return grouped
    
def SMALL_add_count_and_ratio_exp(df, a,sep):  
    """  
    当列扩行透视 
    """  
    df=df.copy()
    df=CLEAN_expand_rows(df,sep,[a])
    df=df.reset_index(drop=True)
    # 对列a进行分组并计算计数  
    grouped = df.groupby(a).size().reset_index(name='计数')  
      
    # 计算构成比并转换为百分数形式，保留两位小数  
    total_count = grouped['计数'].sum()  
    grouped['构成比'] = (grouped['计数'] / total_count) * 100  
    grouped['构成比'] = grouped['构成比'].round(2)  
      
    # 按构成比大小进行排序  
    grouped = grouped.sort_values('构成比', ascending=False)  
    
    lst = []  
    if isinstance(a, str):  
        lst.append(a)  
    elif isinstance(a, list):  
        lst.extend(a)
    
    grouped["报表类型"]="{'group_sep':"+str(lst)+str("}")       
    return grouped    
#######################################
#数据清洗函数 
def AAA_Clean():
    pass
###################################### 
def CLEAN_fill_column_c_based_on_a_and_b(df, col_a, col_b, col_c):  
    """  
    对于DataFrame df中col_a列具有相同值的行，将相关行中col_b列的值（除了所在行的值）  
    填写到col_c列中，以逗号分隔。  
      
    参数:  
    df (DataFrame): 输入的DataFrame。  
    col_a (str): 用于比较的列名。  
    col_b (str): 需要被复制值的列名。  
    col_c (str): 目标列名，用于存储合并后的值。  
      
    返回值:  
    DataFrame: 处理后的DataFrame。  
    """  
    # 创建一个空的list用于存储结果  
    df[col_c] = ''  
      
    # 对col_a列中的每个唯一值进行处理  
    for value_a in df[col_a].unique():  
        # 筛选出col_a列等于当前唯一值的所有行  
        sub_df = df[df[col_a] == value_a]  
          
        # 如果筛选出的子DataFrame有多于一行  
        if len(sub_df) > 1:  
            # 遍历子DataFrame的每一行  
            for index, row in sub_df.iterrows():  
                # 从子DataFrame中去掉当前行，然后获取col_b列的值  
                other_values_b = sub_df.drop(index)[col_b].tolist()  
                # 将这些值用逗号连接成一个字符串，并赋值给当前行的col_c列  
                df.at[index, col_c] = ','.join(map(str, other_values_b))  
      
    return df  

def CLEAN_easystat(df, guize,method):  
    """  
    根据operation_log中的记录，在df上复现操作。  
    """
    mark_column_name=guize['结果列'][0]
    df[mark_column_name]=""

    guize=guize.fillna("XXX---XXXXXXXXX---XXXXX")
    allx=len(guize)
    for ids,cols in guize.iterrows(): 
        print(ids, allx)  
        selected_columns=str(cols["查找位置"]).split("|")
        include_str=cols["值"]
        exclude_str=cols['排除值']
        predefined_column=cols['适用范围列']
        predefined_objects=cols['适用范围']
        
        mark_value=str(cols["值"]).split("|")[0]
        
        CLEAN_filter_and_mark_rows(df, selected_columns, include_str, exclude_str, predefined_column, predefined_objects, mark_column_name, mark_value, method) 
    df[mark_column_name] = df[mark_column_name].replace("", "其他") 
    df[mark_column_name] = df[mark_column_name].str.replace("nan;", "")
    df[mark_column_name] = df[mark_column_name].str.replace("nan", "其他") 
    return df

def CLEAN_replay_operations(df, operation_log):  
    """  
    根据operation_log中的记录，在df上复现操作。  
    """ 
    #operation_log["hash_auto"]="" 
    
    
    for ids,operation in operation_log.iterrows():  
        
        try:

            method = operation["方法"]  
            columns = ast.literal_eval(operation["作用列"] )
            selected_columns=columns
            params = eval(str(operation["参数"]))   
             
            
              
            if method == "复制合并":  
                new_column_name = params["new_column_name"] 

                separator = params["separator"]  
                if new_column_name in df.columns:  
                    print("错误：新列名与现有列名重复")  
                    continue  # 跳过当前操作，继续下一个  
                df[new_column_name] = df[columns].apply(lambda row: separator.join(row.astype(str)), axis=1)  
                print(f"已合并列，并创建新列：{new_column_name}")  

            elif method == "两列运算":  
                new_column_name = params["new_column_name"] 

                separator = params["separator"]  

                if len(columns)!=2: 
                    print("错误", "请选择两个数据列。")  
                    return  
                df = SMALL_apply_operation(df,columns,new_column_name,separator)                
                print(f"已开展相关运算：{new_column_name}")  
                
            elif method == "众数对齐" :  
                group_col = params["group_col"] # 获取目标格式并去除首尾空格  
                if not group_col:  
                    print("错误，对齐源列不能为空")  
                    return  

                try:  
                    df=SMALL_align_values_to_most_frequent(df,group_col,columns)             
                except ValueError as e:  
                    messagebox.showinfo("错误", f"无法完成任务，原因：{str(e)}")  
                    return
      
                  
            elif method == "重命名列":  
                if len(columns) > 1:  
                    print("错误：重命名操作只能选择一列")  
                    continue  # 跳过当前操作，继续下一个  
                new_column_name = params["new_column_name"]  
                if new_column_name in df.columns and new_column_name != columns[0]:  
                    print("错误：新列名与现有列名重复")  
                    continue  # 跳过当前操作，继续下一个  
                df.rename(columns={columns[0]: new_column_name}, inplace=True)  
                print(f"已将列 {columns[0]} 重命名为 {new_column_name}")  
                
            elif method == "最后非空":  

                new_column_name = params["new_column_name"]  
                if new_column_name in df.columns and new_column_name != columns[0]:  
                    print("错误：新列名与现有列名重复")  
                    continue  # 跳过当前操作，继续下一个  
                df=SMALL_last_non_null_value(df,columns,new_column_name) 
                print(f"已设置最后非空为 {new_column_name}")  
                            
                
            elif method == "并用追加":  
      
                new_column_name = params["new_column_name"]  
                new_column_valuex = params["new_column_valuex"]  
                df=CLEAN_fill_column_c_based_on_a_and_b(df,new_column_name,columns[0],new_column_valuex)     
                
            elif method == "新建一列":  
                if len(columns) > 1:  
                    print("错误：重命名操作只能选择一列")  
                    continue  # 跳过当前操作，继续下一个  
                new_column_name = params["new_column_name"]
                new_column_valuex=params["new_column_valuex"]  
                if new_column_name in df.columns and new_column_name != columns[0]:  
                    print("错误：新列名与现有列名重复")  
                    continue  # 跳过当前操作，继续下一个  
                df[new_column_name]=new_column_valuex 
                print(f"已新建{new_column_name}")              
                

            elif method == "替换字符":  
                find_text = params["find_text"]  
                replace_text = params["replace_text"]  
                if not find_text:  
                    print("错误：请输入要查找的字符")  
                    continue  
                for col in columns:  
                    df[col] = df[col].astype(str).str.replace(find_text, replace_text)  
                print(f"已在列 {columns} 中替换字符：{find_text} -> {replace_text}")  

            elif method == "条件赋值":  
                compare_value = params["compare_value"]  
                G_value = params["G_value"]  

                if len(columns)!=1:  
                    print("错误", "请选择1个数据列。")  
                    return               
                df=SMALL_assign_value_based_on_expression(df, columns[0], compare_value, G_value) 
                print(f"已完整条件赋值..") 

                  
            elif method == "填充空值":  
                fill_value = params["fill_value"]  
                if not fill_value:  
                    print("错误：填充值不能为空")  
                    continue  
                try:  
                    fill_value = df[columns[0]].dtype.type(fill_value)  
                except ValueError:  
                    print("错误：填充值与列的数据类型不匹配")  
                    continue  
                for column in columns:  
                    df[column].fillna(fill_value, inplace=True)  
                print(f"已在列 {columns} 中填充空值：{fill_value}")  
                  
            elif method == "转换格式":  
                target_format = params["target_format"]  
                if not target_format:  
                    print("错误：目标格式不能为空")  
                    continue  
                for column in columns:  
                    try: 
                        if target_format=="日期":
                            df[column] = pd.to_datetime(df[column], format='%Y-%m-%d').dt.strftime('%Y-%m-%d')
                        elif target_format=="年份":
                            df[column] = pd.to_datetime(df[column], format='%Y-%m-%d').dt.strftime('%Y') 
                        elif target_format=="月份":
                            df[column] = pd.to_datetime(df[column], format='%Y-%m-%d').dt.strftime('%Y-%m')
                        elif target_format=="季度":
                            quarterx = pd.to_datetime(df[column], format='%Y-%m-%d').dt.quarter
                            yearx=pd.to_datetime(df[column], format='%Y-%m-%d').dt.strftime('%Y') 
                            df[column]  = yearx.astype(str) +"Q" +quarterx.astype(str)     
                        elif target_format=="round2":
                            df[column] = round(df[column].astype(float),2)                
                        else:
                            df[column] = df[column].astype(target_format)  
                    except ValueError as e:  
                        print(f"错误：无法将列 '{column}' 转换为 {target_format} 格式。错误消息：{str(e)}")  
                        continue  
                print(f"已将列 {columns} 的格式转换为 {target_format}")  


            elif method == "保留包含" or  method == "删除包含" :  
                keyword_ct = params["keyword_ct"]  
                if not keyword_ct:  
                    print("错误", "关键词不能为空")  
                    return  
                
                try:  
                    if method == "删除包含":
                        df = df[~df[columns[0]].str.contains(keyword_ct)]
                    elif method == "保留包含":
                        df = df[df[columns[0]].str.contains(keyword_ct,na=False)]               
                except ValueError as e:  
                    messagebox.showinfo("错误", f"无法完成任务，原因：{str(e)}")  
                    return

            elif method == "删空值行" :   
                for column in columns:  
                    df.dropna(subset=[column], inplace=True)  
                                  
            elif method == "扩展多行":  
                word_sep = params["word_sep"]  
                if not word_sep:  
                    print("错误：分隔符不能为空")  
                    continue  
                try:  
                    df = CLEAN_expand_rows(df, word_sep, columns)  # 假设CLEAN_expand_rows函数已定义  
                except ValueError as e:  
                    print(f"错误：无法完成任务，原因：{str(e)}")  
                    continue  
                print(f"已使用分隔符 {word_sep} 在列 {columns} 中扩展多行")  

            elif method == "还原扩行":  
                word_sep = params["word_sep"]  
                if not word_sep:  
                    print("错误：分隔符不能为空")  
                    continue  
                try:  
                    df = CLEAN_expand_rows_REVERT(df, word_sep, columns)  # 假设CLEAN_expand_rows函数已定义  
                except ValueError as e:  
                    print(f"错误：无法完成任务，原因：{str(e)}")  
                    continue  
                print(f"已使用分隔符 {word_sep} 在列 {columns} 中还原扩展多行")  

            elif method == "扩展多列":  #hhhhhhhhh
                word_sep2 = params["word_sep2"]  
                if not word_sep2:  
                    print("错误：分隔符不能为空")  
                    return  
                try:  
                    # 假设CLEAN_expand_cols是一个有效的函数，用于扩展多列  
                    df = CLEAN_expand_cols(df, word_sep2, columns[0])

                except ValueError as e:  
                    print("错误：无法完成任务，原因：" + str(e))  
                    return  

            elif method == "升序排列":  
                df.sort_values(by=columns, ascending= [True] * len(columns)  , inplace=True)  
        
            elif method == "降序排列":  
                df.sort_values(by=columns, ascending= [False] * len(columns)  , inplace=True)   
      
            elif method == "按列去重":  
                df.drop_duplicates(subset=columns, keep='first', inplace=True)  
      
            elif method == "删除本列":  
                df.drop(columns=columns, inplace=True)  
      
            elif method == "选列保留":  
                all_columns = columns  # 获取所有列名  
                columns_to_drop = [col for col in all_columns if col not in columns]  # 找出不在selected_columns中的列名  
                df.drop(columns=columns_to_drop, inplace=True)  # 删除这些列

            elif method == "整体去重":  
                df.drop_duplicates(inplace=True)   
                  
            elif method == "重置索引":  
                df.reset_index(inplace=True)  
      
            elif method == "查看数据":  
                pass  # 不执行任何操作，仅用于查看数据  

            elif method == "批量规整":  
                guize = params["guize"]
                print(guize)
                df=CLEAN_replay_operations(df,pd.DataFrame(eval(str(guize))))
                
      
            elif method == "加关键词" or  method == "赋关键词" :  
                mark_column_name = params["mark_column_name"]   
                include_str = params["include_str"] 
                exclude_str = params["exclude_str"]   
                predefined_column = params["predefined_column"]  
                predefined_objects = params["predefined_objects"]   
                mark_value = params["mark_value"]    
                       
                df=CLEAN_filter_and_mark_rows(df, columns, include_str, exclude_str, predefined_column, predefined_objects, mark_column_name, mark_value, params["method"] ) 
                                       #                                                       
                  
            else:  
                print(f"不支持的操作：{method}")  # 对于不支持的操作类型，打印一条消息并跳过  
            #operation_log.loc[ids,"hash_auto"]=SMALL_get_dataframe_hash(df)#其他地方验证
            print("操作完成：" + str(ids)+method)  # 打印操作完成的消息，以替代messagebox.showinfo的功能   
        except Exception  as e:
            print("操作失败：" + str(ids)+method+" 原因： "+str(e))
                     
    print("所有批量规整工作完成。")
    
    #operation_log["check"]=operation_log["hash_auto"]-operation_log["hash"]#其他地方验证

    return df




def CLEAN_table(df):  
    """  
    合并多列和修改列名函数  
    """ 
    #为验证hash而作。
    df_hash_auto=df.copy()

    operation_log = []  

    def on_confirm(dfsd,method): 
        nonlocal df  #dfsd实际没有作用，没有哪里使用
        
        #增加记录
        operation = {  
                    "方法": "",  
                    "作用列": [],  
                    "参数": {},  
                } 
        selected_columns = [tree.item(item_id, "text") for item_id in tree.selection()]  # 获取多选的列名    
        operation["作用列"] = selected_columns   
        operation["方法"] = method               
        if not selected_columns and method not in ["重置索引","查看数据","查看轨迹","批量规整","新建一列","整体去重"]: 
            messagebox.showinfo("错误", "没有选择任何列")  
            return  # 如果没有选择任何列，则直接返回  
                      
        if method == "复制合并":  
            new_column_name = new_column_entry.get()  # 获取新列名  
            separator = separator_entry.get()  # 获取分隔符             
            # 检查新列名是否为空或与现有列名重复  
            if not new_column_name or new_column_name in df.columns:  
                messagebox.showinfo("错误", "新列名不能为空或与现有列名重复")  
                return  
            df[new_column_name] = df[selected_columns].apply(lambda row: separator.join(row.astype(str)), axis=1)  
            operation["参数"] = {"new_column_name": new_column_name, "separator": separator}  #######################################



        elif method == "条件赋值":  
 
            compare_value=corr_express.get() #获取条件。
            G_value=cor_entry2.get()  #获取赋值。        

            if len(selected_columns)!=1:  
                messagebox.showinfo("错误", "请选择1个数据列。")  
                return               
            df=SMALL_assign_value_based_on_expression(df, selected_columns[0], compare_value, G_value)
            operation["参数"] = {"compare_value":compare_value,"G_value":G_value}  #######################################


        elif method == "两列运算":  
            new_column_name = Enew_column_entry.get()  # 获取新列名  
            separator = Eseparator_entry.get()  # 获取分隔符             
            # 检查新列名是否为空或与现有列名重复  

            if len(selected_columns)!=2:  
                messagebox.showinfo("错误", "请选择两个数据列。")  
                return  
            df = SMALL_apply_operation(df,selected_columns,new_column_name,separator) 
            operation["参数"] = {"new_column_name": new_column_name, "separator": separator}  #######################################
                       
        elif method == "重命名列":  
            if len(selected_columns) > 1:  
                messagebox.showinfo("提示", "请只选择一列进行重命名")  
                return  
            new_column_name = Anew_column_entry.get().strip()  # 获取新列名并去除首尾空格                
            # 检查新列名是否为空或与现有列名（除了被重命名的列）重复  
            if not new_column_name or new_column_name in df.columns and new_column_name != selected_columns[0]:  
                messagebox.showinfo("错误", "新列名不能为空或与现有列名重复")  
                return   
            df.rename(columns={selected_columns[0]: new_column_name}, inplace=True)  
            operation["参数"] = {"new_column_name": new_column_name}  ###################################################           

        elif method == "最后非空":  
            new_column_name = Anew_column_entry.get().strip()  # 获取新列名并去除首尾空格                
            # 检查新列名是否为空或与现有列名（除了被重命名的列）重复  
            if not new_column_name or new_column_name in df.columns and new_column_name != selected_columns[0]:  
                messagebox.showinfo("错误", "新列名不能为空或与现有列名重复")  
                return   
            df=SMALL_last_non_null_value(df,selected_columns,new_column_name) 
            operation["参数"] = {"new_column_name": new_column_name}  ###################################################  



        elif method == "新建一列":  
			
            if len(selected_columns) > 1:  
                messagebox.showinfo("提示", "请只选择一列进行重命名")  
                return  
            new_column_name = AAnew_column_entry.get().strip()  # 获取新列名并去除首尾空格                
            # 检查新列名是否为空或与现有列名（除了被重命名的列）重复 
            new_column_valuex = new_column_valuex_entry.get()       
            # 检查新列名是否为空或与现有列名（除了被重命名的列）重复              
            if not new_column_name or new_column_name in df.columns and new_column_name != selected_columns[0]:  
                messagebox.showinfo("错误", "新列名不能为空或与现有列名重复")  
                return   
            df[new_column_name]=new_column_valuex  
            operation["参数"] = {"new_column_name": new_column_name,"new_column_valuex":new_column_valuex}  ###################################################     

        elif method == "并用追加":  
			
            if len(selected_columns) > 1:  
                messagebox.showinfo("提示", "请只选择一列进行（比如并用药品列）")  
                return  
            new_column_name = SAAnew_column_entry.get().strip()  # 获取新列名并去除首尾空格                
            new_column_valuex = Snew_column_valuex_entry.get()       
            
 
            df=CLEAN_fill_column_c_based_on_a_and_b(df,new_column_name,selected_columns[0],new_column_valuex)
            operation["参数"] = {"new_column_name": new_column_name,"new_column_valuex":new_column_valuex}  ###################################################     

            
        elif method == "替换字符": 
            if not selected_columns:  
                return  # 如果没有选择任何列，则直接返回            
            find_text = find_entry.get()  # 获取要查找的字符  
            replace_text = replace_entry.get()  # 获取要替换的字符            
            if not find_text:  
                messagebox.showinfo("错误", "请输入要查找的字符")  
                return            
            # 在选定的列中执行字符替换操作  
            for col in selected_columns:  
                df[col] = df[col].astype(str).str.replace(find_text, replace_text) 
            operation["参数"] = {"find_text": find_text, "replace_text": replace_text}  #######################################


        elif method == "填充空值":  
            fill_value = fill_entry.get().strip()  # 获取填充值并去除首尾空格  
            if not fill_value:  
                messagebox.showinfo("错误", "填充值不能为空")  
                return  
            try:  
                # 尝试将填充值转换为适当的数据类型（例如，如果列是整数类型，则转换为整数）  
                fill_value = df[selected_columns[0]].dtype.type(fill_value)  
            except ValueError:  
                messagebox.showinfo("错误", "填充值与列的数据类型不匹配")  
                return  
            for column in selected_columns:  
                df[column].fillna(fill_value, inplace=True)  # 使用指定的填充值填充选定列的空值
            operation["参数"] = {"fill_value": fill_value}  ###################################################  
                     
        elif method == "删空值行" :   
            for column in selected_columns:  
                df.dropna(subset=[column], inplace=True)  
                
        elif method == "转换格式":  
            target_format = format_entry.get().strip()  # 获取目标格式并去除首尾空格  
            if not target_format:  
                messagebox.showinfo("错误", "目标格式不能为空")  
                return  
            for column in selected_columns:  
                try: 
                    if target_format=="日期":
                        df[column] = pd.to_datetime(df[column], format='%Y-%m-%d').dt.strftime('%Y-%m-%d')
                    elif target_format=="年份":
                        df[column] = pd.to_datetime(df[column], format='%Y-%m-%d').dt.strftime('%Y') 
                    elif target_format=="月份":
                        df[column] = pd.to_datetime(df[column], format='%Y-%m-%d').dt.strftime('%Y-%m')
                    elif target_format=="季度":
                        quarterx = pd.to_datetime(df[column], format='%Y-%m-%d').dt.quarter
                        yearx=pd.to_datetime(df[column], format='%Y-%m-%d').dt.strftime('%Y') 
                        df[column]  = yearx.astype(str) +"Q" +quarterx.astype(str)
                    elif target_format=="round2":
                        df[column] = round(df[column].astype(float),2) 

                    else: 
                        # 尝试将选定列的数据类型转换为目标格式  
                        df[column] = df[column].astype(target_format)  
                except ValueError as e:  
                    messagebox.showinfo("错误", f"无法将列 '{column}' 转换为 {target_format} 格式。错误消息：{str(e)}")  
                    return
            operation["参数"] = {"target_format": target_format}  ###################################################           

        elif method == "保留包含" or  method == "删除包含" :  
            keyword_ct = ct_entry.get().strip()  # 获取目标格式并去除首尾空格  
            if not keyword_ct:  
                messagebox.showinfo("错误", "关键词不能为空")  
                return  

            try:  
                if method == "删除包含":
                    df = df[~df[selected_columns[0]].str.contains(keyword_ct,na=False)]  #
                elif method == "保留包含":
                    df = df[df[selected_columns[0]].str.contains(keyword_ct,na=False)]               
            except ValueError as e:  
                messagebox.showinfo("错误", f"无法完成任务，原因：{str(e)}")  
                return
            operation["参数"] = {"keyword_ct": keyword_ct}  ###################################################  

        elif method == "众数对齐" :  
            group_col = ffd_entry.get().strip()  # 获取目标格式并去除首尾空格  
            if not group_col:  
                messagebox.showinfo("错误", "对齐源列不能为空")  
                return  

            try: 

                df=SMALL_align_values_to_most_frequent(df,group_col,selected_columns)             
            except ValueError as e:  
                messagebox.showinfo("错误", f"无法完成任务，原因：{str(e)}")  
                return
            operation["参数"] = {"group_col": group_col}  ###################################################  


        elif method == "扩展多行":  

            word_sep = exp_row_entry.get().strip()  # 获取目标格式并去除首尾空格  
            if not word_sep:  
                messagebox.showinfo("错误", "分隔符不能为空")  
                return  

            try:  
                df=CLEAN_expand_rows(df.copy(),word_sep,selected_columns) 

            except ValueError as e:  
                messagebox.showinfo("错误", f"无法完成任务，原因：{str(e)}")  
                return
            operation["参数"] = {"word_sep": word_sep}  ###################################################           
                

        elif method == "还原扩行":  
            Rword_sep = Rexp_row_entry.get().strip()  # 获取目标格式并去除首尾空格  
            if not Rword_sep:  
                messagebox.showinfo("错误", "分隔符不能为空")  
                return  

            try:  
                df=CLEAN_expand_rows_REVERT(df,Rword_sep,selected_columns) 
            except ValueError as e:  
                messagebox.showinfo("错误", f"无法完成任务，原因：{str(e)}")  
                return
            operation["参数"] = {"word_sep": Rword_sep}  ###################################################   
                            
        elif method == "扩展多列":  
            word_sep2 = exp_row_entry.get().strip()  # 获取目标格式并去除首尾空格  
            if not word_sep2:  
                messagebox.showinfo("错误", "分隔符不能为空")  
                return  
            try:  
                df=CLEAN_expand_cols(df,word_sep2,selected_columns[0]) 

            except ValueError as e:  
                messagebox.showinfo("错误", f"无法完成任务，原因：{str(e)}")  
                return
            operation["参数"] = {"word_sep2": word_sep2}  ###################################################           
            
        elif method == "按列去重":  
            df.drop_duplicates(subset=selected_columns, keep='first', inplace=True)


		#sort_value_up_button	
        elif method == "升序排列":  
            df.sort_values(by=selected_columns, ascending= [True] * len(selected_columns)  , inplace=True)  
    
        elif method == "降序排列":  
            df.sort_values(by=selected_columns, ascending= [False] * len(selected_columns)  , inplace=True)     
               
        elif method == "删除本列":  
            df.drop(columns=selected_columns, inplace=True)
            
        elif method == "整体去重":  
            df.drop_duplicates(inplace=True)            
            
        elif method == "选列保留":  
            all_columns = df.columns.tolist()  # 获取所有列名  
            columns_to_drop = [col for col in all_columns if col not in selected_columns]  # 找出不在selected_columns中的列名  
            df.drop(columns=columns_to_drop, inplace=True)  # 删除这些列
        elif method == "重置索引":  
            df.reset_index(inplace=True) 
        elif method == "查看数据":  
            pass  
        elif method == "查看轨迹":  
            if operation_log==[]:
                return
            PROGRAM_display_df_in_treeview(pd.DataFrame(operation_log),0,df)
            return 
            
        elif method == "批量规整":  
            guize=SMALL_read_excel_files()
            df=CLEAN_replay_operations(df,guize)
            
            operation["参数"] = {"guize": guize.to_dict(orient='list')}  #######################################  
                  
        elif method == "加关键词" or  method == "赋关键词":  #XFDD
            mark_column_name = mark_column_entry.get().strip()  # 获取新列名并去除首尾空格  
            predefined_all=fdel_column_entry2.get().strip() 
            
            parts0 = predefined_all.split("|") 
            predefined_column= parts0[0]
            predefined_objects=parts0[1]  
                   
            keyword_include = keyword_entry.get().strip()  # 获取关键词并去除首尾空格   
            parts1 = keyword_include.split("|")  
            mark_value=parts1[0] 
            include_str=keyword_include
              
            exclude_str=fdel_column_entry.get()
            
            df=CLEAN_filter_and_mark_rows(df, selected_columns, include_str, exclude_str, predefined_column, predefined_objects, mark_column_name, mark_value, method) 
                                   #                                                       
            operation["参数"] = {"mark_column_name": mark_column_name, "include_str": include_str, "exclude_str": exclude_str, "predefined_column": predefined_column, "predefined_objects": predefined_objects, "mark_value": mark_value, "method": method}  #######################################

        else:  
            return  

        """更新TreeView以显示DataFrame的列名"""  
        # 清空当前的TreeView内容  
        tree.delete(*tree.get_children())  
          
        # 根据最新的DataFrame列名重新填充TreeView  
        for col in df.columns:  
            tree.insert("", tk.END, text=str(col), values=())


            
            
        operation_log.append(operation)  
        
        if 1==2:
    
            #本步骤hash验证
            hash_gui=SMALL_get_dataframe_hash(df.copy()) 
            dataxx=pd.DataFrame(operation_log)
            with pd.ExcelWriter(r"guize.xls", engine="xlsxwriter") as writer:  
                dataxx.to_excel(writer, sheet_name="导出的数据", index=False)  
            guizex = pd.read_excel(r"guize.xls")  
            df_hash_auto_result=CLEAN_replay_operations(df_hash_auto.copy(),guizex )
            
            hash_auto=SMALL_get_dataframe_hash(df_hash_auto_result.copy()) 
            if hash_gui == hash_auto:
                print("●●●●●●●●●●●●●●●●●●●●●●●●pass")
            else:
                print("●●●●●●●●●●●●●●●●●●●●●●●●fail！！！！！！！！！！！！！")
            with pd.ExcelWriter(r"对比文件.xls", engine="xlsxwriter") as writer:  
                df.to_excel(writer, sheet_name="gui", index=False)          
                df_hash_auto_result.to_excel(writer, sheet_name="auto", index=False)  
        print(operation_log)                
        PROGRAM_display_df_in_treeview(df,0,0)  # 更新树状视图  
 

  
    title = "数据清洗工具" 
    root = tk.Tk()  
    root.title(title)  
  
    sw = root.winfo_screenwidth()  
    sh = root.winfo_screenheight()  
    ww = 1230  # 窗口宽度  
    wh = 630  # 窗口高度  
    x = (sw - ww) // 2  
    y = (sh - wh) // 2  
    root.geometry(f"{ww}x{wh}+{x}+{y}")  
  
    # 创建主框架，用于放置所有控件  
    main_frame = ttk.Frame(root)  
    main_frame.pack(fill="both", expand=True, padx=10, pady=10)  
  
    # 创建左侧树状视图，显示df的列名  
    tree_frame = ttk.Frame(main_frame)  
    tree_frame.pack(side="left", fill="both", expand=True, padx=10, pady=10)  
    tree_label = ttk.Label(tree_frame, text="请选择要操作的列:")  
    tree_label.pack(anchor='w')  
  
    # 添加滚动条  
    scrollbar = ttk.Scrollbar(tree_frame)  
    scrollbar.pack(side="right", fill="y")  
  
    tree = ttk.Treeview(tree_frame, yscrollcommand=scrollbar.set)  
    scrollbar.config(command=tree.yview)  
    tree['columns'] = ('Column',)  
    tree.column('#0', width=200, stretch=tk.NO)  
    tree.heading('#0', text='列名', anchor=tk.W)  
  
    for col in df.columns:  
        tree.insert("", tk.END, text=str(col), values=())  
  
    tree.pack(fill="both", expand=True, padx=5, pady=5)  
  
    # 创建右侧内容区域，使用Grid布局来放置所有控件在同一行  
    right_frame = ttk.Frame(main_frame)  
    right_frame.pack(side="right", fill="both", expand=True, padx=10, pady=10)  

    print("复制列请用复制合并功能。")
    
    # FSM 复制合并   
    separator_label = ttk.Label(right_frame, text="分隔符：")  
    separator_label.grid(row=0, column=0, padx=5, pady=2, sticky='w')  
    separator_entry = tk.Entry(right_frame, width=20)  
    separator_entry.grid(row=0, column=1, padx=5, pady=2, sticky='w')  
    new_column_label = ttk.Label(right_frame, text="新列名：")  
    new_column_label.grid(row=0, column=2, padx=5, pady=2, sticky='w')  
    new_column_entry = tk.Entry(right_frame, width=20)  
    new_column_entry.grid(row=0, column=3, padx=5, pady=2, sticky='w')  
    confirm_button = tk.Button(right_frame, text="复制合并", command=lambda:on_confirm(df,"复制合并"))  
    confirm_button.grid(row=0, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2)  
    # 字符替换 
    find_label = ttk.Label(right_frame, text="查找的字符：")  
    find_label.grid(row=1, column=0, padx=5, pady=2, sticky='w')  
    find_entry = tk.Entry(right_frame, width=20)  
    find_entry.grid(row=1, column=1, padx=5, pady=2, sticky='w')  
      
    replace_label = ttk.Label(right_frame, text="替换值：")  
    replace_label.grid(row=1, column=2, padx=5, pady=2, sticky='w')  
    replace_entry = tk.Entry(right_frame, width=20)  
    replace_entry.grid(row=1, column=3, padx=5, pady=2, sticky='w')  
      
    replace_button = tk.Button(right_frame, text="替换字符", command=lambda:on_confirm(df,"替换字符"))  
    replace_button.grid(row=1, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2) 
         
    # FSM 重命名列   
    Anew_column_label = ttk.Label(right_frame, text="新列名：")  
    Anew_column_label.grid(row=2, column=2, padx=5, pady=2, sticky='w')  
    Anew_column_entry = tk.Entry(right_frame, width=20)  
    Anew_column_entry.grid(row=2, column=3, padx=5, pady=2, sticky='w')  
    rename_button = tk.Button(right_frame, text="修改列名", command=lambda:on_confirm(df,"重命名列"))  # 使用单独的按钮和命令来处理重命名操作  
    rename_button.grid(row=2, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2)  # 使用单独的按钮来处理重命名操作  
    Rrename_button = tk.Button(right_frame, text="最后非空", command=lambda:on_confirm(df,"最后非空"))  # 使用单独的按钮和命令来处理重命名操作  
    Rrename_button.grid(row=2, column=6, padx=(0, 10), pady=2, sticky='e', columnspan=2)  # 使用单独的按钮来处理重命名操作  
    # FSM 空值填充  删空值行
    fill_label = ttk.Label(right_frame, text="填充值：")  
    fill_label.grid(row=3, column=2, padx=5, pady=2, sticky='w')  
    fill_entry = tk.Entry(right_frame, width=20)  
    fill_entry.grid(row=3, column=3, padx=5, pady=2, sticky='w')  
    fill_button = tk.Button(right_frame, text="填充空值", command=lambda: on_confirm(df,"填充空值"))  # 使用单独的按钮和命令来处理空值填充操作  
    fill_button.grid(row=3, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2)  # 使用单独的按钮来处理空值填充操作  
    delfill_button = tk.Button(right_frame, text="删空值行", command=lambda: on_confirm(df,"删空值行"))  # 使用单独的按钮和命令来处理空值填充操作  
    delfill_button.grid(row=3, column=6, padx=(0, 10), pady=2, sticky='e', columnspan=2)  # 使用单独的按钮来处理空值填充操作  

    # FSM 格式转换 
    format_options = ["str", "float", "int","round2", "日期", "月份", "季度", "年份"]   
    format_label = ttk.Label(right_frame, text="目标格式（如str）：")  
    format_label.grid(row=4, column=2, padx=5, pady=2, sticky='w')  
    format_entry =  ttk.Combobox(right_frame, values=format_options, width=20)   
    format_entry.grid(row=4, column=3, padx=5, pady=2, sticky='w')  
    convert_button = tk.Button(right_frame, text="转换格式", command=lambda: on_confirm(df,"转换格式"))   
    convert_button.grid(row=4, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2)  #  


    # FSM 扩展多行   扩展多列
    exp_row_label = ttk.Label(right_frame, text="分隔符：")  
    exp_row_label.grid(row=5, column=2, padx=5, pady=2, sticky='w')  
    exp_row_entry = tk.Entry(right_frame, width=20)  
    exp_row_entry.grid(row=5, column=3, padx=5, pady=2, sticky='w')  
    exp_row_button = tk.Button(right_frame, text="扩展多行", command=lambda:on_confirm(df,"扩展多行"))  #  
    exp_row_button.grid(row=5, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2)  # 
    exp_col_button = tk.Button(right_frame, text="扩展多列", command=lambda:on_confirm(df,"扩展多列"))  #   
    exp_col_button.grid(row=5, column=6, padx=(0, 10), pady=2, sticky='e', columnspan=2)  #  
    
     
    # FSM 含与不含  
    ct_label = ttk.Label(right_frame, text="关键词（仅一列）：")  
    ct_label.grid(row=6, column=2, padx=5, pady=2, sticky='w')  
    ct_entry = tk.Entry(right_frame, width=20)  
    ct_entry.grid(row=6, column=3, padx=5, pady=2, sticky='w')  
    ct_keep_button = tk.Button(right_frame, text="保留包含", command=lambda:on_confirm(df,"保留包含"))  #  
    ct_keep_button.grid(row=6, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2)  # 
    ct_del_button = tk.Button(right_frame, text="删除包含", command=lambda:on_confirm(df,"删除包含"))  #   
    ct_del_button.grid(row=6, column=6, padx=(0, 10), pady=2, sticky='e', columnspan=2)  #   

    # FSM 众数对齐  
    ffd_label = ttk.Label(right_frame, text="众数源列：")  
    ffd_label.grid(row=7, column=2, padx=5, pady=2, sticky='w')  
    ffd_entry = tk.Entry(right_frame, width=20)  
    ffd_entry.grid(row=7, column=3, padx=5, pady=2, sticky='w')  
    ffd_button = tk.Button(right_frame, text="众数对齐", command=lambda:on_confirm(df,"众数对齐"))  #  
    ffd_button.grid(row=7, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2)  # 


      
    # FSM 还原扩行   
    Rexp_row_label = ttk.Label(right_frame, text="分隔符：")  
    Rexp_row_label.grid(row=11, column=2, padx=5, pady=2, sticky='w')  
    Rexp_row_entry = tk.Entry(right_frame, width=20)  
    Rexp_row_entry.grid(row=11, column=3, padx=5, pady=2, sticky='w')  
    Rexp_row_button = tk.Button(right_frame, text="还原扩行", command=lambda:on_confirm(df,"还原扩行"))  #  
    Rexp_row_button.grid(row=11, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2)  # 

    # FSM 两列运算   
    Eseparator_label = ttk.Label(right_frame, text="运算式(a+b)：")  
    Eseparator_label.grid(row=12, column=0, padx=5, pady=2, sticky='w')  
    Eseparator_entry = tk.Entry(right_frame, width=20)  
    Eseparator_entry.grid(row=12, column=1, padx=5, pady=2, sticky='w')  
    Enew_column_label = ttk.Label(right_frame, text="新列名：")  
    Enew_column_label.grid(row=12, column=2, padx=5, pady=2, sticky='w')  
    Enew_column_entry = tk.Entry(right_frame, width=20)  
    Enew_column_entry.grid(row=12, column=3, padx=5, pady=2, sticky='w')  
    Econfirm_button = tk.Button(right_frame, text="两列运算", command=lambda:on_confirm(df,"两列运算"))  
    Econfirm_button.grid(row=12, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2)  
    
    # FSM 新建列    new_column_valuex
    AAnew_column_label = ttk.Label(right_frame, text="新列名：")  
    AAnew_column_label.grid(row=13, column=2, padx=5, pady=2, sticky='w')  
    AAnew_column_entry = tk.Entry(right_frame, width=20)  
    AAnew_column_entry.grid(row=13, column=3, padx=5, pady=2, sticky='w')  
    AAnew_column_label2 = ttk.Label(right_frame, text="赋值：")  
    AAnew_column_label2.grid(row=13, column=0, padx=5, pady=2, sticky='w')  
    new_column_valuex_entry = tk.Entry(right_frame, width=20)  
    new_column_valuex_entry.grid(row=13, column=1, padx=5, pady=2, sticky='w')
    creatnew_button = tk.Button(right_frame, text="新建一列", command=lambda:on_confirm(df,"新建一列"))  # 使用单独的按钮和命令来处理重命名操作  
    creatnew_button.grid(row=13, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2)  # 使用单独的按钮来处理重命名操作  

    # FSM 并用追加
    SAAnew_column_label = ttk.Label(right_frame, text="依据列：")  
    SAAnew_column_label.grid(row=14, column=2, padx=5, pady=2, sticky='w')  
    SAAnew_column_entry = tk.Entry(right_frame, width=20)  
    SAAnew_column_entry.grid(row=14, column=3, padx=5, pady=2, sticky='w')  
    SAAnew_column_label2 = ttk.Label(right_frame, text="新列名：")  
    SAAnew_column_label2.grid(row=14, column=0, padx=5, pady=2, sticky='w')  
    Snew_column_valuex_entry = tk.Entry(right_frame, width=20)  
    Snew_column_valuex_entry.grid(row=14, column=1, padx=5, pady=2, sticky='w')
    Screatnew_button = tk.Button(right_frame, text="并用追加", command=lambda:on_confirm(df,"并用追加"))  #  
    Screatnew_button.grid(row=14, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2)  


    # 条件赋值 

    cor_value_label = ttk.Label(right_frame, text="表达式df['a']：")  
    cor_value_label.grid(row=15, column=0, padx=5, pady=2, sticky='w')  
    corr_express = tk.Entry(right_frame, width=20)  
    corr_express.grid(row=15, column=1, padx=5, pady=2, sticky='w')  
    
    cor_value_labels = ttk.Label(right_frame, text="赋值：")  
    cor_value_labels.grid(row=15, column=2, padx=5, pady=2, sticky='w')  
    cor_entry2 = tk.Entry(right_frame, width=20)  
    cor_entry2.grid(row=15, column=3, padx=5, pady=2, sticky='w')  
       
    Xcor_button = tk.Button(right_frame, text="条件赋值", command=lambda:on_confirm(df,"条件赋值"))  
    Xcor_button.grid(row=15, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2) 
    
    # 找词标记  
 
    column_label = ttk.Label(right_frame, text="被标记列：")  
    column_label.grid(row=17, column=0, padx=5, pady=2, sticky='w')  
    mark_column_entry = tk.Entry(right_frame, width=20)  
    mark_column_entry.grid(row=17, column=1, padx=5, pady=2, sticky='w')     
    keyword_label = ttk.Label(right_frame, text="标记|词1|词2...")  
    keyword_label.grid(row=17, column=2, padx=5, pady=2, sticky='w')      
    keyword_entry = tk.Entry(right_frame, width=20)  
    keyword_entry.grid(row=17, column=3, padx=5, pady=2, sticky='w')  
    fdel_column_label = ttk.Label(right_frame, text="排除值：")  
    fdel_column_label.grid(row=18, column=0, padx=5, pady=2, sticky='w')  
    fdel_column_entry = tk.Entry(right_frame, width=20)  
    fdel_column_entry.grid(row=18, column=1, padx=5, pady=2, sticky='w') 
    fdel_column_label2 = ttk.Label(right_frame, text="作用对象|列内对象：")  
    fdel_column_label2.grid(row=18, column=2, padx=5, pady=2, sticky='w')  
    fdel_column_entry2 = tk.Entry(right_frame, width=20)  
    fdel_column_entry2.insert(0, "所有列|所有对象")  # 在这里插入默认值
    fdel_column_entry2.grid(row=18, column=3, padx=5, pady=2, sticky='w')            
    add_keyword_button = tk.Button(right_frame, text="找词加标", command=lambda:on_confirm(df,"加关键词"))  
    add_keyword_button.grid(row=18, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2)
    add_keyword_button2 = tk.Button(right_frame, text="找词赋标", command=lambda:on_confirm(df,"赋关键词"))  
    add_keyword_button2.grid(row=18, column=6, padx=(0, 10), pady=2, sticky='e', columnspan=2)

    # FSM 升序排列 ## 
    sort_value_up_button = tk.Button(right_frame, text="升序排列", command=lambda: on_confirm(df,"升序排列"))  # 使用单独的按钮和命令来处理空值填充操作  
    sort_value_up_button.grid(row=19, column=6, padx=(0, 10), pady=2, sticky='e', columnspan=2)  # 使用单独的按钮来处理空值填充操作 

    # FSM 降序排列 ## 
    sort_value_down_button = tk.Button(right_frame, text="降序排列", command=lambda: on_confirm(df,"降序排列"))  # 使用单独的按钮和命令来处理空值填充操作  
    sort_value_down_button.grid(row=20, column=6, padx=(0, 10), pady=2, sticky='e', columnspan=2)  # 使用单独的按钮来处理空值填充操作 
     
    # FSM 按列去重 
    drop_duplicates_button = tk.Button(right_frame, text="按列去重", command=lambda: on_confirm(df,"按列去重"))  # 使用单独的按钮和命令来处理空值填充操作  
    drop_duplicates_button.grid(row=19, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2)  # 使用单独的按钮来处理空值填充操作  

    # FSM 删除本列 
    drop_button = tk.Button(right_frame, text="删除本列", command=lambda: on_confirm(df,"删除本列"))  # 使用单独的按钮和命令来处理空值填充操作  
    drop_button.grid(row=19, column=2, padx=(0, 10), pady=2, sticky='e', columnspan=2)  # 使用单独的按钮来处理空值填充操作  

    # FSM 选列保留 
    keep_button = tk.Button(right_frame, text="选列保留", command=lambda: on_confirm(df,"选列保留"))  # 使用单独的按钮和命令来处理空值填充操作  
    keep_button.grid(row=19, column=0, padx=(0, 10), pady=2, sticky='e', columnspan=2)  # 使用单独的按钮来处理空值填充操作  

    # FSM 整体去重 
    dropdu_button = tk.Button(right_frame, text="整体去重", command=lambda: on_confirm(df,"整体去重"))  # 使用单独的按钮和命令来处理空值填充操作  
    dropdu_button.grid(row=19, column=1, padx=(0, 10), pady=2, sticky='e', columnspan=2)  # 使用单独的按钮来处理空值填充操作  

    # FSM 重置索引 
    reset_button = tk.Button(right_frame, text="重置索引", command=lambda: on_confirm(df,"重置索引"))  # 使用单独的按钮和命令来处理空值填充操作  
    reset_button.grid(row=20, column=0, padx=(0, 10), pady=2, sticky='e', columnspan=2)  # 使用单独的按钮来处理空值填充操作 
    # FSM 批量规整 
    sll_button = tk.Button(right_frame, text="批量规整", command=lambda: on_confirm(df,"批量规整"))  # 使用单独的按钮和命令来处理空值填充操作  
    sll_button.grid(row=20, column=1, padx=(0, 10), pady=2, sticky='e', columnspan=2)  # 使用单独的按钮来处理空值填充操作  
    # FSM 查看轨迹 
    rec_button = tk.Button(right_frame, text="查看轨迹", command=lambda:  on_confirm(df,"查看轨迹"))  
    rec_button.grid(row=20, column=2, padx=(0, 10), pady=2, sticky='e', columnspan=2)  # 使用单独的按钮来处理空值填充操作 
    
    # FSM 查看数据 
    adint_button = tk.Button(right_frame, text="查看数据", command=lambda: on_confirm(df,"查看数据"))  # 使用单独的按钮和命令来处理空值填充操作  
    adint_button.grid(row=20, column=4, padx=(0, 10), pady=2, sticky='e', columnspan=2)  # 使用单独的按钮来处理空值填充操作 


    root.lift()
    root.attributes("-topmost", True)
    root.attributes("-topmost", False)


def CLEAN_filter_and_mark_rows(df, columns, include_str, exclude_str, predefined_column=None, predefined_objects=None, mark_column_name='mark', mark_value='marked', method_name='add'):  
    """  
    对DataFrame中的行进行筛选和标记。  
      
    参数：  
        df (pandas.DataFrame): 需要处理的DataFrame。  
        columns (list of str): 要检查的列名列表。  
        include_str (str): 需要包含的字符串。  
        exclude_str (str): 需要排除的字符串。  
        predefined_column (str, optional): 预定义的列名，用于进一步筛选行。默认为None。  
        predefined_objects (list, optional): 预定义的对象列表，与predefined_column配合使用。默认为None。  
        mark_column_name (str, optional): 标记列的列名。默认为'mark'。  
        mark_value (str, optional): 标记值。默认为'marked'。  
        method_name (str, optional): 处理标记列的方法，'加关键词'表示追加标记值，'赋关键词'表示替换为标记值。默认为'加关键词'。  
      
    返回：  
        pandas.DataFrame: 处理后的DataFrame，包含新增或更新的标记列。  
      
    步骤：  
    1. 创建一个布尔掩码，初始化为True，以便保留所有行。  
    2. 如果提供了预定义列和对象，则基于它们更新掩码，仅包括预定义对象所在的行。  
    3. 使用布尔掩码筛选包含include_str且不包含exclude_str的行。  
    4. 如果需要标记的列不存在于DataFrame中，则新建该列，并初始化为空字符串。  
    5. 根据method_name的值处理标记列：  
        - 如果method_name为'add'，则将标记值追加到筛选结果为True的行的标记列中（用分号隔开）。  
        - 如果method_name为'replace'，则将筛选结果为True的行的标记列替换为标记值。  
    6. 返回处理后的DataFrame，注意行数保持不变，未筛选出的行不会被删除。  
    """  
    # 创建一个布尔掩码，初始化为True，以便保留所有行  
    mask = pd.Series(True, index=df.index)  
    
    if predefined_column=="所有列" and predefined_objects=="所有对象":
        predefined_column=None
        predefined_objects=None
      
    # 如果提供了预定义列和对象，则基于它们更新掩码  
    if predefined_column and predefined_objects:  
        if predefined_column not in df.columns:  
            raise ValueError(f"Predefined column '{predefined_column}' not found in DataFrame.")  
        mask &= df[predefined_column] == predefined_objects  # 更新掩码以仅包括预定义对象所在的行  
  
      
    # 使用布尔掩码筛选包含include_str且不包含exclude_str的行  
    # 检查exclude_str是否为空  
    if exclude_str:  
        mask &= df[columns].apply(lambda col: col.str.contains(include_str, na=False) & ~col.str.contains(exclude_str, na=False)).any(axis=1)  
    else:  
        mask &= df[columns].apply(lambda col: col.str.contains(include_str, na=False)).any(axis=1)  
           
    # 如果需要标记的列不存在于DataFrame中，则新建该列，并初始化为空字符串  
    if mark_column_name not in df.columns:  
        df[mark_column_name] = ''  # 初始化为空字符串，确保后续可以正确添加或替换标记值  
      
    # 根据method_name的值处理标记列  
    if method_name == '加关键词':  
        # 追加标记值（用分号隔开），仅对筛选结果为True的行操作  
        df.loc[mask, mark_column_name] = df.loc[mask, mark_column_name].astype(str) + ';' + mark_value  
    elif method_name == '赋关键词':  
        # 替换为标记值，仅对筛选结果为True的行操作  
        df.loc[mask, mark_column_name] = mark_value  
    else:  
        raise ValueError(f"Invalid method_name '{method_name}'. Only 'add' or 'replace' are accepted.")  
    df[mark_column_name] = df[mark_column_name].str.lstrip(';')

    # 返回处理后的DataFrame，注意行数保持不变，未筛选出的行不会被删除  
    return df

def CLEAN_expand_cols(dfs, delimiter, col):  
    """  
    扩展指定列的函数。  只支持单列。
  
    参数:  
    df (DataFrame): 输入的DataFrame。  
    delimiter (str): 用于拆分指定列的分隔符。  
    col (str): 需要拆分的列名。  
  
    返回值:  
    DataFrame: 扩展后的DataFrame，包含原始列和新的拆分列。  
  
    步骤说明:  
    1. 使用分隔符拆分指定的列，并创建虚拟/指示列。  
    2. 通过将原始列名和唯一值附加到一起来重命名列。  
    3. 将原始DataFrame与虚拟DataFrame沿列方向（axis=1）连接。  
    """  
    # Split the column by delimiter and create dummy columns    
    dummy_df = dfs[col].str.get_dummies(sep=delimiter)    
        
    # Rename the columns by appending the original column name and the unique values    
    dummy_df.columns = [f'{col}_{value}' for value in dummy_df.columns]    
        
    # Concatenate the original dataframe with the dummy dataframe    
    result_df = pd.concat([dfs, dummy_df], axis=1)    
   
    return result_df
  


     
def CLEAN_expand_rows(df, sep, cols):  
    """  
    拆分成行的函数，sep是分隔符，cols是需要拆分的列的列表。 支持多列。
    """  
    # 重置索引  
    df = df.reset_index(drop=True)  
      
    # 拆分指定的列  
    frames = []  
    for col in cols:  
        frame = df[col].str.split(sep, expand=True).stack().reset_index(level=1, drop=True).to_frame(col)  
        frames.append(frame)  
      
    # 合并拆分的列  
    result = pd.concat(frames, axis=1)  
      
    # 获取原始数据框中非拆分列的部分  
    other_cols = [col for col in df.columns if col not in cols]  
    remaining_df = df[other_cols]  
      
    # 如果有非拆分列，则合并拆分的列与非拆分列的部分  
    if other_cols:  
        result = pd.merge(result, remaining_df, left_index=True, right_index=True)  
    result['groupby_column_0'] = result.index.copy()  
    # 清理结果并返回  
    result = result.reset_index(drop=True)  
    return result    
   
def CLEAN_expand_rows_REVERT(df, delimiter, merge_columns):  
    """  
    根据指定的列合并df的行，并将特定列的值用分号连接起来，同时保留其他列的第一个值。  
      
    参数:  
    df (DataFrame): 要合并的DataFrame。  
    groupby_column (str): 作为合并依据的列名。  
    merge_columns (list of str): 需要合并值的列名列表。  
    delimiter (str): 用于合并字符串值的分隔符。  
      
    返回值:  
    merged_df (DataFrame): 合并后的DataFrame。  
    """ 
    groupby_column="groupby_column_0" 
    # 创建一个新的DataFrame来存储合并后的结果  
    merged_df = pd.DataFrame()  
      
    # 对于需要合并的字符串列，使用groupby和agg来合并它们的值  
    string_merged = df.groupby(groupby_column)[merge_columns].agg(lambda x: delimiter.join(x.astype(str)))  
    # 对于其他列，使用first()保留每个分组的第一个值  
    other_cols = [col for col in df.columns if col not in merge_columns and col != groupby_column]  
    first_values = df.groupby(groupby_column)[other_cols].first()  
      
    # 将两部分结果合并到一个DataFrame中  
    merged_df = pd.concat([string_merged, first_values], axis=1)  
      
    # 将groupby_column从索引移回普通列  
    merged_df.reset_index(inplace=True)  
      
    return merged_df  
  

#######################################
#功能和实用工具函数
def AAAA_Funtion_Tools():
    pass
######################################





#df暂存器


def TOOLS_temp_save_df(df):  
    def save_file():  
        filename = entry.get().strip()  
        if filename and filename not in global_dfs:  
            global_dfs[filename] = df  
            refresh_tree()  
            entry.delete(0, tk.END)  # 清空输入框  
        elif filename in global_dfs:  
            messagebox.showwarning("警告", "该名称已存在，请另选名称。")  
  
    def refresh_tree():  
        tree.delete(*tree.get_children())  
        for filename in global_dfs:  
            tree.insert("", tk.END, text=filename)  
    def concat_dfs_with_same_number_of_columns(dfs):  
        """  
        拼接具有相同数量列的多个DataFrame，使用第一个DataFrame的列名。  
      
        参数:  
            dfs (list of DataFrame): 要拼接的DataFrame列表。  
      
        返回:  
            DataFrame: 拼接后的DataFrame，使用第一个DataFrame的列名。  
        """  
        if not isinstance(dfs, list) or len(dfs) == 0:  
            raise ValueError("输入必须是一个包含至少一个DataFrame的非空列表。")  
      
        # 检查所有DataFrame的列数是否相同  
        num_columns = dfs[0].shape[1]  
        for df in dfs:  
            if df.shape[1] != num_columns:  
                raise ValueError("所有DataFrame必须具有相同数量的列。")  
      
        # 使用第一个DataFrame的列名作为模板  
        first_columns = dfs[0].columns  
      
        # 将每个DataFrame的列名更改为第一个DataFrame的列名  
        aligned_dfs = []  
        for df in dfs:  
            df.columns = first_columns  
            aligned_dfs.append(df)  
      
        # 拼接DataFrame列表，由于列名已经对齐，这里不需要忽略索引或列名  
        concatenated = pd.concat(aligned_dfs)  
      
        return concatenated  
    
    def concat_files():  
        selected_items = tree.selection()  
        if selected_items:  # 检查是否有选中的项目  
            filenames = []  
            new_dfs = []  
            for item in selected_items:  
                filename = tree.item(item, "text")  
                filenames.append(filename)  
                new_dfs.append(global_dfs[filename])  
              
            new_df = concat_dfs_with_same_number_of_columns(new_dfs) # 合并所有选中的DataFrame  
            new_filename = entry.get().strip()  
              
            if new_filename and new_filename not in global_dfs:  
                global_dfs[new_filename] = new_df  
                refresh_tree()  
            else:  
                messagebox.showwarning("警告", "该新名称已存在，请另选名称。")  
        else:  
            messagebox.showinfo("提示", "请选择要合并的行。")

    def find_common_columns(df1, df2):  
        """Find common columns between two DataFrames."""  
        return set(df1.columns).intersection(df2.columns)  
      
    def get_main_df_index(dfs):  
        """Ask the user for the index of the main DataFrame."""  
        xroot = tk.Tk()  
        xroot.withdraw()  # Hide the main window    
        prompt = "请选择以哪个文件为主文件 (0 to {}):".format(len(dfs) - 1)  
        main_df_index = simpledialog.askstring("Select", prompt, parent=xroot)  
        if main_df_index:  
            main_df_index = int(main_df_index)  
            if 0 <= main_df_index < len(dfs):  
                return main_df_index  
        return None  
      
    def merge_dfs(dfs):  
        """Merge all DataFrames into the main DataFrame specified by the user."""  
        main_df_index = get_main_df_index(dfs)  
        if main_df_index is None:  
            return None  
      
        main_df = dfs[main_df_index]  
        dfs_to_merge = [df for i, df in enumerate(dfs) if i != main_df_index]  
      
        for df in dfs_to_merge:  
            common_cols = find_common_columns(main_df, df)  
            if not common_cols:  
                print(f"No common columns found between main_df and df{dfs.index(df)}")  
                continue  
            # By default, 'inner' join is used, which keeps only common rows  
            main_df = pd.merge(main_df, df, on=list(common_cols),how='left')  
      
        return main_df  

  
    def merge_files():  
        selected_items = tree.selection()   
        filenames = [tree.item(item, "text") for item in selected_items] 
        selected_elements = [global_dfs[i] for i in filenames]   
        new_df=merge_dfs(selected_elements)  
        new_filename = entry.get().strip()  
              
        if new_filename and new_filename not in global_dfs:  
            global_dfs[new_filename] = new_df  
            refresh_tree()  
        else:  
            messagebox.showwarning("警告", "该新名称已存在，请另选名称。")  
            
    def on_double_click(event):  
        selected_item = tree.selection()[0]  
        filename = tree.item(selected_item, "text")  
        # 假设您有一个函数来显示DataFrame (这里应该是实际逻辑，而不是print)  
        # PROGRAM_display_df_in_treeview(global_dfs[filename], 1, 0)  
        PROGRAM_display_df_in_treeview(global_dfs[filename],0,0)  # 此处仅为演示  
  
    # 创建GUI窗口  
    mroot = tk.Tk()  
    mroot.title("表单暂存器")  
    # 得到屏幕宽度和高度，并设置窗口位置和大小（这里可以自定义）          
    sw = mroot.winfo_screenwidth()          
    sh = mroot.winfo_screenheight()          
    ww = 500  # 窗口宽度          
    wh = 400  # 窗口高度          
    x = (sw - ww) // 2          
    y = (sh - wh) // 2          
    mroot.geometry(f"{ww}x{wh}+{x}+{y}")   
  
    # 使用Frame来更好地组织控件  
    top_frame = ttk.Frame(mroot, padding="10")  
    top_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))  
    bottom_frame = ttk.Frame(mroot)  
    bottom_frame.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))  
    mroot.grid_rowconfigure(1, weight=1)  # 让TreeView占据多余空间  
    mroot.grid_columnconfigure(0, weight=1)  # 让窗口在水平方向上可伸缩  
  
    # 输入框用于填写文件名或新文件名（合并时）  
    ttk.Label(top_frame, text="表单名/新表单名:").grid(row=0, column=0, sticky=tk.W, pady=5)  
    entry = ttk.Entry(top_frame)  
    entry.grid(row=0, column=1, sticky=(tk.W, tk.E), padx=10)  


    # concat按钮（垂直合并）  
    concat_button = ttk.Button(top_frame, text="追行拼接", command=concat_files)  
    concat_button.grid(row=1, column=4, columnspan=2, pady=5)  
    # merge按钮（这里只是打印提示，实际应添加merge逻辑）  
    merge_button = ttk.Button(top_frame, text="追列拼接", command=merge_files)  
    merge_button.grid(row=1, column=2, columnspan=2, pady=5)  
      
    # 保存按钮（暂存）  
    save_button = ttk.Button(top_frame, text="暂存", command=save_file)  
    save_button.grid(row=1, column=0, columnspan=2, pady=5)  

  
    # TreeView控件显示已保存的文件列表  
    ttk.Label(bottom_frame, text="已保存的文件:").grid(row=0, column=0, sticky=tk.W, pady=5)  
    tree = ttk.Treeview(bottom_frame)  
    tree_scrollbar = ttk.Scrollbar(bottom_frame, orient="vertical", command=tree.yview)  
    tree.configure(yscrollcommand=tree_scrollbar.set)  
    tree.grid(row=1, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))  
    tree_scrollbar.grid(row=1, column=1, sticky=(tk.N, tk.S))  
    bottom_frame.grid_rowconfigure(1, weight=1)  # 让TreeView在Frame内垂直方向上可伸缩  
    bottom_frame.grid_columnconfigure(0, weight=1)  # 让TreeView在Frame内水平方向上可伸缩  
    tree.bind("<Double-1>", on_double_click)  # 绑定双击事件以显示DataFrame内容（需实现显示逻辑）  
    tree.bind("<<TreeviewSelect>>", lambda event: entry.delete(0, tk.END))  # 清除输入框当选择树中的项时（为新名称做准备）  
  
    refresh_tree()  # 初始化TreeView控件  
    mroot.mainloop()  

#统计选定的列，并且拼接在一起。    
def TOOLS_stat_all_gui(dfs):  
    df=dfs.fillna(0)
     
    # 创建主窗口  
    root = tk.Tk()  
    root.title("批量透视")  
      
    # 定义窗口的大小和位置  
    sw = root.winfo_screenwidth()  
    sh = root.winfo_screenheight()  
    ww = 400  # 窗口宽度  
    wh = 450  # 窗口高度  
    x = (sw - ww) // 2  
    y = (sh - wh) // 2  
    root.geometry(f"{ww}x{wh}+{x}+{y}")  
      
    # 定义下拉框的选项  
    aggregation_methods = ['nunique', 'sum', 'count']  
      
    # 创建标签和下拉框选择计数方法和a列  
    ttk.Label(root, text="请选择方法:").pack()  
    aggregation_var = tk.StringVar(root)  
    aggregation_var.set(aggregation_methods[0])  
    aggregation_dropdown = ttk.Combobox(root, textvariable=aggregation_var, values=aggregation_methods)  
    aggregation_dropdown.pack()  
      
    ttk.Label(root, text="请选择计数列（如：报告编码）:").pack()  
    a_column_var = tk.StringVar(root)  
    a_column_var.set(df.columns.tolist()[0] if not df.columns.empty else '')  
    a_column_dropdown = ttk.Combobox(root, textvariable=a_column_var)  
    a_column_dropdown['values'] = df.columns.tolist()  
    a_column_dropdown.pack()  

    ttk.Label(root, text="透视列:").pack()  
    pvot_column_var = tk.StringVar(root)  
    pvot_column_var.set('无需透视')  
    pvot_column_dropdown = ttk.Combobox(root, textvariable=pvot_column_var)  
    pvot_column_dropdown['values'] = df.columns.tolist()  
    pvot_column_dropdown.pack()  
   
    # 创建Treeview来显示列名及滚动条  
    tree_frame = ttk.Frame(root)  
    tree_frame.pack(side='top', fill='both', expand=True)  
      
    # Treeview  
    tree = ttk.Treeview(tree_frame, columns=("column_name",), show='headings')  
    vertical_scrollbar = ttk.Scrollbar(tree_frame, orient='vertical', command=tree.yview)  
    tree.configure(yscrollcommand=vertical_scrollbar.set)  
  
    tree.heading("#0", text="列名")  
    tree.column("#0", width=100, anchor='w')  
    tree.column("column_name", width=100, anchor='w')  
      
    # 不再需要这一行，因为我们已经在ttk.Treeview中定义了列  
    # tree["columns"] = ("column_name",)  
      
    # Pack Treeview and Scrollbar  
    tree.pack(side='left', fill='both', expand=True)  
    vertical_scrollbar.pack(side='right', fill='y')  
      
    for column in df.columns:  
        tree.insert("", 'end', text=column, values=(column,))  
  
    def calculate():  
        selected_items = tree.selection()  
        if not selected_items:  
            return  
  
        selected_columns = [tree.set(item, "column_name") for item in selected_items]  
        aggregation_method = aggregation_var.get()  
        a_column = a_column_var.get()  
        
        #add
        pvot_column=pvot_column_var.get()
        ratio_col=str(a_column)+'合计'
        
  
        if not all(col in df.columns for col in selected_columns) or a_column not in df.columns or aggregation_method not in aggregation_methods:  
            return  
  
        result_dfs = []  
        for selected_column in selected_columns: 
            df[selected_column]=df[selected_column].astype(str)
                       
            if pvot_column=="无需透视":
                result_temp=SMALL_add_count_and_ratio(df,selected_column,a_column_var.get(),aggregation_var.get())
            else:
                df[pvot_column]=df[pvot_column].astype(str) 
                result_temp=TOOLS_create_pivot_tool(df,[[selected_column], [pvot_column], [a_column_var.get()], [aggregation_var.get()], '', [ratio_col]])
            
            
            result_temp.rename(columns={selected_column: "项目"}, inplace=True)
            result_temp["列名"]=selected_column
            result_dfs.append(result_temp.copy())
        # 合并所有结果 DataFrame  
        if result_dfs:  
 
            final_result_df = pd.concat(result_dfs, ignore_index=True).reset_index(drop=True)

            PROGRAM_display_df_in_treeview(final_result_df,0,0)  
  
    # 创建按钮来执行计算  
    calculate_button = ttk.Button(root, text="执行", command=calculate)  
    calculate_button.pack()  
  
    root.mainloop()                     				

#秩和检验，包括TOOLS_rank_sum_test_cout，TOOLS_rank_sum_test两个函数#######################
def TOOLS_rank_sum_test_cout(group1, group2):
    # 将两个样本组合并
    combined = np.concatenate((group1, group2))
    
    # 对合并后的样本进行排序，并获取排序后的索引
    sorted_indices = np.argsort(combined)
    
    # 根据排序后的索引，计算每个样本的秩次
    ranks = np.empty_like(sorted_indices)
    ranks[sorted_indices] = np.arange(1, len(combined) + 1)
    
    # 分别提取两个样本组的秩次
    ranks_group1 = ranks[:len(group1)]
    ranks_group2 = ranks[len(group1):]
    
    # 计算两个样本组的秩和
    rank_sum_group1 = np.sum(ranks_group1)
    rank_sum_group2 = np.sum(ranks_group2)
    
    # 计算检验统计量U
    U = min(rank_sum_group1, rank_sum_group2)
    
    # 计算p值
    p_value = ranksums(ranks_group1, ranks_group2)[1]
    
    return U, p_value
 

def TOOLS_rank_sum_test(df):  
    if len(df.columns) < 3:  
        print("DataFrame must have at least 3 columns.")  
        return  
  
    def print_selected_values():  
        # 获取选择的列名和目标列名  
        selected_column = column_var.get()  
        target_column = target_var.get()  
        value_column = value_var.get()  
  
        # 检查选择的列是否有效  
        if selected_column == target_column:  
            print("Selected column and target column cannot be the same.")  
            return  
  
        # 从DataFrame中提取选定的列和目标列的值  
        column_values = df[selected_column].unique()  
        target_values = df[target_column].unique()  
  
        # 获取选择的两个不同元素  
        element1 = element1_var.get()  
        element2 = element2_var.get()  
  
        # 找到与选定元素对应的行，并提取目标数值列的值  
        selected_rows1 = df[df[selected_column] == element1]  
        selected_rows2 = df[df[selected_column] == element2]  
        values1 = selected_rows1[value_column].values  
        values2 = selected_rows2[value_column].values  
        result=TOOLS_rank_sum_test_cout(values1, values2)  
        # 打印提取的数值列  
        data1=f"Values for {element1} in {target_column}: {values1}\n"
        data2=f"Values for {element2} in {target_column}: {values2}\n"  
        data3="U:"+str(result[0])+"\n"
        data4="P:"+str(result[1])+"\n"       

        PROGRAM_display_content_in_textbox("秩和检验结果：\n"+data1+data2+data3+data4)
  
    def update_elements(event=None):    
        selected_column = column_var.get()    
        element1_dropdown['values'] = sorted(df[selected_column].unique())    
        element2_dropdown['values'] = sorted(df[selected_column].unique()) 
  
    # 创建GUI窗口和组件    
    root = tk.Tk()    
    root.title('秩和检验工具')    
  
    # 得到屏幕宽度和高度，并设置窗口位置和大小（这里可以自定义）          
    sw = root.winfo_screenwidth()          
    sh = root.winfo_screenheight()          
    ww = 710  # 窗口宽度          
    wh = 200  # 窗口高度          
    x = (sw - ww) // 2          
    y = (sh - wh) // 2          
    root.geometry(f"{ww}x{wh}+{x}+{y}")   
  
    # 第一行：选择比较对象所处的列    
    ttk.Label(root, text="请选择比较对象所处的列(如：通用名称)").grid(row=0, column=0, padx=5, pady=5)    
    column_var = tk.StringVar(root)    
    column_var.set(df.columns[0])  # 默认选择第一列    
    column_dropdown = ttk.Combobox(root, textvariable=column_var, values=df.columns.to_list())    
    column_dropdown.grid(row=0, column=1, padx=5, pady=5)    
    column_dropdown.bind('<<ComboboxSelected>>', update_elements)    
    
    # 第二行：选择第一行选择结果列内两个不同元素    
    ttk.Label(root, text="请选择第一个比较对象:").grid(row=1, column=0, padx=5, pady=5)    
    element1_var = tk.StringVar(root)    
    element1_dropdown = ttk.Combobox(root, textvariable=element1_var)    
    element1_dropdown.grid(row=1, column=1, padx=5, pady=5)    
    ttk.Label(root, text="请选择第二个比较对象:").grid(row=1, column=2, padx=5, pady=5)    
    element2_var = tk.StringVar(root)    
    element2_dropdown = ttk.Combobox(root, textvariable=element2_var)    
    element2_dropdown.grid(row=1, column=3, padx=5, pady=5)    
    
    # 第三行：选择比较目标列    
    ttk.Label(root, text="请选择比较目标列(如：年龄段)").grid(row=2, column=0, padx=5, pady=5)    
    target_var = tk.StringVar(root)    
    target_var.set(df.columns[0])  # 默认选择第二列    
    target_dropdown = ttk.Combobox(root, textvariable=target_var, values=df.columns.to_list())    
    target_dropdown.grid(row=2, column=1, padx=5, pady=5)    
    
    # 第四行：选择比较数值列    
    ttk.Label(root, text="请选择比较目标的数值列（如：计数）:").grid(row=3, column=0, padx=5, pady=5)    
    value_var = tk.StringVar(root)    
    value_var.set(df.columns[0])  # 默认选择第三列    
    value_dropdown = ttk.Combobox(root, textvariable=value_var, values=df.columns.to_list())    
    value_dropdown.grid(row=3, column=1, padx=5, pady=5)    
    
    # 第五行：打印按钮    
    print_button = ttk.Button(root, text="确定", command=print_selected_values)    
    print_button.grid(row=4, column=3, padx=5, pady=5)    
    
    update_elements()  # 初始化时更新一次    
    root.mainloop()  # 运行GUI窗口


    
def TOOLS_ROR_from_DB_get_abcd(df, name_char, drug_col='drugname', pt_col='pt', soc_col='soc_name', a_col='a'):  
    # 筛选出drugname包含name_char的行  
    df_filtered = df[df[drug_col].str.contains(name_char, na=False)].copy()  
  
    # 创建用于存储结果的列  
    df_filtered['b'] = 0  
    df_filtered['c'] = 0  
    df_filtered['d'] = 0  
  
    # 创建一个数据框，包含每种药物名称下 'a' 值的总和，不考虑患者类型和 SOC 名称  
    drug_totals = df.groupby(drug_col)[a_col].sum().reset_index()  
  
    # 合并药物名称的 'a' 值总和到筛选后的数据框  
    df_filtered = df_filtered.merge(drug_totals, on=drug_col, how='left', suffixes=('', '_total'))  
  
    # 计算 'b' 值  
    df_filtered['b'] = df_filtered[f'{a_col}_total'] - df_filtered[a_col]  
  
    # 计算 'c' 值  
    # 首先创建一个辅助数据框，包含每个患者类型和 SOC 名称下不同药物的 'a' 值总和  
    group_totals = df.groupby([pt_col, soc_col])[a_col].sum().reset_index()  
  
    # 然后将这个辅助数据框与筛选后的数据框合并  
    df_filtered = df_filtered.merge(group_totals, on=[pt_col, soc_col], how='left', suffixes=('', '_group_total'))  
  
    df_filtered['c'] = df_filtered[f'{a_col}_group_total'] - df_filtered[a_col]  
  
    # 确保 'c' 值不为负（根据业务逻辑调整）  
    df_filtered['c'] = df_filtered['c'].clip(lower=0)  
  
    # 计算 'd' 值  
    total_a = df[a_col].sum()  
    df_filtered['d'] = total_a - (df_filtered[a_col] + df_filtered['b'] + df_filtered['c'])  
  
    return df_filtered     
    
def TOOLS_ROR_from_DB(conn, name_field, pt_field, soc_field, pid, name_char, result_mode, display_mode=None, count_mode=None,additional_where_clause=None,abcd_value=None):  
    """    
    从SQLite数据库中高效获取数据并计算相关统计信息。    
    
    参数:    
    conn (sqlite3.Connection): 数据库连接对象。    
    name_field (str): 名称字段名。    
    pt_field (str): pt字段名。    
    soc_field (str): soc字段名。    
    pid (str): 患者ID字段名。  
    name_char (str): 要筛选的名称字符。    
    additional_where_clause (str, optional): 额外的WHERE子句条件。默认为None。  
    
    返回:    
    DataFrame: 包含name, pt, soc, a, b, c, d列的DataFrame。    
    """    
    # 构建基本的WHERE子句  
    where_clause = f"{name_field} LIKE '%{name_char}%'"  
    
  
    # 如果提供了额外的WHERE子句，则追加它  
    if additional_where_clause and additional_where_clause.strip(): 
        where_clause += f" AND ({additional_where_clause})"  
    print(where_clause)
    # 第一步：获取包含name_char的所有行，作为目标数据，用于查看原始数据   
    print("正在读取目标对象的数据作为溯源。")
    time1=time.time()
    query_target = f"""    
    SELECT    
        *    
    FROM    
        table1    
    WHERE    
        {where_clause}    
    """  

    dfs = pd.read_sql_query(query_target, conn)  
    
  
    # 第二步：计算所有药品的a（这部分代码没有变化）
    print(time.time() -time1) 
    print("正在读取数据库计算所有品种的a值...。")
    if abcd_value==None or  abcd_value=="":
        if additional_where_clause and additional_where_clause.strip():
            print("限定条件："+str(additional_where_clause)) 
            query_a = f"""    
            SELECT     
                {name_field},     
                {soc_field},     
                {pt_field},     
                COUNT(DISTINCT {pid}) AS a    
            FROM     
                table1  
            WHERE    
                {additional_where_clause}     
            GROUP BY     
                {name_field},     
                {soc_field},     
                {pt_field}    
            """  
        else:
            print("限定条件：无。") 
            query_a = f"""    
            SELECT     
                {name_field},     
                {soc_field},     
                {pt_field},     
                COUNT(DISTINCT {pid}) AS a    
            FROM     
                table1    
            GROUP BY     
                {name_field},     
                {soc_field},     
                {pt_field}    
            """       
        list_all = pd.read_sql_query(query_a, conn)    
        print(time.time() -time1) 
        # 假设这两个函数是在其他地方定义的 
        print("正在计算abcd值...")  
        result = TOOLS_ROR_from_DB_get_abcd(list_all, name_char, name_field, pt_field, soc_field)  
        print(time.time() -time1)     
    else:
        pass
    result = result.fillna(0) 
    print("正在计算ROR值...")  
    ROR_result = TOOLS_ROR_STAT_0(result)  
    print(time.time() -time1) 
    if result_mode=="a>=3&ROR_CI_95_low>1":
       ROR_result=ROR_result[(ROR_result["a"]>=3)&(ROR_result["ROR_CI_95_low"]>1)] 
       ROR_result=ROR_result.sort_values(by="ROR", ascending=False).reset_index(drop=True)
    ROR_result["报表类型"]="{'grouped':"+str([name_field, pt_field, soc_field])+str("}")
    if display_mode=="对比表":
       ROR_result=TOOLS_create_pivot_tool(ROR_result,[[soc_field, pt_field], [name_field], ['a', 'b', 'c', 'd', 'ROR', 'ROR_CI_95_low', 'PRR', 'PRR_CI_95_low'], ['sum', 'sum', 'sum', 'sum', 'sum', 'sum', 'sum', 'sum'], '', ''])
       del ROR_result["PRR合计"]       
       del ROR_result["ROR合计"]
    return ROR_result,dfs
        
def TOOLS_ROR_from_DB_GUI(df,table_name='table1'):  
    """  
    弹出一个对话框让用户选择数据库文件，并返回选择的文件路径以及过滤参数。  
    """  
    # 选择数据库文件  
    root = tk.Tk()  
    root.withdraw()  
    file_path = filedialog.askopenfilename(title="Select Database File", filetypes=[("SQLite Database", "*.db;*.sqlite;*.sqlite3")])  
    if not file_path:  
        return None, None, None, None  
    root.destroy()  
  
    # 连接到SQLite数据库  
    conn = sqlite3.connect(file_path)  
  
    # 获取字段名列表  
    cursor = conn.cursor()  
    cursor.execute(f"PRAGMA table_info({table_name})")  
    columns = [row[1] for row in cursor.fetchall()]  
  
    # 创建新窗口以选择字段和输入字符  
    root = tk.Tk()  
    root.title("ROR和PRR计算器（from数据库）")  
    sw = root.winfo_screenwidth()  
    sh = root.winfo_screenheight()  
    ww = 450  # 窗口宽度  
    wh = 500  # 窗口高度  
    x = (sw - ww) // 2  
    y = (sh - wh) // 2  
    root.geometry(f"{ww}x{wh}+{x}+{y}")  
  
    var_name = tk.StringVar(root)  
    var_pt = tk.StringVar(root)  
    var_soc = tk.StringVar(root)  
    var_char = tk.StringVar(root)  
    var_pid = tk.StringVar(root)  
    def on_ok():  
        name_field = var_name.get()  
        pt_field = var_pt.get()  
        soc_field = var_soc.get()  
        name_char = var_char.get() 
        pid=var_pid.get() 
        
        result_mode=xcode_var.get() 
        display_mode=dcode_var.get()
        count_mode=  Mcode_var.get() 
        abcd_value=abcd_var.get()
        additional_where_clause= text_box.get("1.0", tk.END)  
        # 这里调用你的函数，我假设它存在并可以接受这些参数  
        result,dfs = TOOLS_ROR_from_DB(conn, name_field, pt_field, soc_field, pid,name_char,result_mode,display_mode,count_mode,additional_where_clause,abcd_value) 
        
        PROGRAM_display_df_in_treeview(result,1,dfs)
    def add_ok(event=None):  
        condition = SQL_create_query_gui(df,text_box,conn) 

        
    # 创建并布局标签和下拉框  
    def create_labeled_dropdown(root, label_text, var, values, row):  
        label = ttk.Label(root, text=label_text)  
        label.grid(row=row, column=0, sticky=tk.W, padx=5, pady=5)  
          
        dropdown = ttk.Combobox(root, textvariable=var, values=values)  
        dropdown.grid(row=row, column=1, sticky=tk.W+tk.E, padx=5, pady=5)  
          
        return dropdown  
      
    # 参数设置标签  
    tk.Label(root, text="参数设置").grid(row=0, column=0, columnspan=2, pady=10)  
      
    # 列设置  
    var_name = tk.StringVar(root)  
    name_field_dropdown = create_labeled_dropdown(root, "计算对象所在的列（比如药品名称列）", var_name, columns, 1)  
    name_field_dropdown.set("NAME")  # 设置默认选项  
      
    var_pt = tk.StringVar(root)  
    pt_field_dropdown = create_labeled_dropdown(root, "PT:", var_pt, columns, 2)  
    pt_field_dropdown.set("PT")  # 设置默认选项  
      
    var_soc = tk.StringVar(root)  
    soc_field_dropdown = create_labeled_dropdown(root, "SOC:", var_soc, columns, 3)  
    soc_field_dropdown.set("SOC")  # 设置默认选项  
      
    var_pid = tk.StringVar(root)  
    pid_field_dropdown = create_labeled_dropdown(root, "PID:", var_pid, columns, 4)  
    pid_field_dropdown.set("PID")  # 设置默认选项  
      
    # 计算模式  
    Mcode_var = tk.StringVar(root)  
    Mcode_var.set("count")  
    Mcode_dropdown = create_labeled_dropdown(root, "计算模式:", Mcode_var, ["count"], 5)  
      
    # 信号标准  
    xcode_var = tk.StringVar(root)  
    xcode_var.set("全部")  
    xcode_dropdown = create_labeled_dropdown(root, "信号标准:", xcode_var, ["a>=3&ROR_CI_95_low>1", "全部"], 6)  
      
    # 结果展示  
    dcode_var = tk.StringVar(root)  
    dcode_var.set("详细表")  
    dcode_dropdown = create_labeled_dropdown(root, "结果展示:", dcode_var, ["对比表", "详细表"], 7)  
      
    # 模糊查询  
    tk.Label(root, text="对象范围(模糊查询):", pady=10).grid(row=8, column=0, columnspan=2, sticky=tk.W)  
    var_char = tk.StringVar(root)  
    name_char_entry = tk.Entry(root, textvariable=var_char)  
    name_char_entry.grid(row=8, column=1, columnspan=2, padx=5, pady=5)  
    # where  
    # 全局范围标签  
    #tk.Label(root, text="a+b+c+d:", pady=10).grid(row=9, column=0, sticky=tk.W)  
    abcd_var = tk.StringVar(root)  
    abcd_char_entry = tk.Entry(root, textvariable=abcd_var)  
    #abcd_char_entry.grid(row=9, column=1, columnspan=2, padx=5, pady=5)  
            
    # 创建一个带滚动条的文本框  
    where_char = tk.StringVar(root)  
    text_frame = tk.Frame(root)  # 创建一个框架来容纳文本框和滚动条  
    text_frame.grid(row=10, column=0, columnspan=2, padx=5, pady=5, sticky=tk.W+tk.E+tk.N+tk.S)  
      
    text_box = tk.Text(text_frame, height=5, width=20, wrap=tk.WORD)  # 创建一个文本框，默认3行  
    text_box.pack(side=tk.LEFT, fill=tk.BOTH, expand=True)  
    text_box.insert(tk.END,"(role_cod LIKE '%PS%')")    
    # 创建一个垂直滚动条，并与文本框关联  
    scroll_bar = ttk.Scrollbar(text_frame, command=text_box.yview)  
    scroll_bar.pack(side=tk.RIGHT, fill=tk.Y)  
      
    # 将文本框的滚动与滚动条关联  
    text_box['yscrollcommand'] = scroll_bar.set  
    # 执行按钮  
    ok_button = tk.Button(root, text="执行", command=on_ok)  
    ok_button.grid(row=11, column=1, columnspan=2, pady=10)  
    add_button = tk.Button(root, text="生成前置", command=add_ok)  
    add_button.grid(row=11, column=0, columnspan=2, pady=10)        
    root.mainloop()


#ROR计算，包括TOOLS_ROR_STAT_0，TOOLS_ROR_from_df,TOOLS_ROR_from_df_with_gui三个函数#######################
def TOOLS_ROR_STAT_0(df):  
    """  
    根据DataFrame的四个列a, b, c, d，计算TOOLS_ROR_STAT_0及置信区间。  
    直接在原始的df上添加ROR, ROR_CI_95_low, ROR_CI_95_high, PRR, PRR_CI_95_low, PRR_CI_95_high, X2, chi2, p, dof列。  
    """  
      
    def calculate_chi2(row):  
        contingency_table = [[row['a'], row['b']], [row['c'], row['d']]]  
        chi2, p, dof, expected = stats.chi2_contingency(contingency_table)  
        return chi2, p, dof  
      
    # 验证输入是否是DataFrame且包含a, b, c, d四列  
    if not isinstance(df, pd.DataFrame) or not all(col in df.columns for col in ['a', 'b', 'c', 'd']):  
        raise ValueError("输入必须是一个包含a, b, c, d四列的DataFrame")  
      
    # 计算PRR和PRR的标准误  
    df['PRR'] = (df['a'] / (df['a'] + df['b'])) / (df['c'] / (df['c'] + df['d']))  
    df['PRR_SE'] = np.sqrt(1 / df['a'] - 1 / (df['a'] + df['b']) + 1 / df['c'] - 1 / (df['c'] + df['d']))  
    df['PRR_CI_95_low'] = df['PRR'] * np.exp(-1.96 * df['PRR_SE'])  
    df['PRR_CI_95_high'] = df['PRR'] * np.exp(1.96 * df['PRR_SE'])  
      
    # 计算ROR和ROR的标准误  
    df['ROR'] = (df['a'] / df['c']) / (df['b'] / df['d'])  
    df['ROR_SE'] = np.sqrt(1 / df['a'] + 1 / df['b'] + 1 / df['c'] + 1 / df['d'])  
    df['ROR_CI_95_low'] = df['ROR'] * np.exp(-1.96 * df['ROR_SE'])  
    df['ROR_CI_95_high'] = df['ROR'] * np.exp(1.96 * df['ROR_SE'])  
    
    # 计算卡方值  
    #df['X2'] = ((df['a'] * df['d'] - df['b'] * df['c']) ** 2 * (df['a'] + df['b'] + df['c'] + df['d'])) / (  
    #    (df['a'] + df['b']) * (df['c'] + df['d']) * (df['a'] + df['c']) * (df['b'] + df['d'])  
    #)  
        
    # 使用calculate_chi2函数计算卡方值、p值和自由度，并添加到DataFrame中  
    #df[['chi2', 'chi2-p', 'chi2-dof']] = df.apply(calculate_chi2, axis=1, result_type='expand')  
      
    return df.round(2)



def TOOLS_ROR_from_df(dfs, target_column, event_column, extra_event_column, code_column,result_mode,display_mode,count_mode): 
	
    if extra_event_column=="":
        dfs["事件列分类"]="无分类"
        extra_event_column ="事件列分类" 
    df=dfs.copy()
    #df.drop_duplicates(subset=[target_column, event_column, extra_event_column, code_column], keep='first', inplace=True)
    # 生成List1  目标药品报告数量
    if count_mode=="count":
        print(count_mode)
        list1 = df.groupby(target_column)[code_column].count().reset_index()          
    else:
        list1 = df.groupby(target_column)[code_column].nunique().reset_index()  
    list1.columns = [target_column, 'a_plus_b']  
    list1['abcd'] = list1['a_plus_b'].sum()  # abcd在返回后其他地方使用  
      
    # 生成List2  所有药品目标adr报告数量
    if count_mode=="count":
        list2 = df.groupby([extra_event_column,event_column])[code_column].count().reset_index()          
    else:
        list2 = df.groupby([extra_event_column,event_column])[code_column].nunique().reset_index()  
    list2.columns = [extra_event_column,event_column,'a_plus_c']  
      
    # 生成List3  目标药品目标adr报告数量
    if count_mode=="count":
        list3 = df.groupby([target_column, extra_event_column, event_column])[code_column].count().reset_index()          
    else:
        list3 = df.groupby([target_column, extra_event_column, event_column])[code_column].nunique().reset_index()  
    list3.columns = [target_column, extra_event_column, event_column, 'a']  
      
    # 拼接List1和List3  
    list3 = pd.merge(list3, list1, on=target_column, how='left')  
      
    # 拼接List2和List3  
    list3 = pd.merge(list3, list2, on=[extra_event_column, event_column], how='left')  
      
    list3["b"] = list3['a_plus_b'] - list3["a"]  
    list3["c"] = list3['a_plus_c'] - list3["a"]      
    list3["d"] = list3['abcd'] - list3["a"] - list3["b"] - list3["c"]   
    del list3['a_plus_b']  
    del list3['a_plus_c']   
    del list3['abcd']  
    ROR_result= TOOLS_ROR_STAT_0(list3)
    if result_mode=="a>=3&ROR_CI_95_low>1":
       ROR_result=ROR_result[(ROR_result["a"]>=3)&(ROR_result["ROR_CI_95_low"]>1)] 
       ROR_result=ROR_result.sort_values(by="ROR", ascending=False).reset_index(drop=True)
    ROR_result["报表类型"]="{'grouped':"+str([target_column, event_column, extra_event_column])+str("}")
    if display_mode=="对比表":
       ROR_result=TOOLS_create_pivot_tool(ROR_result,[[extra_event_column, event_column], [target_column], ['a', 'b', 'c', 'd', 'ROR', 'ROR_CI_95_low', 'PRR', 'PRR_CI_95_low'], ['sum', 'sum', 'sum', 'sum', 'sum', 'sum', 'sum', 'sum'], '', ''])
       del ROR_result["PRR合计"]       
       del ROR_result["ROR合计"]
   
    PROGRAM_display_df_in_treeview(ROR_result, 1, dfs)  # 这个函数在原始代码中没有提供，因此注释掉  



def TOOLS_ROR_from_df_with_gui(df):   
    print("ROR计算器：您应该使用一份规整过的原始数据工作。a:目标对象目标事件报告数，b：目标对象不出现目标事件报告数，c:非目标对象目标事件报告数，d：非目标对象不出现目标事件报告数") 
    def on_submits():    
        target_col = target_var.get()    
        event_col = event_var.get()    
        extra_event_col = extra_event_var.get()  # 获取额外的事件列选择  
        code_col = code_var.get()    
        result_mode=xcode_var.get() 
        display_mode=dcode_var.get()
        count_mode=  Mcode_var.get()         
        # 在这里添加对默认选项的检查，确保用户已选择有效选项    
        if any([col == "请选择" for col in [target_var.get(), event_var.get(), extra_event_var.get(), code_var.get()]]):    
            print("请选择所有选项后再提交")    
            return    
    
        result = TOOLS_ROR_from_df(df, target_col, event_col, extra_event_col, code_col,result_mode,display_mode,count_mode)  # 调用处理函数，并传入额外的事件列参数  
            
        # 显示结果（这里只是简单地打印到控制台，您可以根据需要修改）    
        print(result)    
    
    root = tk.Tk()    
    root.title("ROR和PRR计算器")    
    sw = root.winfo_screenwidth()          
    sh = root.winfo_screenheight()          
    ww = 400  # 窗口宽度          
    wh = 450  # 窗口高度          
    x = (sw - ww) // 2          
    y = (sh - wh) // 2          
    root.geometry(f"{ww}x{wh}+{x}+{y}")    
    columns = list(df.columns)    
    target_options = ["请选择"] + columns  # 添加默认选项    
    event_options = ["请选择"] + columns  # 添加默认选项    
    extra_event_options = [""] + columns  # 添加默认选项，用于额外的事件列下拉列表  
    code_options = ["请选择"] + columns  # 添加默认选项    

    
    # 目标列下拉列表    
    target_label = ttk.Label(root, text="目标对象列(如产品名称):")    
    target_label.pack()    
    target_var = tk.StringVar(root)    
    target_var.set("")   
    target_dropdown = ttk.Combobox(root, textvariable=target_var, values=target_options)    
    target_dropdown.pack()    
    
    # 事件列下拉列表    
    event_label = ttk.Label(root, text="事件列(比如不良反应):")    
    event_label.pack()    
    event_var = tk.StringVar(root)  # 初始化为默认选项    
    event_var.set("")   
    event_dropdown = ttk.Combobox(root, textvariable=event_var, values=event_options)    
    event_dropdown.pack()    
    
    # 额外的事件列下拉列表（新增）  
    extra_event_label = ttk.Label(root, text="额外事件列(非必填，比如器官系统):")  # 新增标签和下拉列表变量等。此列与事件列在同一行。  
    extra_event_label.pack()  # 新增的列与事件列在同一行。
    extra_event_var = tk.StringVar(root)  # 初始化为默认选项    
    extra_event_var.set("")   
    extra_event_dropdown = ttk.Combobox(root, textvariable=extra_event_var, values=extra_event_options)    
    extra_event_dropdown.pack()    
    # 代码列下拉列表  
    code_label = ttk.Label(root, text="计数列（比如报告编码）:")  
    code_label.pack()  
    code_var = tk.StringVar(root)  # 初始化为默认选项  
    code_var.set("报告编码")     
    code_dropdown = ttk.Combobox(root, textvariable=code_var, values=code_options)  
    code_dropdown.pack()  
    # 计算方法  
    Mcode_label = ttk.Label(root, text="计算模式:")  
    Mcode_label.pack()  
    Mcode_var = tk.StringVar(root)  # 初始化为默认选项  
    Mcode_var.set("nunique")     
    Mcode_dropdown = ttk.Combobox(root, textvariable=Mcode_var, values=["nunique","count"])  
    Mcode_dropdown.pack()  
    # 信号标准 
    xcode_label = ttk.Label(root, text="信号标准:")  
    xcode_label.pack()  
    xcode_var = tk.StringVar(root)  # 初始化为默认选项  
    xcode_var.set("a>=3&ROR_CI_95_low>1")     
    xcode_dropdown = ttk.Combobox(root, textvariable=xcode_var, values=["a>=3&ROR_CI_95_low>1","全部"])  
    xcode_dropdown.pack()  

    # 结果展示 
    dcode_label = ttk.Label(root, text="结果展示:")  
    dcode_label.pack()  
    dcode_var = tk.StringVar(root)  # 初始化为默认选项  
    dcode_var.set("详细表")     
    dcode_dropdown = ttk.Combobox(root, textvariable=dcode_var, values=["对比表","详细表"])  
    dcode_dropdown.pack()  

    # 提交按钮  
    submit_button = ttk.Button(root, text="确定", command=on_submits)  
    submit_button.pack(pady=20)  

    separator_label = ttk.Label(root, text="采用报告数为基准。 如：(目标对象出现该事件的报告数量/其他对象出现该事件不良反应的报告数量)\n/ (目标对象不出现该事件的报告数量/其他对象不出现该事件的报告数量)")  
    separator_label.pack() 
  
    root.mainloop()

def TOOLS_merge_dataframes(df1):  
    # 创建主窗口  
    
    # 确定按钮  
    def on_ok(df1,df2,left_on,right_on):  
 
        if not left_on or not right_on:  
            messagebox.showerror("错误", '没有正确选择列。')  
            return  
  
        # 可选：检查列名是否存在于各自的DataFrame中  
        if left_on not in df1.columns or right_on not in df2.columns:  
            messagebox.showerror("错误", "没有正确选择列。")  
            return  
  
        try:  
            # 合并DataFrame  
            merged_df = pd.merge(df1, df2, left_on=left_on, right_on=right_on, how='left')  
            # 打印合并后的DataFrame（您可以替换为其他功能）  
            PROGRAM_display_df_in_treeview(merged_df,0,0)  
            # 关闭窗口  
            bsroot.destroy()  
        except Exception as e:  
            messagebox.showerror("错误", f"合并失败，原因: {e}")  
    
    bsroot = tk.Tk()  
    bsroot.title("追加合并")    
    sw = bsroot.winfo_screenwidth()          
    sh = bsroot.winfo_screenheight()          
    ww = 600  # 窗口宽度          
    wh = 150  # 窗口高度          
    x = (sw - ww) // 2          
    y = (sh - wh) // 2          
    bsroot.geometry(f"{ww}x{wh}+{x}+{y}")   
    # 隐藏主窗口直到文件对话框弹出  
    bsroot.withdraw()  
  
    # 弹出文件选择对话框，选择xlsx或xls文件  
    file_path = filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx *.xls")])  
    if not file_path:  
        messagebox.showerror("Error", "No file selected!")  
        return df1  
  
    # 读取选择的文件为DataFrame  
    try:  
        df2 = pd.read_excel(file_path)  
    except Exception as e:  
        messagebox.showerror("错误", f"Failed to read file: {e}")  
        return df1  
  
    # 显示主窗口  
    bsroot.deiconify()  
  
    # 获取df1和df2的列名  
    columns_df1 = df1.columns.to_list()  
    columns_df2 =df2.columns.to_list()  
  
    # 变量来存储选中的列名  
    selected_col_df1 = tk.StringVar(bsroot)  
    selected_col_df2 = tk.StringVar(bsroot)  
  
    # 创建并布局界面元素  
    label_df1 = tk.Label(bsroot, text="请选择主文件的合并依据列 (left_on):")  
    label_df1.grid(row=0, column=0, padx=10, pady=10)  
  
    combo_df1 = ttk.Combobox(bsroot, textvariable=selected_col_df1, values=columns_df1)  
    combo_df1.grid(row=0, column=1, padx=10)  
  
    label_df2 = tk.Label(bsroot, text="请选择需要追加文件的合并依据列 (right_on):")  
    label_df2.grid(row=1, column=0, padx=10, pady=10)  
  
    combo_df2 = ttk.Combobox(bsroot, textvariable=selected_col_df2, values=columns_df2)  
    combo_df2.grid(row=1, column=1, padx=10)  
  

  
    button_ok = tk.Button(bsroot, text="确定", command=lambda:on_ok(df1,df2.drop_duplicates(selected_col_df2.get()),selected_col_df1.get(),selected_col_df2.get()))  
    button_ok.grid(row=2, column=0, columnspan=2, pady=20)  
  
    # 取消按钮  
    def on_cancel():  
        bsroot.destroy()  
  
    button_cancel = tk.Button(bsroot, text="取消", command=on_cancel)  
    button_cancel.grid(row=2, column=1, columnspan=2, pady=10)  
  
    # 运行主循环  
    bsroot.mainloop()  


def TOOLS_drug_reaction_CH2(df, a_drug_freq_col, b_drug_freq_col, a_drug_users_col, b_drug_users_col):  
    #a_drug_users_col=8400
    #b_drug_users_col=1760
    """  
    对DataFrame中的两种药物不良反应进行卡方检验和Fisher精确检验。  
      
    参数:  
        df (pd.DataFrame): 包含不良反应数据的DataFrame。  
        a_drug_freq_col (str): a药发生频次的列名。  
        b_drug_freq_col (str): b药发生频次的列名。  
        a_drug_users_col (str): a药使用人数的列名（或固定值，如果对所有反应都一样）。  
        b_drug_users_col (str): b药使用人数的列名（或固定值，如果对所有反应都一样）。  
            
    返回:  
        pd.DataFrame: 在原DataFrame基础上增加卡方检验和Fisher精确检验的结果列，以及期望计数相关信息。  
    """  
    # 检查传入列名是否存在  
    if not all(col in df.columns for col in [a_drug_freq_col, b_drug_freq_col]):  
        raise ValueError("One or more column names are not found in the input DataFrame.")  
        
    # 尝试获取固定的用户数量，如果失败则假定是列名  
    try:  
        a_drug_users = int(df[a_drug_users_col].iloc[0]) if a_drug_users_col in df.columns else int(a_drug_users_col)  
        b_drug_users = int(df[b_drug_users_col].iloc[0]) if b_drug_users_col in df.columns else int(b_drug_users_col)  
    except ValueError:  
        raise ValueError("a_drug_users_col and b_drug_users_col must be either column names or fixed integer values.")  
        
    # 创建一个新的DataFrame，防止修改原始数据  
    new_df = df.copy()  
    new_df = new_df.fillna(0)  
    new_df[a_drug_freq_col]= new_df[a_drug_freq_col].astype(int)
    new_df[b_drug_freq_col]= new_df[b_drug_freq_col].astype(int)    
      
    # 初始化结果列表  
    results = []  
      
    # 对每一行（每一种不良反应）进行卡方检验和Fisher精确检验  
    for index, row in new_df.iterrows():  
        # 提取当前行的不良反应频次和使用人数  
        a_freq = int(row[a_drug_freq_col])  
        b_freq = int(row[b_drug_freq_col])  
            
        # 假设未使用药物的人没有不良反应  
        a_no_reaction = a_drug_users - a_freq  
        b_no_reaction = b_drug_users - b_freq  
            
        # 构建2x2的频数表  
        contingency_table = [[a_freq, a_no_reaction],  
                             [b_freq, b_no_reaction]]  
            

          
        # 卡方检验  
        chi2, p, dof, exp = chi2_contingency(contingency_table,correction=False)  
        
        
        # 计算期望计数  
        expected_counts = exp  
        expected_counts = expected_counts.reshape(2, 2)  # Ensure it's a 2x2 array  
          
        # 计算小于5的期望计数数量和百分比  
        less_than_5 = (expected_counts < 5).sum()  
        percent_less_than_5 = (less_than_5 / expected_counts.size) * 100  
          
        # 找到最小期望计数  
        min_expected_count = expected_counts.min()  

        # 卡方检验  
        Fix_chi2, Fix_p, Fix_dof, Fix_exp = chi2_contingency(contingency_table)  
  
        # Fisher精确检验  
        oddsratio, p_value = fisher_exact(contingency_table, alternative='two-sided')  
            
        # 收集结果  
        result = {  
            'chi2': chi2,  
            'chi2_p': p,  
            'chi2_dof': dof,  
            'chi2_exp':exp,
            'Fix_chi2': Fix_chi2,  
            'Fix_p': Fix_p,  
            'Fix_dof': Fix_dof,  
            'Fix_exp':Fix_exp,            
            'fisher_oddsratio': oddsratio,  
            'fisher_p_value': p_value,  
            'less_than_5_count': less_than_5,  
            'percent_less_than_5': percent_less_than_5,  
            'min_expected_count': min_expected_count  
        }  
        results.append(result)  
        
    # 将结果添加到DataFrame中  
    new_df = pd.concat([new_df, pd.DataFrame(results, index=new_df.index)], axis=1)  
        
    return new_df


def TOOLS_drug_reaction_CH2_create_gui(df):  
    print("若每个条件下的期望频数均大于5，且总样本量大于40,则可使用Pearson卡方检验，否则需使用相应的矫正方法。\nYates矫正卡方检验适用于两个分类变量均只有两个水平，即两个分类变量共能组成2*2个条件的情况下，且存在某个条件的期望频数小于5但大于1时，此时应使用Yates矫正卡方检验。\nFisher精确卡方检验适用于样本量小于40，或者两个分类变量的各个条件下的期望频数存在小于1的情况，此时应使用Fisher精确检验。")
    # 创建主窗口  
    root = tk.Tk()  
    root.title("卡方检验和fisher精确检验") 
    
    sw = root.winfo_screenwidth()    
    sh = root.winfo_screenheight()    
    ww = 300  # 窗口宽度          
    wh = 220  # 窗口高度          
    x = (sw - ww) // 2          
    y = (sh - wh) // 2          
    root.geometry(f"{ww}x{wh}+{x}+{y}")    
    # 创建用于选择列名的下拉菜单  
    reaction_label = ttk.Label(root, text="事件列:")  
    reaction_label.grid(row=0, column=0, padx=5, pady=5)  
    reaction_var = tk.StringVar(root)  
    reaction_var.set(df.columns[0])  # 设置默认选项  
    reaction_dropdown = ttk.Combobox(root, textvariable=reaction_var, state='readonly')  
    reaction_dropdown['values'] = df.columns.tolist()   
    reaction_dropdown.grid(row=0, column=1, padx=5, pady=5)  
  
    a_drug_label = ttk.Label(root, text="对象A阳性频次:")  
    a_drug_label.grid(row=1, column=0, padx=5, pady=5)  
    a_drug_var = tk.StringVar(root)  
    a_drug_var.set(df.columns[3])  # 设置默认选项  
    a_drug_dropdown = ttk.Combobox(root, textvariable=a_drug_var, state='readonly')  
    a_drug_dropdown['values'] = df.columns.tolist()   
    a_drug_dropdown.grid(row=1, column=1, padx=5, pady=5)  
  
    b_drug_label = ttk.Label(root, text="对象B阳性频次:")  
    b_drug_label.grid(row=2, column=0, padx=5, pady=5)  
    b_drug_var = tk.StringVar(root)  
    b_drug_var.set(df.columns[4])  # 设置默认选项  
    b_drug_dropdown = ttk.Combobox(root, textvariable=b_drug_var, state='readonly')  
    b_drug_dropdown['values'] = df.columns.tolist()  
    b_drug_dropdown.grid(row=2, column=1, padx=5, pady=5)  
  
    # 用户数量是固定的，所以我们使用Entry来输入  
    a_users_label = ttk.Label(root, text="对象A基数:")  
    a_users_label.grid(row=3, column=0, padx=5, pady=5)  
    a_users_entry = ttk.Entry(root)  
    a_users_entry.grid(row=3, column=1, padx=5, pady=5)  
  
    b_users_label = ttk.Label(root, text="对象B基数:")  
    b_users_label.grid(row=4, column=0, padx=5, pady=5)  
    b_users_entry = ttk.Entry(root)  
    b_users_entry.grid(row=4, column=1, padx=5, pady=5)  
  
    # 创建一个按钮来执行统计函数  
    calculate_button = ttk.Button(root, text="统计", command=lambda: PROGRAM_display_df_in_treeview(TOOLS_drug_reaction_CH2(  
        df,a_drug_var.get(), b_drug_var.get(), a_users_entry.get(), b_users_entry.get()),0,0))  
    calculate_button.grid(row=5, column=0, columnspan=2, padx=5, pady=10)  
  
    # 运行主循环  
    root.mainloop() 









    
#数据分组和透视，TOOLS_create_pivot_tool_gui#######################    
def TOOLS_create_pivot_create_multiselect_pivot_gui(df, result_text):  
    """  
    创建一个新的GUI，用于多选值透视功能  
    """  
    def on_confirm():  
        """  
        确认按钮的回调函数，用于获取用户的选择并返回  
        """  
        # 获取用户选择的列名  
        selected_ids = tree.selection()  
        selected_col_names = [tree.item(item)["text"] for item in selected_ids]  
        selected_method = method_var.get()  
  
        # 创建一个字典，其中key是列名，value是选定的方法  
        result_dict = {col_name: selected_method for col_name in selected_col_names}  
  
        # 将字典转换为字符串并插入到result_text中  
        result_str = str(result_dict)  
        result_text.insert(tk.END, result_str)  
  
        # 关闭当前窗口  
        top.destroy()  
  
    # 创建新的顶层窗口  
    top = tk.Toplevel()  
    top.title("多选值透视工具")  
    # 设置窗口位置和大小（这里可以自定义）    
    sw = top.winfo_screenwidth()    
    sh = top.winfo_screenheight()    
    ww = 320  # 窗口宽度    
    wh = 300  # 窗口高度    
    x = (sw - ww) // 2    
    y = (sh - wh) // 2    
    top.geometry(f"{ww}x{wh}+{x}+{y}")   
    # 在窗口中创建Treeview，显示df的所有列  
    tree = ttk.Treeview(top)  
    ysb = ttk.Scrollbar(top, orient='vertical', command=tree.yview)  # 创建垂直滚动条  
    xsb = ttk.Scrollbar(top, orient='horizontal', command=tree.xview)  # 创建水平滚动条  
    tree.configure(yscrollcommand=ysb.set, xscrollcommand=xsb.set)  # 将滚动条的视图与Treeview绑定  
  
    tree.column("#0", width=200, minwidth=200, stretch=tk.NO)  
    tree.heading("#0", text="列名", anchor=tk.W)  
    for i, col in enumerate(df.columns):  
        tree.insert("", i, text=col)  # 删除values参数以修复Treeview显示问题  
  
    # 安排Treeview和滚动条的位置  
    tree.grid(row=0, column=0, sticky='nsew')  
    ysb.grid(row=0, column=1, sticky='ns')  
    xsb.grid(row=1, column=0, sticky='ew')  
  
    top.grid_rowconfigure(0, weight=1)  
    top.grid_columnconfigure(0, weight=1)  
  
    # 创建方法选择的下拉菜单  
    methods = ["sum", "count", "nunique", "unique", "mean", "median", "std", "size", "cumsum", "SMALL_count_mode", "count_CH_semicolon", "count_EN_semicolon", "count_CH_comma", "count_EN_comma", "count_CH_commas", "count_ALL"]  
    method_var = tk.StringVar(top)  
    method_var.set(methods[0])  # 设置默认方法  
    method_dropdown = ttk.Combobox(top, textvariable=method_var, values=methods)  
    method_dropdown.grid(row=2, column=0 ) 
  
    # 创建确认按钮，并绑定回调函数on_confirm  
    confirm_button = ttk.Button(top, text="确认", command=on_confirm)  
    confirm_button.grid(row=3, column=0)
  
    # 运行新窗口的主循环，等待用户操作  
    top.wait_window()
      
def TOOLS_create_pivot_tool(df,methon): 
    """    
    数据透视工具的的功能函数   
    """
    row_labels=methon[0]
    col_labels=methon[1] 
    value_cols=methon[2] 
    agg_methods=methon[3] 
    text_content=methon[4] 
    all_ratio=methon[5] 


    
    try:
        for i, value in enumerate(agg_methods):  
            if value == "SMALL_count_mode":  
                agg_methods[i] = lambda x: SMALL_count_mode(x, "")  
            elif value == "count_CH_semicolon":  
                agg_methods[i] = lambda x: SMALL_count_mode(x, "；")  
            elif value == "count_EN_semicolon":  
                agg_methods[i] = lambda x: SMALL_count_mode(x, ";")  
            elif value == "count_CH_comma":  
                agg_methods[i] = lambda x: SMALL_count_mode(x, "，")  
            elif value == "count_EN_comma":  
                agg_methods[i] = lambda x: SMALL_count_mode(x, ",")  
            elif value == "count_CH_commas":  
                agg_methods[i] = lambda x: SMALL_count_mode(x, "、")  
            elif value == "count_ALL":  
                agg_methods[i] = lambda x: SMALL_count_mode(x, "。|；|，|、|,|;|/")
    except:
        print("SMALL_count_mode error")
        
    try:
        for i, value in text_content.items():  
            if value == "SMALL_count_mode":  
                text_content[i] = lambda x: SMALL_count_mode(x, "")  
            elif value == "count_CH_semicolon":  
                text_content[i] = lambda x: SMALL_count_mode(x, "；")  
            elif value == "count_EN_semicolon":  
                text_content[i] = lambda x: SMALL_count_mode(x, ";")  
            elif value == "count_CH_comma":  
                text_content[i] = lambda x: SMALL_count_mode(x, "，")  
            elif value == "count_EN_comma":  
                text_content[i] = lambda x: SMALL_count_mode(x, ",")  
            elif value == "count_CH_commas":  
                text_content[i] = lambda x: SMALL_count_mode(x, "、")  
            elif value == "count_ALL":  
                text_content[i] = lambda x: SMALL_count_mode(x, "。|；|，|、|,|;|/")
    except:
        pass 
    pivot_table = pd.pivot_table(df, values=value_cols, index=row_labels, columns=col_labels, aggfunc={col: agg for col, agg in zip(value_cols, agg_methods)},margins=True, margins_name='合计') .reset_index() 
    #解除多级索引  
    pivot_table.columns = [''.join(col).strip() for col in pivot_table.columns.values] 	           
    if len(text_content)>=1:
        try:
            pivot_table1 = pd.pivot_table(df, values=text_content.keys(), index=row_labels, columns=[], aggfunc=text_content,margins=True, margins_name='合计') .reset_index() 
            pivot_table1.columns = [''.join(col).strip() for col in pivot_table1.columns.values] 	
            pivot_table = pd.pivot_table(df, values=value_cols, index=row_labels, columns=col_labels, aggfunc={col: agg for col, agg in zip(value_cols, agg_methods)},margins=True, margins_name='合计') .reset_index() 
            pivot_table.columns = [''.join(col).strip() for col in pivot_table.columns.values] 
            #del pivot_table1["合计"]
            pivot_table=pd.merge(pivot_table,pivot_table1, on=row_labels, how='left')    

        except dfd:
            print("多选值列配置错误。") 
    #增加构成比
    if len(all_ratio)!=0: 
        pivot_table=SMALL_add_composition_ratio(pivot_table,all_ratio[0])  #SMALL_calculate_ratios
    #增加严重比、超时比等。
    if len(all_ratio)==3: 
        pivot_table=SMALL_calculate_ratios(pivot_table,all_ratio[1],all_ratio[2])		               
    pivot_table["报表类型"]="{'grouped':"+str(row_labels)+str("}")  
    return pivot_table
        
def TOOLS_create_pivot_tool_gui(df, ori):    
    """    
    创建数据透视工具的图形用户界面(GUI)    
    
    参数:    
        df (DataFrame): 输入的数据框    
        ori (object): 原始对象，用于后续处理结果    
    """    
    
    def on_submit():    
        """    
        提交按钮的回调函数，用于处理用户的选择并执行数据透视操作    
        """    
        # 获取用户选择的行标签、列标签、值列和聚合方法    
        row_labels = [var.get() for var in row_vars if var.get() != ""]    
        col_labels = [var.get() for var in col_vars if var.get() != ""]    
        value_cols = [var.get() for var in value_vars if var.get() != ""]    
        agg_methods = [var.get() for var in agg_method_vars if var.get() != ""]        
        

   
        # 创建数据透视表，使用字典指定每个值列的聚合方法    
        
        text_content=""
        try:
           text_content = eval(result_text.get("1.0", tk.END)) 
        except:
            pass
        methon=[row_labels,col_labels,value_cols,agg_methods,text_content,""]
        print(methon)
        pivot_table=TOOLS_create_pivot_tool(df,methon)
        PROGRAM_display_df_in_treeview(pivot_table, 0, ori)    

  
    
    # 创建主窗口    
    root = tk.Tk()    
    root.title("数据分组和透视工具")  # 设置窗口标题为“数据透视工具”    
    
    # 设置窗口位置和大小（这里可以自定义）    
    sw = root.winfo_screenwidth()    
    sh = root.winfo_screenheight()    
    ww = 820  # 窗口宽度    
    wh = 300  # 窗口高度    
    x = (sw - ww) // 2    
    y = (sh - wh) // 2    
    root.geometry(f"{ww}x{wh}+{x}+{y}")    
    
    # 获取数据框的列名列表    
    columnslist = df.columns.to_list()    
    
    # 创建行标签的下拉菜单    
    row_vars = []    
    row_labels = ["行标签1", "行标签2", "行标签3", "行标签4", "行标签5"]    
    for i, label in enumerate(row_labels):    
        tk.Label(root, text=label).grid(row=0, column=i)    
        var = tk.StringVar(root)    
        var.set("")    
        dropdown = ttk.Combobox(root, textvariable=var, values=columnslist)    
        dropdown.grid(row=1, column=i)    
        row_vars.append(var)    
    
    # 创建列标签的下拉菜单    
    col_vars = []    
    col_labels = ["透视列1", "透视列2", "透视列3", "透视列4", "透视列5"]    
    for i, label in enumerate(col_labels):    
        tk.Label(root, text=label).grid(row=2, column=i)    
        var = tk.StringVar(root)    
        var.set("")    
        dropdown = ttk.Combobox(root, textvariable=var, values=columnslist)    
        dropdown.grid(row=3, column=i)    
        col_vars.append(var)    
    
    # 创建值列的下拉菜单和对应的聚合方法的下拉菜单    
    value_vars = []    
    agg_method_vars = []    
    value_labels = ["值列1", "值列2", "值列3", "值列4", "值列5"]    
    for i, label in enumerate(value_labels):    
        tk.Label(root, text=label).grid(row=4, column=i)    
        var = tk.StringVar(root)    
        var.set("")    
        dropdown = ttk.Combobox(root, textvariable=var, values=columnslist)    
        dropdown.grid(row=5, column=i)    
        value_vars.append(var)    
        # 创建聚合方法的下拉菜单    
        methods = ["sum", "count", "nunique", "unique", "mean", "median", "std", "size", "cumsum", "SMALL_count_mode", "count_CH_semicolon", "count_EN_semicolon", "count_CH_comma", "count_EN_comma", "count_CH_commas", "count_ALL"]  # 可根据需要添加其他聚合方法    

        agg_var = tk.StringVar(root)    
        agg_var.set("")  # 设置默认值为第一个聚合方法    
        agg_dropdown = ttk.Combobox(root, textvariable=agg_var, values=methods)    
        agg_dropdown.grid(row=6, column=i)  # 将下拉菜单放置到网格中，与对应的值列下拉菜单对齐显示    
        agg_method_vars.append(agg_var)    
    
    # 创建提交按钮，并绑定回调函数on_submit，用于处理用户的选择和执行数据透视操作    
    submit_button = ttk.Button(root, text="提交", command=on_submit)  # 创建提交按钮并设置文本和回调函数    
    submit_button.grid(row=7, column=0, columnspan=3)  # 将提交按钮放置到网格中，并跨3列显示    
    multiselect_button = ttk.Button(root, text="拼接数据分组", command=lambda: TOOLS_create_pivot_create_multiselect_pivot_gui(df,result_text))  # 创建按钮并绑定功能函数create_multiselect_pivot_gui()      
    multiselect_button.grid(row=7, column=3, columnspan=3)  # 将按钮放置到网格中，并跨3列显示      
    # 运行主循环，显示窗口并等待用户操作  
    # 创建文本框，用于显示多选值透视的结果  
    result_text = tk.Text(root, height=4, width=100, wrap=tk.WORD)  
    # 创建垂直滚动条  
    scrollbar = tk.Scrollbar(root, command=result_text.yview)     
    # 将滚动条与文本框关联  
    result_text['yscrollcommand'] = scrollbar.set  
    # 将文本框放置到网格中，并跨5列显示  
    result_text.grid(row=8, column=0, columnspan=5, sticky='nsew')  
    # 将滚动条放置在文本框右侧，与文本框同高  
    scrollbar.grid(row=8, column=5, rowspan=4, sticky='ns')  
    root.mainloop() #启动Tkinter事件循环，显示窗口并等待用户操作或关闭窗口时结束程序
    
#趋势分析函数，TOOLS_trend_analysis_GUI，,TOOLS_trend_analysis_with_3_sd,TOOLS_trend_analysis_with_3_sd#######################    
def TOOLS_trend_analysis_GUI(df):  
    # 创建Tkinter窗口  
    # 更新分析对象下拉菜单的去重清单  
    def update_unique_values(event):  
        selected_column = target_column_var.get()  
        if selected_column == "不筛选":  
            analysis_object_var.set("不筛选")  
            analysis_object_dropdown['values'] = []  
            return  
          
        unique_values = df[selected_column].unique()  
        sorted_values = sorted(unique_values, key=lambda x: df[selected_column].value_counts()[x], reverse=True)  
          
        # 格式化所有的唯一值，包括值、计数和百分比  
        formatted_values = [str([value, df[selected_column].value_counts()[value], f"{df[selected_column].value_counts(normalize=True)[value]*100:.2f}%"]) for value in sorted_values]


          
        # 更新分析对象变量和下拉菜单的值  
        analysis_object_var.set(formatted_values[0] if formatted_values else [])  
        analysis_object_dropdown['values'] = formatted_values  

     
    
    
    root = tk.Tk()  
    root.title("趋势分析预处理")  
      
    # 得到屏幕宽度和高度，并设置窗口位置和大小（这里可以自定义）  
    sw = root.winfo_screenwidth()  
    sh = root.winfo_screenheight()  
    ww = 350  # 窗口宽度  
    wh = 250  # 窗口高度  
    x = (sw - ww) // 2  
    y = (sh - wh) // 2  
    root.geometry(f"{ww}x{wh}+{x}+{y}")  
      
    # 获取df的列名以供选择  
    columns = df.columns.tolist()  
      
    # 设置默认参数  
    default_date_column = "事件发生日期" if "事件发生日期" in columns else columns[1]  
    default_event_column = "报告编码" if "报告编码" in columns else columns[1]  
    default_windows = 12  
    default_method = "nunique"  
    default_freq = "M"  
      
    # 创建标签和输入框  
    mainframe = ttk.Frame(root, padding="10")  
    mainframe.grid(column=0, row=0, sticky=(tk.N, tk.W, tk.E, tk.S))  
      
    date_label = ttk.Label(mainframe, text="日期列:")  
    date_label.grid(column=1, row=1, sticky=tk.W)  
    date_column_var = tk.StringVar(root)  
    date_column_var.set(default_date_column)  # 初始设置  
    date_column_dropdown = ttk.Combobox(mainframe, textvariable=date_column_var, values=columns, width=30)  # 设置宽度为30个字符  
    date_column_dropdown.grid(column=2, row=1, sticky=(tk.W, tk.E))  
      
    event_label = ttk.Label(mainframe, text="报告编号列:")  
    event_label.grid(column=1, row=2, sticky=tk.W)  
    event_column_var = tk.StringVar(root)  
    event_column_var.set(default_event_column)  # 初始设置  
    event_column_dropdown = ttk.Combobox(mainframe, textvariable=event_column_var, values=columns, width=30)  # 设置宽度为30个字符  
    event_column_dropdown.grid(column=2, row=2, sticky=(tk.W, tk.E))  
      
    # 分析目标列下拉菜单  
    target_label = ttk.Label(mainframe, text="目标筛选列:")  
    target_label.grid(column=1, row=3, sticky=tk.W)  
    target_column_var = tk.StringVar(root)  
    target_column_var.set("不筛选")  # 设置默认值为“不筛选”  
    target_column_dropdown = ttk.Combobox(mainframe, textvariable=target_column_var, values=["不筛选"] + columns, width=30)  # 设置宽度为30个字符  
    target_column_dropdown.grid(column=2, row=3, sticky=(tk.W, tk.E))  
    target_column_dropdown.bind("<<ComboboxSelected>>", update_unique_values)  # 绑定事件，当选定列时更新分析对象下拉菜单  
      
    # 分析对象下拉菜单（显示选定列的去重清单）  
    analysis_object_label = ttk.Label(mainframe, text="目标对象:")  
    analysis_object_label.grid(column=1, row=4, sticky=tk.W)  
    analysis_object_var = tk.StringVar(root)  
    analysis_object_dropdown = ttk.Combobox(mainframe, textvariable=analysis_object_var, width=30)  # 设置宽度为30个字符  
    analysis_object_dropdown.grid(column=2, row=4, sticky=(tk.W, tk.E))  
    update_unique_values(None)  # 初始化时更新分析对象下拉菜单  


    #分析窗口  
    windows_label = ttk.Label(mainframe, text="分析窗口:")  
    windows_label.grid(column=1, row=5, sticky=tk.W)  
    windows_entry = ttk.Entry(mainframe)  
    windows_entry.insert(0, str(default_windows))  
    windows_entry.grid(column=2, row=5, sticky=(tk.W, tk.E))  
  
    method_label = ttk.Label(mainframe, text="方法:")  
    method_label.grid(column=1, row=6, sticky=tk.W)  
    method_var = tk.StringVar(root)  
    method_var.set(default_method)  # 初始设置  
    method_dropdown = ttk.Combobox(mainframe, textvariable=method_var, values=["nunique", "count", "sum"])  # 下拉菜单选择方法  
    method_dropdown.grid(column=2, row=6, sticky=(tk.W, tk.E))  
  
    freq_label = ttk.Label(mainframe, text="频率:")  
    freq_label.grid(column=1, row=7, sticky=tk.W)  
    freq_var = tk.StringVar(root)  
    freq_var.set(default_freq)  # 初始设置  
    freq_dropdown = ttk.Combobox(mainframe, textvariable=freq_var, values=["M", "Q"])  # 下拉菜单选择频率  
    freq_dropdown.grid(column=2, row=7, sticky=(tk.W, tk.E))  

    # 添加控制限选择下拉框  
    control_limit_label = ttk.Label(mainframe, text="控制限:")  
    control_limit_label.grid(column=1, row=8, sticky=tk.W)  
    control_limit_var = tk.StringVar(root)  
    control_limit_var.set("标准差")  # 默认设置为标准差  
    control_limit_dropdown = ttk.Combobox(mainframe, textvariable=control_limit_var, values=["标准差", "IQR"])  
    control_limit_dropdown.grid(column=2, row=8, sticky=(tk.W, tk.E))  
  
    # 创建确认按钮，用于启动趋势分析函数 
     
    #趋势图
    def confirm_1():  
        # 获取用户输入的参数值，并进行趋势分析函数调用  
        date_column = date_column_var.get()  # 获取日期列名称  
        event_column = event_column_var.get()  # 获取事件列名称  
        windows = int(windows_entry.get()) or default_windows  # 获取分析周期，若为空则使用默认值（需转换为整数）  
        method = method_var.get()  # 获取方法（计数或求和）  
        freq = freq_var.get()  # 获取数据聚合频率（月或季度）  
        target_column = target_column_var.get()  
        analysis_object=analysis_object_var.get()

        # 根据target_column的值决定传递哪个df  
        if target_column == "不筛选":  
            df_to_use = df  
        else: 
            
            pattern = r"'(.*?)'"  # 匹配单引号内的任意文本 
            analysis_object = re.findall(pattern,analysis_object) 
            df_to_use = df[df[target_column] == analysis_object[0]]  # 假设这是筛选逻辑  


      
        if control_limit_var.get() == "标准差":  
            TOOLS_trend_analysis_with_3_sd(df_to_use, date_column, event_column, windows, method, "draw", freq)  # 调用趋势分析函数，并传递参数进行趋势分析。注意：这里假设trend函数可以接受这些参数并正确处理。如果不能，则需要在trend函数内部进行适当修改。  
        elif control_limit_var.get() == "IQR":  
            TOOLS_trend_analysis_with_1_5IQR(df_to_use, date_column, event_column, windows, method, "draw", freq)  # 调用趋势分析函数，并传递参数进行趋势分析。注意：这里假设trend函数可以接受这些参数并正确处理。如果不能，则需要在trend函数内部进行适当修改。  
      
        root.destroy()  # 关闭窗口  
      
    # 趋势表  
    def confirm_2():  
        # 获取用户输入的参数值，并进行趋势分析函数调用  
        date_column = date_column_var.get()  # 获取日期列名称  
        event_column = event_column_var.get()  # 获取事件列名称  
        windows = int(windows_entry.get()) or default_windows  # 获取分析周期，若为空则使用默认值（需转换为整数）  
        method = method_var.get()  # 获取方法（计数或求和）  
        freq = freq_var.get()  # 获取数据聚合频率（月或季度）  
        target_column = target_column_var.get()  
        analysis_object=analysis_object_var.get()

        # 根据target_column的值决定传递哪个df  
        if target_column == "不筛选":  
            df_to_use = df  
        else: 
            pattern = r"'(.*?)'"  # 匹配单引号内的任意文本 
            analysis_object = re.findall(pattern,analysis_object)  

			          
            df_to_use = df[df[target_column] == analysis_object[0]]  # 假设这是筛选逻辑  

      
        if control_limit_var.get() == "标准差":  
            PROGRAM_display_df_in_treeview(TOOLS_trend_analysis_with_3_sd(df_to_use, date_column, event_column, windows, method, "data", freq), 1, 0)  # 调用趋势分析函数，并传递参数进行趋势分析。注意：这里假设trend函数可以接受这些参数并正确处理。如果不能，则需要在trend函数内部进行适当修改。  
        elif control_limit_var.get() == "IQR":  
            PROGRAM_display_df_in_treeview(TOOLS_trend_analysis_with_1_5IQR(df_to_use, date_column, event_column, windows, method, "data", freq), 1, 0)  # 调用趋势分析函数，并传递参数进行趋势分析。注意：这里假设trend函数可以接受这些参数并正确处理。如果不能，则需要在trend函数内部进行适当修改。  
      
        root.destroy()  # 关闭窗口   
    confirm_button = ttk.Button(mainframe, text="趋势图", command=confirm_1)  # 创建确认按钮，用于启动趋势分析函数，看图  
    confirm_button.grid(column=2, row=9, sticky=(tk.W, tk.E))  
    confirm_button = ttk.Button(mainframe, text="趋势表", command=confirm_2)  # 创建确认按钮，用于启动趋势分析函数，看表  
    confirm_button.grid(column=1, row=9, sticky=(tk.W, tk.E))    
    # 运行Tkinter主循环，等待用户操作  
    root.mainloop()
    
def TOOLS_trend_analysis_with_3_sd(df, date_column, event_column, windows, method,draw_or_data, freq='M'):  
    """  
    df: pandas DataFrame, 包含日期和事件的数据  
    date_column: str, 日期列的名称  
    event_column: str, 事件列的名称  ，比如报告表编码
    windows: 分析周期  
    method: 'count' 或 'sum'  或 'unique'
    freq: 数据聚合频率，'M' 表示按月，'Q' 表示按季度  
    draw_or_data:draw绘图，data返回数据
    """  
    # 确保日期列为datetime，并按日期排序  
    df.loc[:, date_column] = pd.to_datetime(df[date_column]) 
    df = df.sort_values(by=date_column)  
  
    # 按指定频率对事件进行分组和计数/求和，即使数据为0也要纳入  
    if method == "count":  
        monthly_counts = df.groupby(pd.Grouper(key=date_column, freq=freq))[event_column].count()  
    elif method == "sum":  
        monthly_counts = df.groupby(pd.Grouper(key=date_column, freq=freq))[event_column].sum()  
    elif method == "nunique":  
        monthly_counts = df.groupby(pd.Grouper(key=date_column, freq=freq))[event_column].nunique()  
    else:  
        raise ValueError("Invalid method. Please use 'count', 'sum', or 'unique'.")  
  
    monthly_counts = monthly_counts.resample(freq).asfreq().fillna(0)  # Resample to ensure monthly/quarterly data, fill NaNs with 0  
  
    # 计算移动平均线和标准差控制限  
    rolling_mean = monthly_counts.rolling(window=windows).mean()  
    rolling_std = monthly_counts.rolling(window=windows).std()  
    UCL_2sd = rolling_mean + 2 * rolling_std  
    LCL_2sd = rolling_mean - 2 * rolling_std  
    UCL_3sd = rolling_mean + 3 * rolling_std  
    LCL_3sd = rolling_mean - 3 * rolling_std  
  
    result_df = pd.DataFrame({  
        'Date': monthly_counts.index,  
        'Monthly Counts': monthly_counts.values,  
        'Rolling Mean': rolling_mean.values,  
        'UCL (2 SD)': UCL_2sd.values,  
        'LCL (2 SD)': LCL_2sd.values,  
        'UCL (3 SD)': UCL_3sd.values,  
        'LCL (3 SD)': LCL_3sd.values  
    })  
    if draw_or_data=="data": 
        return result_df    

    # 创建Tkinter窗口和matplotlib图形    
    root = tk.Tk()    
    root.title("Trend Analysis")    
    # 得到屏幕宽度和高度，并设置窗口位置和大小（这里可以自定义）      
    sw = root.winfo_screenwidth()      
    sh = root.winfo_screenheight()      
    ww = 1300  # 窗口宽度      
    wh = 700  # 窗口高度      
    x = (sw - ww) // 2      
    y = (sh - wh) // 2      
    root.geometry(f"{ww}x{wh}+{x}+{y}")    
    fig, ax = plt.subplots(figsize=(10, 6))    
        
    # 解决汉字乱码问题    
    plt.rcParams["font.sans-serif"] = ["SimHei"]  # 使用指定的汉字字体类型（此处为黑体）    
    plt.rcParams['axes.unicode_minus'] = False  # 用来正常显示负号    
        
    # 绘制趋势图    
    ax.plot(result_df['Date'], result_df['Monthly Counts'], marker='o', label='Monthly Counts')    
    ax.plot(result_df['Date'], result_df['Rolling Mean'], color='red', label='Rolling Mean ({} Months)'.format(windows))    
    ax.plot(result_df['Date'], result_df['UCL (2 SD)'], color='blue', linestyle='--', label='UCL (2 SD)')  # 2 SD line in blue    
    ax.plot(result_df['Date'], result_df['LCL (2 SD)'], color='blue', linestyle='--')  # 2 SD line in blue    
    ax.plot(result_df['Date'], result_df['UCL (3 SD)'], color='green', linestyle='--', label='UCL (3 SD)')  # 3 SD line in green    
    ax.plot(result_df['Date'], result_df['LCL (3 SD)'], color='green', linestyle='--')  # 3 SD line in green    
    ax.set_xlabel('Date')    
    ax.set_ylabel('Event Counts')    
    ax.set_title('Trend Analysis with Control Limits (2 and 3 Standard Deviations)')    
    ax.legend()    
    ax.grid(True)    
        
    # 使用FigureCanvasTkAgg将图形嵌入到Tkinter窗口中    
    canvas = FigureCanvasTkAgg(fig, master=root)    
    canvas.draw()    
    canvas.get_tk_widget().pack(side=tk.TOP, fill=tk.BOTH, expand=1)    
        
    # 创建和显示工具条  
    toolbar = NavigationToolbar2Tk(canvas, root)  
    toolbar.update()  
    canvas.get_tk_widget().pack()  

    # 运行Tkinter事件循环  
    root.mainloop()

def TOOLS_trend_analysis_with_1_5IQR(df, date_column, event_column, windows, method, draw_or_data, freq='M'):  
    """  
    df: pandas DataFrame, 包含日期和事件的数据  
    date_column: str, 日期列的名称  
    event_column: str, 事件列的名称  
    windows: 分析周期  
    method: 'count' 或 'sum'  或 'unique'  
    freq: 数据聚合频率，'M' 表示按月，'Q' 表示按季度  
    draw_or_data:draw绘图，data返回数据  
    """  
    # 确保日期列为datetime，并按日期排序  
    df.loc[:, date_column] = pd.to_datetime(df[date_column])  
    df = df.sort_values(by=date_column)  
  
    # 按指定频率对事件进行分组和计数/求和，即使数据为0也要纳入  
    if method == "count":  
        monthly_counts = df.groupby(pd.Grouper(key=date_column, freq=freq))[event_column].count()  
    elif method == "sum":  
        monthly_counts = df.groupby(pd.Grouper(key=date_column, freq=freq))[event_column].sum()  
    elif method == "nunique":  
        monthly_counts = df.groupby(pd.Grouper(key=date_column, freq=freq))[event_column].nunique()  
    else:  
        raise ValueError("Invalid method. Please use 'count', 'sum', or 'unique'.")  
  
    monthly_counts = monthly_counts.resample(freq).asfreq().fillna(0)  # Resample to ensure monthly/quarterly data, fill NaNs with 0  
  
    # 计算移动平均线和IQR控制限  
    rolling_mean = monthly_counts.rolling(window=windows).mean()  
    rolling_25p = monthly_counts.rolling(window=windows).quantile(0.25)  # 25th percentile  
    rolling_50p = monthly_counts.rolling(window=windows).quantile(0.5)  # 50th percentile (median)  
    rolling_75p = monthly_counts.rolling(window=windows).quantile(0.75)  # 75th percentile  
    IQR = rolling_75p - rolling_25p  # Interquartile range  
    UCL = rolling_75p + 1.5 * IQR  # Upper control limit (1.5 IQR above 75th percentile)  
    LCL = rolling_25p - 1.5 * IQR  # Lower control limit (1.5 IQR below 25th percentile)  
  
    result_df = pd.DataFrame({  
        'Date': monthly_counts.index,  
        'Monthly Counts': monthly_counts.values,  
        'Rolling Mean': rolling_mean.values,  
        'Rolling 25th Percentile': rolling_25p.values,  
        'Rolling Median': rolling_50p.values,  
        'Rolling 75th Percentile': rolling_75p.values,  
        'UCL (1.5 IQR)': UCL.values,  
        'LCL (1.5 IQR)': LCL.values  
    })  
    if draw_or_data=="data":   
        return result_df      
  
    # 创建Tkinter窗口和matplotlib图形      
    root = tk.Tk()      
    root.title("Trend Analysis")      
    # 得到屏幕宽度和高度，并设置窗口位置和大小（这里可以自定义）        
    sw = root.winfo_screenwidth()        
    sh = root.winfo_screenheight()        
    ww = 1300  # 窗口宽度        
    wh = 700  # 窗口高度        
    x = (sw - ww) // 2        
    y = (sh - wh) // 2        
    root.geometry(f"{ww}x{wh}+{x}+{y}")      
    fig, ax = plt.subplots(figsize=(10, 6))      
          
    # 解决汉字乱码问题      
    plt.rcParams["font.sans-serif"] = ["SimHei"]  # 使用指定的汉字字体类型（此处为黑体）      
    plt.rcParams['axes.unicode_minus'] = False  # 用来正常显示负号      
          
    # 绘制趋势图      
    ax.plot(result_df['Date'], result_df['Monthly Counts'], marker='o', label='Monthly Counts')      
    ax.plot(result_df['Date'], result_df['Rolling Mean'], color='red', label='Rolling Mean ({} Months)'.format(windows))      
    ax.plot(result_df['Date'], result_df['Rolling Median'], color='orange', linestyle='--', label='Rolling Median ({} Months)'.format(windows))  # Median line in orange      
    ax.plot(result_df['Date'], result_df['Rolling 25th Percentile'], color='green', linestyle='--', label='Rolling 25th Percentile ({} Months)'.format(windows))  # 25th percentile line in green      
    ax.plot(result_df['Date'], result_df['Rolling 75th Percentile'], color='blue', linestyle='--', label='Rolling 75th Percentile ({} Months)'.format(windows))  # 75th percentile line in blue       
    ax.plot(result_df['Date'], result_df['UCL (1.5 IQR)'], color='purple', linestyle='--', label='UCL (1.5 IQR)')  # UCL line in purple      
    ax.plot(result_df['Date'], result_df['LCL (1.5 IQR)'], color='purple', linestyle='--')  # LCL line in purple       
    ax.set_xlabel('Date')      
    ax.set_ylabel('Event Counts')      
    ax.set_title('Trend Analysis with Control Limits (1.5 IQR)')      
    ax.legend()      
    ax.grid(True)      
          
    # 使用FigureCanvasTkAgg将图形嵌入到Tkinter窗口中      
    canvas = FigureCanvasTkAgg(fig, master=root)      
    canvas.draw()      
    canvas.get_tk_widget().pack(side=tk.TOP, fill=tk.BOTH, expand=1)      
          
    # 创建和显示工具条    
    toolbar = NavigationToolbar2Tk(canvas, root)    
    toolbar.update()    
    canvas.get_tk_widget().pack()    
  
    # 运行Tkinter事件循环    
    root.mainloop() 




################################
#移植过来的模块

def TOOLS_data_masking(data):  # 2222222
    """-数据脱敏，meithon= 药品 或者 器械"""
    from random import choices
    from string import ascii_letters, digits

    data = data.reset_index(drop=True)
    print('适用于配置文件中的其他选项卡')
    methon = "其他"
    umeu = peizhidir+'share（easy_adrmdr）数据脱敏.xls'
    try:
        masking_set = pd.read_excel(
            umeu, sheet_name=methon, header=0, index_col=0
        ).reset_index()
    except:
        showinfo(title="错误信息", message="该功能需要配置文件才能使用！")
        return 0
    x = 0
    x2 = len(data)
    data["abcd"] = "□"
    for i in masking_set["要脱敏的列"]:
        x = x + 1
        print( "\n正在对以下列进行脱敏处理：")
        print(i)
        try:
            ids = set(data[i])
        except:
            showinfo(title="提示", message="脱敏文件配置错误，请修改配置表。")
            return 0
        id_mapping = {si: "".join(choices(digits, k=10)) for si in ids}
        data[i] = data[i].map(id_mapping)
        data[i] = data["abcd"] + data[i].astype(str)
    
    del data["abcd"]

    print("\n脱敏操作完成。")

    return data

def TOOLS_easyreadT(bos):  # 查看表格
    """行列互换查看表格"""
    bos[
        "#####分隔符#########"
    ] = "######################################################################"
    bos2 = bos.stack(dropna=False)
    bos2 = pd.DataFrame(bos2).reset_index()
    bos2.columns = ["序号", "条目", "详细描述T"]
    bos2["逐条查看"] = "逐条查看"
    bos2["报表类型"] = "逐条查看"
    return bos2


def TOOLS_get_new_GUI(df):  
    # 创建主窗口（这里使用Toplevel窗口以避免关闭整个应用程序）  
    root = tk.Toplevel()  
    root.title("新的不良反应检测（集合法）")  
    sw = root.winfo_screenwidth()  
    sh = root.winfo_screenheight()  
    ww = 430  # 窗口宽度  
    wh = 180  # 窗口高度（增加以适应标签）  
    x = (sw - ww) // 2  
    y = (sh - wh) // 2  
    root.geometry(f"{ww}x{wh}+{x}+{y}")  
  
    # 创建标签和下拉菜单  
    label_column = tk.Label(root, text="请选择对象：")  
    label_column.grid(row=0, column=0, padx=10, pady=10)  
      
    column_var = tk.StringVar(root,value="证号")  # 预设值为"证号"  
    column_options = ['证号', '品种', '页面']  # 使用预设的选项  
    column_menu = tk.OptionMenu(root, column_var, *column_options)  
    column_menu.grid(row=0, column=1, padx=10, pady=10)  
  
    # 创建标签和输入框  
    label_number = tk.Label(root, text="请输入灵敏度：")  
    label_number.grid(row=1, column=0, padx=10, pady=5)  
      
    number_var = tk.StringVar(value="1")  
    number_entry = tk.Entry(root, textvariable=number_var)  
    number_entry.grid(row=1, column=1, padx=10, pady=5)  
  
    # 创建确定按钮  
    def on_ok_click():  
        selected_column = column_var.get()  
        input_number = number_var.get()  
          
        # 简单的输入验证  
        try:  
            input_number = int(input_number)  
        except ValueError:  
            messagebox.showerror("输入错误", "请输入一个有效的数字。")  
            return  
          
        # 这里可以处理所选内容、输入的数字以及DataFrame  
        # 例如，打印结果或将其传递给其他函数  
        result = (selected_column, input_number, df)  
        TOOLS_get_new(df,selected_column,input_number) # 在实际应用中，应该替换为适当的处理逻辑  
          
        # 关闭窗口  
        root.destroy()  
  
    ok_button = tk.Button(root, text="确定", command=on_ok_click)  
    ok_button.grid(row=2, column=1, columnspan=2, pady=20)  
  
    # 启动Tkinter的事件循环  
    root.mainloop()  


def TOOLS_get_new(df1,methon,n):
    """监测新的不良反应-不需要说明书的版本 df1:原始数据，df2：规则文件，df4重点关注"""  
    
    def drug(data,n2=1):
        """药品不良反应名称统计"""	#仅用于页面方法
        data=data.drop_duplicates("报告表编码") 
		   
        rm=str(Counter(SMALL_get_list("use(不良反应名称（规整）).file",data,1000))).replace("Counter({", "{")
        rm=rm.replace("})", "}")
        import ast
        user_dict = ast.literal_eval(rm)    
        df = TOOLS_easyreadT(pd.DataFrame([user_dict]))
        df=df[(df['详细描述T']!='######################################################################')]
        df=df[(df['详细描述T'].astype(int)>=n2)]
        df = df.rename(columns={"逐条查看": "不良反应（规整）"})
        return df

    def filter_adverse_effects(df, column_name, n):  
        # 创建一个空列表来存储符合条件的行  
        result_rows = []  
          
        # 遍历 DataFrame 的每一行  
        for index, row in df.iterrows():  
            adverse_effects_str = row['不良反应']  
              
            # 使用正则表达式来匹配和提取症状及其对应的数字  
            matches = re.findall(r'([^（]+)（(\d+)）', adverse_effects_str)  
              
            # 过滤出数字大于等于 n 的症状  
            filtered_effects = [(symptom, int(count)) for symptom, count in matches if int(count) >= n]  
              
            # 如果还有符合条件的不良反应，则构建新的字符串  
            if filtered_effects:  
                new_adverse_effects_str = '、'.join(f"{symptom}（{count}）" for symptom, count in filtered_effects)  
                new_adverse_effects_str = new_adverse_effects_str.replace("、、", '、')  
                  
                # 创建一个新的 Series，包含当前行的所有数据和更新后的不良反应字符串  
                new_row = row.copy()  
                new_row['不良反应'] = new_adverse_effects_str  
                  
                # 将新的 Series 添加到结果列表中  
                result_rows.append(new_row)  
          
        # 将结果列表转换为 DataFrame  
        result_df = pd.DataFrame(result_rows)  
          
        return result_df  

    if methon=="证号":
        df2=TOOLS_create_pivot_tool(df1[(df1['报告类型-新的']!="新的")],[['-注册证备案证'], [], ['不良反应名称（规整）'], ['SMALL_count_mode'], '', ''])
        df2.rename(columns={"-注册证备案证": "批准文号","不良反应名称（规整）": "不良反应"},inplace=True)
        df2=SMALL_expand_dict_like_columns(df2)
        df2=filter_adverse_effects(df2,'批准文号',n)
        df2['不良反应']=df2['不良反应'].astype(str)
        TOOLS_analyze_products(df1,df2,biaozhun['药品重点关注库'],methon)
    if methon=="品种":
        df2=TOOLS_create_pivot_tool(df1[(df1['报告类型-新的']!="新的")],[['-产品名称'], [], ['不良反应名称（规整）'], ['SMALL_count_mode'], '', ''])
        df2.rename(columns={"-产品名称": "通用名称","不良反应名称（规整）": "不良反应"},inplace=True)
        df2=SMALL_expand_dict_like_columns(df2)
        df2=filter_adverse_effects(df2,'通用名称',n)
        df2['不良反应']=df2['不良反应'].astype(str)
        TOOLS_analyze_products(df1,df2,biaozhun['药品重点关注库'],methon)  
    if methon=="页面":
        data=df1
        new=""
        old=""
        data_new=data.loc[data["报告类型-新的"].str.contains("新", na=False)].copy()
        data_old=data.loc[~data["报告类型-新的"].str.contains("新", na=False)].copy()
        list_new=drug(data_new,n)
        list_old=drug(data_old)
        if 1==1:
            for idc,colc in list_old.iterrows():
                    if  "分隔符" not in  colc["条目"]:
                        kde="'"+str(colc["条目"])+"':"+str(colc["详细描述T"])+","                
                        old=old+kde                                              
            for idc,colc in list_new.iterrows():
                if str(colc["条目"]).strip() not in old and "分隔符" not in str(colc["条目"]):
                    kde="'"+str(colc["条目"])+"':"+str(colc["详细描述T"])+","
                    new=new+kde    
        old="{"+old+"}"
        new="{"+new+"}"
        allon="\n可能是新的不良反应：\n\n"+new+"\n\n\n可能不是新的不良反应：\n\n"+old
        PROGRAM_display_content_in_textbox(allon)
        
def TOOLS_analyze_products(df1, df2, df4,methon):  
    """监测新的不良反应-匹配库法（优化版） df1:原始数据，df2：规则文件，df4重点关注"""  
  
    # 定义一个函数来检查不良反应是否存在于df2中  ，注册证号模式用
    def is_new_adverse_reaction(row):  
        matched_rows = df2[df2['批准文号'] == row['-注册证备案证']]  
        if not matched_rows.empty and str(row['不良反应名称（规整）']) not in matched_rows['不良反应'].str.cat(sep='|'):  
            return row['不良反应名称（规整）']  
        return None  
        
    # 定义一个函数来检查不良反应是否存在于df2中 通用名称模式用 
    def is_new_adverse_reaction_pz(row):  
        matched_rows = df2[df2['通用名称'] == row['-产品名称']]  
        if not matched_rows.empty and str(row['不良反应名称（规整）']) not in matched_rows['不良反应'].str.cat(sep='|'):  
            return row['不良反应名称（规整）']  
        return None  
  
    # 定义一个函数来检查不良反应是否为重点关注  
    def is_key_adverse_reaction(row):  
        for key in df4['重点关注库']:  
            if key in str(row['不良反应名称（规整）']):  
                return row['不良反应名称（规整）']  
        return None  

    def app_3(d): 
        if any(value >= 3 for value in d.values()):  
            return 3  
        else:  
            return 0  # 如果没有，可以返回0或其他默认值 
    def app_2(d): 
        if any(value >=1 for value in d.values()):  
            return 2  
        else:  
            return 0  # 如果没有，可以返回0或其他默认值         

    # 应用函数到df1的每一行  
    if methon=='证号':
        df1['疑似新的不良反应'] = df1.apply(is_new_adverse_reaction, axis=1) 
    elif methon=='品种':  
        df1['疑似新的不良反应'] = df1.apply(is_new_adverse_reaction_pz, axis=1) 
    df1['重点关注的不良反应'] = df1.apply(is_key_adverse_reaction, axis=1)  

  
    # 调用 TOOLS_create_pivot_tool 函数（确保它已正确定义并接受适当的参数）  
    # 注意：这里假设 TOOLS_create_pivot_tool 的参数和返回值是已知的，并且与下面的代码兼容  
    if methon=='证号':
        df3=TOOLS_create_pivot_tool(df1,[['-注册人备案人', '产品类别', '-产品名称', '-注册证备案证'], ['-伤害'], ['报告编码'], ['nunique'], {'不良反应名称（规整）': 'SMALL_count_mode', '疑似新的不良反应': 'SMALL_count_mode', '重点关注的不良反应': 'SMALL_count_mode'}, ['报告编码合计']])
    elif methon=='品种':
        df3=TOOLS_create_pivot_tool(df1,[['产品类别', '-产品名称'], ['-伤害'], ['报告编码'], ['nunique'], {'不良反应名称（规整）': 'SMALL_count_mode', '疑似新的不良反应': 'SMALL_count_mode', '重点关注的不良反应': 'SMALL_count_mode'}, ['报告编码合计']])

  
    # 如果 TOOLS_create_pivot_tool 没有包含 '参考不良反应'，则在这里添加  
    if '参考不良反应' not in df3.columns:
        if methon=='证号':
            df3['参考不良反应'] = df3.apply(lambda row: df2[df2['批准文号'] == row['-注册证备案证']]['不良反应'].str.cat(sep='|') if any(df2['批准文号'] == row['-注册证备案证']) else '', axis=1)  
        elif methon=='品种':
            df3['参考不良反应'] = df3.apply(lambda row: df2[df2['通用名称'] == row['-产品名称']]['不良反应'].str.cat(sep='|') if any(df2['通用名称'] == row['-产品名称']) else '', axis=1)  
      
    # 应用评分函数  
    df3['疑似新的不良反应'] = df3['疑似新的不良反应'].apply(lambda x: x if pd.notna(x) and '测试XX' not in str(x) else None)  
    df3['重点关注的不良反应'] = df3['重点关注的不良反应'].apply(lambda x: x if pd.notna(x) and '测试XX' not in str(x) else None)  


    df3['新的评分'] = df3['疑似新的不良反应'].apply(app_3)  
    df3['重点关注评分'] = df3['重点关注的不良反应'].apply(app_2)  
    del  df3["报表类型"]
    
    # 计算总评分  
    df3['总评分'] = df3['新的评分'] + df3['重点关注评分']  
  
    # 排序  
    df3 = df3.sort_values(by=['总评分', '报告编码合计'], ascending=[False, False]).reset_index(drop=True)  
  
    # 假设 '报表类型' 是一个固定的字符串，不需要动态生成  
    if methon=='证号':
        df3["报表类型"] = "{'grouped':['-注册人备案人', '产品类别', '-产品名称', '-注册证备案证']}"  
    elif methon=='品种':
        df3["报表类型"] = "{'grouped':['产品类别', '-产品名称']}"  
          
    PROGRAM_display_df_in_treeview(df3,0,df1)
 



def TOOLS_keti(data):
	"""日期预警功能"""
	
	def STAT_countx(x):
		"""所有成分关键字计数,返回一个字典""" 
		return x.value_counts().to_dict()
		
	def STAT_countpx(x,y):
		"""特定成分关键字计数,返回一个数值""" 
		return len(x[(x==y)])#.values	

	def STAT_get_mean(df):
		"""返回平均值""" 
		return round(df.value_counts().mean(),2)
		
	def STAT_get_std(df):
		"""返回标准差""" 
		return round(df.value_counts().std(ddof=1),2)
		

	def df_findrisk(df,target):	
		"""预警模块,针对批号、月份、季度""" 
		if target=="产品批号":
			return STAT_find_risk(df[(df["产品类别"]!="有源")],["上市许可持有人名称","产品类别","产品名称","注册证编号/曾用注册证编号"],"注册证编号/曾用注册证编号",target)
		else:
			return STAT_find_risk(df,["上市许可持有人名称","产品类别","产品名称","注册证编号/曾用注册证编号"],"注册证编号/曾用注册证编号",target)
			
	def STAT_find_risk(df,cols_list,main_col,target):	
			"""评分及预警模块,cols_list为所要引入的列（列表形式），main_col统计对象列（关键字），target为月份、季度或者批号等""" 
			df=df.drop_duplicates(["报告编码"]).reset_index(drop=True)
			dfx_findrisk1=df.groupby(cols_list).agg(
				证号总数量=(main_col,"count"),	
				包含元素个数=(target,"nunique"),
				历史数据=(target,STAT_countx),		
				均值=(target,STAT_get_mean),
				标准差=(target,STAT_get_std),			
				).reset_index()
						
			cols_list2=cols_list.copy()
			cols_list2.append(target)
			dfx_findrisk2=df.groupby(cols_list2).agg(
				计数=(target,"count"),
				严重伤害数=("伤害",lambda x: STAT_countpx(x.values,"严重伤害")),
				死亡数量=("伤害",lambda x: STAT_countpx(x.values,"死亡")),	
				单位个数=("单位名称","nunique"),	
				单位列表=("单位名称",STAT_countx),							
				).reset_index()				

			dfx_findrisk=pd.merge(dfx_findrisk2,dfx_findrisk1,on=cols_list,how="left")#.reset_index()	
					
			return dfx_findrisk
		
	
	
	def keti(timex,time_windows,time_base,data0,kx):
		print(time_windows,time_base)
		
		if time_base=='发生日期':
			time_base='事件发生日期'
			time_base_m='事件发生月份'
			time_base_q='事件发生季度'
		if time_base=='报告日期':
			time_base='报告日期'
			time_base_m='报告月份'
			time_base_q='报告季度'			
											
			
		data0['规整后品类']=data0['产品名称']
		data0['报告日期'] = pd.to_datetime(data0['报告日期'], format='%Y-%m-%d', errors='coerce') 	
		data0['事件发生日期'] = pd.to_datetime(data0['事件发生日期'], format='%Y-%m-%d', errors='coerce') 	
		data0["伤害与评价"]=data0["伤害"]+data0["持有人报告状态"]				
		data0["报告月份"] = data0["报告日期"].dt.to_period("M").astype(str)	
		data0["报告季度"] = data0["报告日期"].dt.to_period("Q").astype(str)	
		data0["报告年份"] = data0["报告日期"].dt.to_period("Y").astype(str)	#	品种评价		
		data0["事件发生月份"] = data0["事件发生日期"].dt.to_period("M").astype(str)			
		data0["事件发生季度"] = data0["事件发生日期"].dt.to_period("Q").astype(str)				
		data0["事件发生年份"] = data0["事件发生日期"].dt.to_period("Y").astype(str)
		
	
			
						
		k1=kx["权重"][0] #严重比
		k2=kx["权重"][1] #单位数量		
		k3=kx["权重"][2] #批号或型号集中度权重
		k4=kx["权重"][3] #高度关注关键字（一级）		
		k4_values=kx["值"][3] #高度关注关键字（一级） 值	
		
		k5=kx["权重"][4] #高度关注关键字（二级）		
		k5_values=kx["值"][4] #高度关注关键字（二级） 值
		
		k6=kx["权重"][5] #高风险品种		
		k6_values=kx["值"][5] 

		k7=kx["权重"][6] #低风险品种	
		k7_values=kx["值"][6] 
		
		k8=kx["权重"][7] #低风险问题	
		k8_values=kx["值"][7] 
	
						
		lastdayfrom = pd.to_datetime(timex)
		data2=data0.copy().set_index(time_base)
		data2=data2.sort_index()
		data2['高度关注关键字']=''
		data2['二级敏感词']=''
		data2['高风险品种']=0
		data2['低风险品种']=0
		data2['低风险问题']=0
		if ini["模式"]=="器械":		
			data2["关键字查找列"]=data2["器械故障表现"].astype(str)+data2["伤害表现"].astype(str)+data2["使用过程"].astype(str)+data2["事件原因分析描述"].astype(str)+data2["初步处置情况"].astype(str)#器械故障表现|伤害表现|使用过程|事件原因分析描述|初步处置情况
		else:
			data2["关键字查找列"]=data2["器械故障表现"].astype(str)
		
		# 定义一个函数来检查匹配并更新列
		def update_keywords(row, patterns):
			matched_keywords = [pattern.pattern for pattern in patterns if pattern.search(row["关键字查找列"])]
			if matched_keywords:
				return ";".join(matched_keywords)
			return ""
			
		k4_set = k4_values.split('|')
		k5_set = k5_values.split('|')
				
		compiled_patterns_k4 = [re.compile(pattern) for pattern in k4_set]
		compiled_patterns_k5 = [re.compile(pattern) for pattern in k5_set]		
		data2["高度关注关键字"] = data2.apply(lambda row: update_keywords(row, compiled_patterns_k4), axis=1)
		data2["二级敏感词"] = data2.apply(lambda row: update_keywords(row, compiled_patterns_k5), axis=1)

		data2.loc[data2["关键字查找列"].str.contains(k8_values, na=False), "低风险问题"]  = 1	
		data2.loc[data2["产品名称"].str.contains(k6_values, na=False), "高风险品种"]  = 1	
		data2.loc[data2["关键字查找列"].str.contains(k7_values, na=False), "低风险问题"]  = 1	
		
		#月度数据
		if time_windows=='月份窗口':					
			data30 = data2.loc[lastdayfrom - pd.Timedelta(days=30):lastdayfrom].reset_index()

		#季度数据
		if time_windows=='季度窗口':	
			data30 = data2.loc[lastdayfrom - pd.Timedelta(days=90):lastdayfrom].reset_index()
			
		#增加对使用所有数据的兼容性
		if time_windows=='所有数据':
			data30=data2.copy()
		
		
		
		#对于月度窗口或全部数据窗口所对应的全部数据：
		if time_windows!='季度窗口':	
			timex_date = pd.to_datetime(timex)
			start_date = (timex_date.replace(day=1) - pd.DateOffset(months=12)).date()
			end_date = (timex_date.replace(day=1) - timedelta(days=1)).date()
			data365 = data2.loc[start_date:end_date]
			#data365=data2.loc[lastdayfrom - pd.Timedelta(days=365):lastdayfrom].reset_index()
		#季度窗口的情况
		else:
			# 将timex转换为Pandas的datetime对象
			timex_date = pd.to_datetime(timex)
			# 将timex_date转换为季度频率的Period对象
			current_quarter_period = pd.Period(timex_date, freq='Q')
			# 获取前一个季度的Period对象
			previous_quarter_period = current_quarter_period - 1
			# 将前一个季度的Period对象转换为时间戳，并指定为季度末的时间点
			# 注意：这里我们不需要再减去一天，因为to_timestamp('M', how='end')直接给我们季度的最后一天
			previous_quarter_end_timestamp = previous_quarter_period.to_timestamp('M', how='end')
			# 如果只需要日期部分，可以使用.date()方法
			end_date = previous_quarter_end_timestamp.date()
			start_date = end_date - timedelta(days=365)
			data365 = data2.loc[start_date:end_date]
			
		#当前时间窗数据评分
		df301=data30.groupby(["上市许可持有人名称","产品类别","规整后品类","产品名称","注册证编号/曾用注册证编号"]).agg(
			证号计数=("报告编码","nunique"),
			批号个数=("产品批号","nunique"),
			批号列表=("产品批号",STAT_countx),	
			型号个数=("型号","nunique"),
			型号列表=("型号",STAT_countx),		
			规格个数=("规格","nunique"),	
			规格列表=("规格",STAT_countx),		
			).sort_values(by="证号计数", ascending=[False], na_position="last").reset_index()	

		df302=data30.drop_duplicates(["报告编码"]).groupby(["上市许可持有人名称","产品类别","规整后品类","产品名称","注册证编号/曾用注册证编号"]).agg(
			严重伤害数=("伤害",lambda x: STAT_countpx(x.values,"严重伤害")),
			死亡数量=("伤害",lambda x: STAT_countpx(x.values,"死亡")),	
			单位个数=("单位名称","nunique"),	
			单位列表=("单位名称",STAT_countx),					
			待评价数=("持有人报告状态",lambda x: STAT_countpx(x.values,"待评价")),
			严重伤害待评价数=("伤害与评价",lambda x: STAT_countpx(x.values,"严重伤害待评价")),
			高度关注关键字=("高度关注关键字","sum"),	
			二级敏感词=("二级敏感词","sum"),
			高风险品种=("高风险品种","sum"),			
			低风险品种=("低风险品种","sum"),			
			低风险问题=("低风险问题","sum"),			
			).reset_index()	
		
		df30=pd.merge(df301,  df302,on=["上市许可持有人名称","产品类别","规整后品类","产品名称","注册证编号/曾用注册证编号"], how="left")	


		
	
		df30xinghao=data30.groupby(["上市许可持有人名称","产品类别","产品名称","注册证编号/曾用注册证编号","型号"]).agg(
			型号计数=("报告编码","nunique"),	
			).sort_values(by="型号计数", ascending=[False], na_position="last").reset_index()		
		df30xinghao=df30xinghao.drop_duplicates("注册证编号/曾用注册证编号")		
			
		df30wu=data30.drop_duplicates(["报告编码"]).groupby(["上市许可持有人名称","产品类别","产品名称","注册证编号/曾用注册证编号","产品批号"]).agg(
			批号计数=("报告编码","nunique"),	
			严重伤害数=("伤害",lambda x: STAT_countpx(x.values,"严重伤害")),
			).sort_values(by="批号计数", ascending=[False], na_position="last").reset_index()

					
			
		df30wu["风险评分-影响"]=0		
		df30wu["评分说明"]=""		
		df30wu.loc[((df30wu["批号计数"]>=3)&(df30wu["严重伤害数"]>=1)&(df30wu["产品类别"]!="有源"))|((df30wu["批号计数"]>=5)&(df30wu["产品类别"]!="有源")), "风险评分-影响"] = df30wu["风险评分-影响"]+3	
		df30wu.loc[(df30wu["风险评分-影响"]>=3), "评分说明"] = df30wu["评分说明"]+"●符合预警无源规则+3;"	
		

		
		df30wu=df30wu.sort_values(by="风险评分-影响", ascending=[False], na_position="last").reset_index(drop=True)	
		df30wu=df30wu.drop_duplicates("注册证编号/曾用注册证编号")
		df30xinghao=df30xinghao[["上市许可持有人名称","产品类别","产品名称","注册证编号/曾用注册证编号","型号","型号计数"]]	
		df30wu=df30wu[["上市许可持有人名称","产品类别","产品名称","注册证编号/曾用注册证编号","产品批号","批号计数","风险评分-影响","评分说明"]]
		df30=pd.merge(df30, df30xinghao, on=["上市许可持有人名称","产品类别","产品名称","注册证编号/曾用注册证编号"], how="left")
	
		df30=pd.merge(df30, df30wu, on=["上市许可持有人名称","产品类别","产品名称","注册证编号/曾用注册证编号"], how="left")		

		#TABLE_tree_Level_2(df30,1,data30)		
		#符合省中心规则打分（因为是针对证号，按有源标准）
		df30.loc[((df30["证号计数"]>=3)&(df30["严重伤害数"]>=1)&(df30["产品类别"]=="有源"))|((df30["证号计数"]>=5)&(df30["产品类别"]=="有源")), "风险评分-影响"] = df30["风险评分-影响"]+3	
		df30.loc[(df30["风险评分-影响"]>=3)&(df30["产品类别"]=="有源"), "评分说明"] = df30["评分说明"]+"●符合预警有源规则+3;"	


				
		#针对死亡
		df30.loc[(df30["死亡数量"]>=1), "风险评分-影响"] = df30["风险评分-影响"]+50	
		df30.loc[(df30["风险评分-影响"]>=10), "评分说明"] = df30["评分说明"]+"存在死亡报告;"	
		
		#严重比评分
		fen_yanzhong=round(k1*(df30["严重伤害数"]/df30["证号计数"]),2)
		df30["风险评分-影响"] = df30["风险评分-影响"]+	fen_yanzhong
		df30["评分说明"] = df30["评分说明"]+"严重比评分"+fen_yanzhong.astype(str)+";"			
		
		#报告单位数评分
		fen_danwei=round(k2*(np.log(df30["单位个数"])),2)
		df30["风险评分-影响"] = df30["风险评分-影响"]+	fen_danwei
		df30["评分说明"] = df30["评分说明"]+"报告单位评分"+fen_danwei.astype(str)+";"				
		
		#批号型号集中度评分
		df30.loc[(df30["产品类别"]=="有源")&(df30["证号计数"]>=3), "风险评分-影响"] = df30["风险评分-影响"]+k3*df30["型号计数"]/df30["证号计数"]			
		df30.loc[(df30["产品类别"]=="有源")&(df30["证号计数"]>=3), "评分说明"] = df30["评分说明"]+"型号集中度评分"+(round(k3*df30["型号计数"]/df30["证号计数"],2)).astype(str)+";"	
		df30.loc[(df30["产品类别"]!="有源")&(df30["证号计数"]>=3), "风险评分-影响"] = df30["风险评分-影响"]+k3*df30["批号计数"]/df30["证号计数"]			
		df30.loc[(df30["产品类别"]!="有源")&(df30["证号计数"]>=3), "评分说明"]  = df30["评分说明"]+"批号集中度评分"+(round(k3*df30["批号计数"]/df30["证号计数"],2)).astype(str)+";"			

		#高度关注关键字（一级）
		df30.loc[(df30["高度关注关键字"]!=''), "风险评分-影响"]  = df30["风险评分-影响"]+k4
		df30.loc[(df30["高度关注关键字"]!=''), "评分说明"] = df30["评分说明"]+"●含有高度关注关键字评分"+str(k4)+"；"									

		#二级敏感词
		df30.loc[(df30["二级敏感词"]!=''), "风险评分-影响"]  = df30["风险评分-影响"]+k5
		df30.loc[(df30["二级敏感词"]!=''), "评分说明"] = df30["评分说明"]+"含有二级敏感词评分"+str(k5)+"；"		
		
		#高风险品种
		df30.loc[(df30["高风险品种"]>=1), "风险评分-影响"]  = df30["风险评分-影响"]+k6
		df30.loc[(df30["高风险品种"]>=1), "评分说明"] = df30["评分说明"]+"高风险品种"+str(k6)+"；"	
		
		#低风险品种
		df30.loc[(df30["低风险品种"]>=1), "风险评分-影响"]  = df30["风险评分-影响"]+k7
		df30.loc[(df30["低风险品种"]>=1), "评分说明"] = df30["评分说明"]+"高风险品种"+str(k7)+"；"	
		
		
		#低风险问题
		df30.loc[(df30["低风险问题"]>=1), "风险评分-影响"]  = df30["风险评分-影响"]+k8
		df30.loc[(df30["低风险问题"]>=1), "评分说明"] = df30["评分说明"]+"减分项评分"+str(k8)+"；"	
		

		
				
		#历史比较（月份或季度）
		if time_windows=='月份窗口' or time_windows=='所有数据':	
			df365month=df_findrisk(data365,time_base_m)
		if time_windows=='季度窗口':	
			df365month=df_findrisk(data365,time_base_q)			
		df365month=df365month.drop_duplicates("注册证编号/曾用注册证编号")	
		df365month=df365month[["注册证编号/曾用注册证编号",'历史数据',"均值","标准差"]]
		df30=pd.merge(df30, df365month, on=["注册证编号/曾用注册证编号"], how="left")	
		df30["均值"]=round(df30["均值"],2)	
		df30["标准差"]=round(df30["标准差"],2)
		df30["风险评分-历史"]=1
		df30["mfc"]=""

		#增加对使用所有数据的兼容性
		if time_windows!='所有数据':
			
			df30.loc[(df30["证号计数"]>=3), "风险评分-历史"]  = df30["风险评分-历史"]+2
			df30.loc[(df30["证号计数"]>=3),  "mfc"] = "数量超过3例；"		
			
			
			df30.loc[((df30["证号计数"]>df30["均值"])&(df30["风险评分-历史"]>=3)&(df30["标准差"].astype(str)!="nan")), "风险评分-历史"]  = df30["风险评分-历史"]+1
			df30.loc[((df30["证号计数"]>df30["均值"])&(df30["风险评分-历史"]>=3)&(df30["标准差"].astype(str)!="nan")), "mfc"] = "月份计数超过历史均值"+df30["均值"].astype(str)+"；"	
				
			df30.loc[(df30["证号计数"]>=(df30["均值"]+df30["标准差"]))&(df30["风险评分-历史"]>=4), "风险评分-历史"]  = df30["风险评分-历史"]+2
			df30.loc[(df30["证号计数"]>=(df30["均值"]+df30["标准差"]))&(df30["风险评分-历史"]>=4), "mfc"] = "超过历史均值一个标准差；"			
			
				
			df30.loc[(df30["证号计数"]>=(df30["均值"]+2*df30["标准差"]))&(df30["风险评分-历史"]>=6), "风险评分-历史"]  = df30["风险评分-历史"]+2
			df30.loc[(df30["证号计数"]>=(df30["均值"]+2*df30["标准差"]))&(df30["风险评分-历史"]>=6), "mfc"] = "超过历史均值两个标准差；"				

		df30["评分说明"]=df30["评分说明"]+"●●证号数量："+df30["证号计数"].astype(str)+";"+ df30["mfc"]	
		del df30["mfc"]
		df30=df30.rename(columns={"均值": "历史均值","标准差": "历史标准差","历史数据": "历史时间数据"})
		

		#历史比较（批号）
		df365month=df_findrisk(data365,"产品批号")
		df365month=df365month.drop_duplicates("注册证编号/曾用注册证编号")	
		df365month=df365month[["注册证编号/曾用注册证编号","历史数据","均值","标准差"]]
		df30=pd.merge(df30, df365month, on=["注册证编号/曾用注册证编号"], how="left")	
	
		df30["风险评分-批号"]=1
		df30.loc[(df30["产品类别"]!="有源"), "评分说明"] =df30["评分说明"]+"●●高峰批号数量："+df30["批号计数"].astype(str)+";"
		
		#增加对使用所有数据的兼容性
		if time_windows!='所有数据':
			df30.loc[(df30["批号计数"]>=3), "风险评分-批号"]  = df30["风险评分-批号"]+2
			df30.loc[(df30["批号计数"]>=3),  "mfc"] = "数量超过3例；"		
			
			
			df30.loc[((df30["批号计数"]>df30["均值"])&(df30["风险评分-批号"]>=3)&(df30["标准差"].astype(str)!="nan")), "风险评分-批号"]  = df30["风险评分-批号"]+1
			df30.loc[((df30["批号计数"]>df30["均值"])&(df30["风险评分-批号"]>=3)&(df30["标准差"].astype(str)!="nan")), "mfc"] = "高峰批号计数超过历史均值"+df30["均值"].astype(str)+"；"	
				
			df30.loc[(df30["批号计数"]>=(df30["均值"]+df30["标准差"]))&(df30["风险评分-批号"]>=4), "风险评分-批号"]  = df30["风险评分-批号"]+2
			df30.loc[(df30["批号计数"]>=(df30["均值"]+df30["标准差"]))&(df30["风险评分-批号"]>=4), "mfc"] = "高峰批号超过历史均值一个标准差；"			
			
				
			df30.loc[(df30["批号计数"]>=(df30["均值"]+2*df30["标准差"]))&(df30["风险评分-批号"]>=6), "风险评分-批号"]  = df30["风险评分-批号"]+2
			df30.loc[(df30["批号计数"]>=(df30["均值"]+2*df30["标准差"]))&(df30["风险评分-批号"]>=6), "mfc"] = "高峰批号超过历史均值两个标准差；"		
		
			
		df30=df30.rename(columns={"均值": "历史批号均值","标准差": "历史批号标准差","历史数据": "历史批号数据"})

		
		df30["风险评分-影响"]=round(df30["风险评分-影响"],2)
		df30["风险评分-历史"]=round(df30["风险评分-历史"],2)
		df30["风险评分-批号"]=round(df30["风险评分-批号"],2)
		
		df30["总体评分"]=df30["风险评分-影响"].copy()
		df30["关注建议"]=""
		df30.loc[(df30["风险评分-影响"]>=3),  "关注建议"]=df30["关注建议"]+"●建议关注(影响范围)；" 
		df30.loc[(df30["风险评分-历史"]>=5),  "关注建议"]=df30["关注建议"]+"●建议关注(当期数量异常)；"
		df30.loc[(df30["风险评分-批号"]>=3),  "关注建议"]=df30["关注建议"]+"●建议关注(高峰批号数量异常)。"		
		df30.loc[(df30["风险评分-历史"]>=df30["风险评分-批号"]),  "总体评分"]=df30["风险评分-影响"]*df30["风险评分-历史"]
		df30.loc[(df30["风险评分-历史"]<df30["风险评分-批号"]),  "总体评分"]=df30["风险评分-影响"]*df30["风险评分-批号"]

		df30["总体评分"]=round(df30["总体评分"],2)		
		df30["评分说明"]=df30["关注建议"]	+df30["评分说明"]		
		df30=df30.sort_values(by=["总体评分","风险评分-影响"], ascending=[False,False], na_position="last").reset_index(drop=True)
		
			
			
		df30=df30[["上市许可持有人名称","产品类别","规整后品类","产品名称","注册证编号/曾用注册证编号","总体评分","风险评分-影响","风险评分-历史","风险评分-批号","评分说明","证号计数","严重伤害数","死亡数量","单位个数","单位列表","批号个数","批号列表","型号个数","型号列表","规格个数","规格列表","待评价数","严重伤害待评价数","高度关注关键字","二级敏感词",'历史时间数据',"历史均值","历史标准差",'历史批号数据',"历史批号均值","历史批号标准差","型号","型号计数","产品批号","批号计数"]]
		df30["报表类型"]="{'grouped':['上市许可持有人名称','产品类别','产品名称','注册证编号/曾用注册证编号']}"

		PROGRAM_display_df_in_treeview(df30,1,data30)
		pass			
	

	if 1==1:
		ini={}
		ini['模式']='器械'
		if ini["模式"]=="药品":
			kx = pd.read_excel(peizhidir+"easy_预警参数.xlsx", header=0, sheet_name="药品").reset_index(drop=True)
		if ini["模式"]=="器械":
			kx = pd.read_excel(peizhidir+"easy_预警参数.xlsx", header=0, sheet_name="器械").reset_index(drop=True)	
		if ini["模式"]=="化妆品":
			kx = pd.read_excel(peizhidir+"（范例）预警参数.xlsx", header=0, sheet_name="化妆品").reset_index(drop=True)

	se = tk.Toplevel()
	se.title('风险预警')
	sw_se = se.winfo_screenwidth()
    #得到屏幕宽度
	sh_se = se.winfo_screenheight()
    #得到屏幕高度
	ww_se = 250
	wh_se = 250
    #窗口宽高为100
	x_se = (sw_se-ww_se) / 2
	y_se = (sh_se-wh_se) / 2
	se.geometry("%dx%d+%d+%d" %(ww_se,wh_se,x_se,y_se)) 

	import_se=tk.Label(se,text="预警日期：")
	import_se.grid(row=1, column=0, sticky="w", padx=10, pady=10)
	import_se_entry=tk.Entry(se, width = 30)
	import datetime
	import_se_entry.insert(0,datetime.date.today())
	import_se_entry.grid(row=1, column=1, sticky="w")

	# 创建第一个下拉框：请选择时间窗
	time_window_var = tk.StringVar(value="月份窗口")  # 设置默认选项
	time_window_menu_options = ["月份窗口",'季度窗口','所有数据']  # 定义选项列表
	time_window_label = tk.Label(se, text="请选择时间窗：")
	time_window_label.grid(row=2, column=0, sticky="w", padx=10, pady=10)
	time_window_optionmenu = tk.OptionMenu(se, time_window_var, *time_window_menu_options)
	time_window_optionmenu.grid(row=2, column=1, sticky="w", padx=10, pady=10)
	 
	# 创建第二个下拉框：请选择时间依据
	time_basis_var = tk.StringVar(value="报告日期")  # 设置默认选项
	time_basis_menu_options = ["报告日期", "发生日期"]  # 定义选项列表
	time_basis_label = tk.Label(se, text="请选择时间依据：")
	time_basis_label.grid(row=3, column=0, sticky="w", padx=10, pady=10)
	time_basis_optionmenu = tk.OptionMenu(se, time_basis_var, *time_basis_menu_options)
	time_basis_optionmenu.grid(row=3, column=1, sticky="w", padx=10, pady=10)

	# 创建第3个下拉框：请选择规则
	workbook = load_workbook(peizhidir+"easy_预警参数.xlsx", read_only=True)
	sheet_names = workbook.sheetnames
	kx_basis_var = tk.StringVar(value="器械")  # 设置默认选项
	kx_basis_menu_options = sheet_names  # 定义选项列表
	kx_basis_label = tk.Label(se, text="请选择预警参数：")
	kx_basis_label.grid(row=4, column=0, sticky="w", padx=10, pady=10)
	kx_basis_optionmenu = tk.OptionMenu(se, kx_basis_var, *kx_basis_menu_options)
	kx_basis_optionmenu.grid(row=4, column=1, sticky="w", padx=10, pady=10)


	
	btn_se=tk.Button(se,text="确定",width=10,command=lambda:PROGRAM_display_df_in_treeview(keti(import_se_entry.get(),time_window_var.get(),time_basis_var.get(),data,pd.read_excel(peizhidir+"easy_预警参数.xlsx", header=0, sheet_name=kx_basis_var.get()).reset_index(drop=True)),1,data))
	btn_se.grid(row=5, column=1, sticky="w", padx=10, pady=10)

	
	pass
	














#######################################
#绘图函数
def AAAA_DRAW():
    pass
    
def DRAW_plot_df(ori_df):
    df=ori_df.copy() 
    df=df.fillna(0)
    print("绘图工具使用您的一份副本工作。请确保您的数据是经过汇总整理、且相关绘图要素无重复项。")
    cycle_iterator = itertools.cycle([1, 2, 3]) 

    def create_dropdown(frame, label_text, variable, options):  
        """Create a dropdown menu."""  
        label = ttk.Label(frame, text=label_text)  
        label.pack(side=tk.LEFT, pady=5)  
        dropdown = ttk.Combobox(frame, textvariable=variable, values=options, state='readonly',width= 15)  
        dropdown.pack(side=tk.LEFT,padx=3)  
        return dropdown  

    def clear_plot():  
        """Clears the current plot and resets axis properties."""  
        ax.clear()  # 清除轴上的所有内容  
        # 重置轴的限制和视图（可能需要根据具体情况调整）  
        ax.relim()  
        ax.autoscale_view()  
        # 可能还需要其他设置来确保下一个图能够正确显示  clear_legend
        canvas.draw()  # 重新绘制画布以更新显示 

    def clear_legend():  
        """Clears the current plot and resets axis properties."""  
         
        iss=next(cycle_iterator)
        if iss==1:
            ax.legend_ = None
        elif iss==2:
            ax.legend(loc='upper right', bbox_to_anchor=(1.11, 1.0), fontsize=8, borderaxespad=0.0) 
        elif iss==3:  
            ax.legend()   
                    
        # 清除轴上的所有内容  
        # 重置轴的限制和视图（可能需要根据具体情况调整）  
        ax.relim()  
        ax.autoscale_view()  
        # 可能还需要其他设置来确保下一个图能够正确显示  clear_legend
        canvas.draw()  # 重新绘制画布以更新显示 
        
    def draw_text(ax,texts):
        # 遍历每个条形并添加文本  
        for xxbar in ax.patches: 
            print(xxbar) 
            # 获取条形的x位置、宽度和高度  
            x = xxbar.get_x()  
            width = xxbar.get_width()  
            height = xxbar.get_height()  
            # 计算文本的位置（在条形的顶部中心）  
            text_x = x + width / 2  
            text_y = height     
            # 获取当前条形的索引（用于从texts中获取相应的文本）  
            index = int(xxbar.get_x()+1)    
            # 在条形顶部添加文本  
            ax.text(text_x, text_y, texts[index], ha='center', va='bottom', fontsize=8)     
        # 调整y轴的上限，以确保文本完全可见  
        ax.relim()  
        ax.autoscale_view()
  
 

  
    def draw_plot(x_label,x_value, y_label, y_value, z_column, m_column, plot_type, my_colors): 
        """Draw the plot based on user's selection."""  
        dfs=df.copy()
        #print(x_label,x_value, y_label, y_value, z_column, m_column, plot_type, my_colors)
        if x_value=="-自动生成-":
            dfs[x_value] = pd.factorize(dfs[x_label])[0]+1 
            dfs[x_value] =dfs[x_value].astype(int)
        if y_label=="-不显示-":
            dfs[y_label] = None
        next_color = next(my_colors)  # 获取下一个颜色
        if plot_type == "---------":
            return

 
        #################################################          

        if plot_type == "帕累托图(XY)":  
            clear_plot()  # 清除之前的图形  
            ax2 = ax.twinx()  # 创建第二个坐标轴对象用于累计百分比  
      
            # 确保使用传入的dfs而不是未定义的df  
            dfs = dfs.sort_values(by=y_value, ascending=False)  # 按y_value列的值降序排序  

            dfs['Cumulative Percentage'] = dfs[y_value].cumsum() / dfs[y_value].sum() * 100  # 计算累计百分比  
      
            # 绘制条形图  
            dfs.plot(kind='bar', x=x_label, y=y_value, ax=ax, color=next_color, legend=False)  
            next_color = next(my_colors)  # 获取下一个颜色  
      
            # 绘制累计百分比曲线  
            dfs.plot(kind='line',  x=x_label, y='Cumulative Percentage', ax=ax2, color=next_color, marker='D', ms=4)  

            draw_text(ax,dfs[y_label])

            # 设置坐标轴格式和标签  
            ax.set_xlabel(x_label)  # 设置X轴标签  
   
            ax.set_ylabel('count')  # 设置左侧Y轴标签  
            ax2.set_ylabel('Cumulative Percentage (%)')  # 设置右侧Y轴标签  
            ax2.yaxis.set_ticks_position('right')  # 将累计百分比的坐标轴放在右边  
            ax.set_xticklabels(dfs[x_label], rotation=90)  # 设置X轴标签  

            ax2.set_ylim(0, 100)  # 设置累计百分比的y轴范围  
            ax.set_title('Pareto Chart')  # 设置图表标题  

        elif plot_type == "热力图":  
            clear_plot()
            dfs=SMALL_pre_process_dataframe(dfs.copy())
            # 计算相关性矩阵  
            corr = dfs.corr()     
            # 生成一个mask用于隐藏上半部分的热力图（因为相关性矩阵是对称的）  
            mask = np.triu(np.ones_like(corr, dtype=bool))
            sns.heatmap(corr, mask=mask, cmap='coolwarm', annot=True, fmt=".2f")  
        

              
        ######################################            
        elif plot_type == "散点图(XY-Z分层-M大小)":  #分层散点图(XYZ)
            clear_plot()
            # 设置点的大小
            if m_column !="-不选择-":
                min_size = dfs[m_column].min()  
                max_size = dfs[m_column].max()  
                if min_size < 0 or max_size > 1:  
                    dfs['normalized_size'] = (dfs[m_column] - min_size) / (max_size - min_size)*100  
                    size_column = 'normalized_size'
            else:
                dfs['normalized_size']=100
                size_column = 'normalized_size'
            
            #如果不分层
            if z_column =="-不选择-":
                dfs.plots=dfs.plot(kind='scatter', x=x_value, y=y_value, ax=ax,s=size_column, label=x_value) 
                jittered_x=x_value
                jittered_y=y_value

            #如果分层
            else:
                clear_plot()

  
                # 添加抖动和透明度参数  
                jitter_amount = 0.2  # 抖动量，根据需要调整  
                alpha_value = 0.7  # 透明度值，根据需要调整 
                jittered_x="X("+str(x_value)+")"
                jittered_y="X("+str(y_value)+")"
                dfs[jittered_x]=0
                dfs[jittered_y]=0
                dfs[jittered_x] = dfs[x_value] + np.random.uniform(-jitter_amount, jitter_amount, size=len(dfs))  
                dfs[jittered_y] = dfs[y_value] + np.random.uniform(-jitter_amount, jitter_amount, size=len(dfs))
                
                
                for i in dfs[z_column].unique():  
                    product_data = dfs[dfs[z_column] == i]  
  
                    product_data.plot(kind='scatter', x=jittered_x, y=jittered_y, ax=ax, c=next_color, alpha=0.7, label=i,s=size_column) 
                    next_color = next(my_colors) 


            labels = dfs[y_label]  
            texts = [ax.text(x00, y00, z00, color='black', size=8) for x00, y00, z00 in zip(dfs[jittered_x], dfs[jittered_y], labels)] 
        
            ax.set_xticks(dfs[x_value])  # 设置X轴刻度为index列的值  
            if x_value!="-自动生成-":
                ax.set_xticklabels(dfs[x_value])  # 设置X轴标签为x_label_column列的值  
            else:
                ax.set_xticklabels(dfs[x_label])
            ax.legend()


        elif plot_type == "折线图(XY-Z分层)": 
            clear_plot()   
            if z_column =="-不选择-":  
                dfs.plot(kind='line', x=x_value, y=y_value, ax=ax, color=next_color, label=x_label, marker='D', ms=3)  
            else:
                clear_plot() 
                for i in df[z_column].unique():  
                    product_data = dfs[dfs[z_column] == i].sort_values(by=x_value)  
                    product_data.plot(kind='line', x=x_value, y=y_value, ax=ax, color=next_color, alpha=0.7, label=i, marker='D', ms=3)
                    next_color = next(my_colors) 
            ax.set_xticks(dfs[x_value])  # 设置X轴刻度为index列的值  
            if x_value!="-自动生成-":
                ax.set_xticklabels(dfs[x_value])  # 设置X轴标签为x_label_column列的值  
            else:
                ax.set_xticklabels(dfs[x_label])
            ax.legend()

            dfs.plot(kind='scatter', x=x_value, y=y_value, ax=ax, s=0.01) 
            labels = dfs[y_label]  
            texts = [ax.text(x00, y00, z00, color='black', size=8) for x00, y00, z00 in zip(dfs[x_value], dfs[y_value], labels)]   

                      
            
        elif plot_type == "左右比对条形图(XYM)":         
            clear_plot()
            # 确保数据是numpy数组以便于计算  
            data1 = dfs[y_value].values  
            data2 = dfs[m_column].values  
            labels= dfs[x_value].values  
              
            # 计算每侧数据的总和，用于缩放条形图的高度以保持总面积不变  
            total1 = data1.sum()  
            total2 = data2.sum()  
            max_total = max(total1, total2)  # 用于确定条形的最大高度  
            bar_width = 0.4  # 条形宽度  
            gap = 0.05  # 间隔宽度  
            n = len(labels)  # 标签的数量  
              
            # 计算条形的y轴位置  
            y_pos = np.arange(n) - gap / 2  # 条形图左侧的位置  
            y_pos_right = y_pos + bar_width + gap  # 
              
            # 绘制左侧的条形图（负方向）    
            ax.barh(y_pos, -data1 / max_total * 100, height=bar_width, align='center', edgecolor='white', label=y_value)    
            # 绘制右侧的条形图（正方向），注意这里应该使用 m_column 作为标签    
            ax.barh(y_pos, data2 / max_total * 100, height=bar_width, align='center', edgecolor='white', label=m_column)  
              
            # 设置y轴的刻度位置和标签，以及网格线等属性  
            ax.set_yticks(y_pos)  # 将y轴刻度设置在条形的中间位置  
            ax.set_yticklabels(labels)  # 设置y轴的刻度标签为传入的标签列表  
            ax.invert_yaxis()  # 反转y轴，使得标签从底部开始向上排列  
            ax.axvline(x=0, linewidth=1)  # 添加中心线  
            ax.set_title("横向对比图")  # 图表标题 
            
            
            figure.tight_layout()  # 自动调整子图参数，使之填充整个图像区域（可选） 
            # 添加图例，并设置其位置和属性  
            ax.legend()    
        

                    
        elif plot_type == "添加网格":         
            ax.grid(True) 
                                    
        
        elif plot_type == "柱状图(XY-Z分层)":  
            clear_plot() 
            if x_value!="-自动生成-":
                messagebox.showinfo("提示", "X数值仅支持-自动生成-")
                return
                
            if z_column=="-不选择-": 
                dfs.plot(kind='bar', x=x_label, y=y_value, ax=ax, color=next_color)  
            else:
                pivot_dfs = dfs.pivot(index=z_column, columns=x_label, values=y_value)  
                pivot_dfs.plot(kind='bar', width=0.8, ax=ax)   
                   

            draw_text(ax,dfs[y_label])
            
        elif plot_type == "堆叠柱状图(XYM-小到大)":  
            # 绘制堆叠柱状图 
            clear_plot() 
            dfs.plot(kind='bar', x=x_label, y=m_column, ax=ax, color=next_color) 
            #draw_text(ax,dfs[y_label])
            next_color = next(my_colors) 
            dfs.plot(kind='bar', x=x_label, y=y_value, ax=ax, color=next_color) 


            
        elif plot_type == "横向条形图(XY-Z分层)":  
            clear_plot() 
            if x_value!="-自动生成-":
                messagebox.showinfo("提示", "X数值仅支持-自动生成-")
                return
                
            if z_column=="-不选择-": 
                dfs.plot(kind='barh', x=x_label, y=y_value, ax=ax, color=next_color)  
                
                
            else:
                pivot_dfs = dfs.pivot(index=z_column, columns=x_label, values=y_value)  
                pivot_dfs.plot(kind='barh', width=0.8, ax=ax)           
                  
                # 为每个条形添加数值标签  
            if y_label!="-不显示-":
                for p in ax.patches:  
                    ax.annotate(str(p.get_width()), (p.get_width(), p.get_y() + p.get_height() / 2), ha='left', va='center')  



      
            
            
        elif plot_type == "横向堆叠条形图(XYM-小到大)":  
            # 绘制横向条形图 
            clear_plot() 
            dfs.plot(kind='barh', x=x_label, y=m_column, ax=ax, color=next_color)    
            next_color = next(my_colors)
            dfs.plot(kind='barh', x=x_label, y=y_value, ax=ax, color=next_color) 
            if y_label!="-不显示-":
                for p in ax.patches:  
                    ax.annotate(str(p.get_width()), (p.get_width(), p.get_y() + p.get_height() / 2), ha='left', va='center')              



            
        ################################################# 
        elif plot_type == "饼图(XY)":  
            clear_plot() 
            ax.pie(dfs[y_value], labels=dfs[x_label], autopct='%1.1f%%', startangle=90)  
            ax.axis('equal')  
        #################################################    
                        
        else:  
            print("不支持的图表类型。请检查您的输入。")  
        
		#格式设置
       
        labels = df[x_label].tolist() 
        fontsize = 8  # 初始字体大小  
        rotation = 90  # 初始旋转角度  
        plt.xticks(fontsize=fontsize)  # 更新字体大小  
        plt.xticks(rotation=rotation)  # 设置标签旋转角度  
      
          
        ax.yaxis.tick_left()  # Ensure y-axis is on the left (reset if needed)   
        ax.relim()  
        ax.autoscale_view()  
        if ax.legend_ and plot_type!="左右比对条形图(XYM)":  # 只在有图例时才显示图例  
            ax.legend(loc='upper right', bbox_to_anchor=(1.11, 1.0), fontsize=8, borderaxespad=0.0) 
        canvas.draw()  
  
    # Create the main window  
    root = tk.Tk()  
    root.title("数据可视化工具")  
    sw = root.winfo_screenwidth()          
    sh = root.winfo_screenheight()          
    ww = 1400  # 窗口宽度          
    wh = 600  # 窗口高度          
    x = (sw - ww) // 2          
    y = (sh - wh) // 2          
    root.geometry(f"{ww}x{wh}+{x}+{y}")   
    main_frame = ttk.Frame(root, padding="0")  
    main_frame.pack(fill=tk.BOTH, expand=True)  
    
    # Set up matplotlib with Chinese font support  
    plt.rcParams["font.sans-serif"] = ["SimHei"]  # Use specified Chinese font type (here: SimHei)  
    plt.rcParams['axes.unicode_minus'] = False  # Display negative signs properly  
    my_colors = itertools.cycle(plt.rcParams['axes.prop_cycle'].by_key()['color'])  # 创建一个颜色循环器 
    # Create the figure and axes (only once)  
    figure, ax = plt.subplots(figsize=(10, 4))  # Adjust figure size as needed  
    canvas = FigureCanvasTkAgg(figure, master=main_frame)  
    canvas.get_tk_widget().pack(fill=tk.BOTH, expand=True, pady=10)  
    toolbar = NavigationToolbar2Tk(canvas, main_frame)  
    toolbar.update()  
    canvas._tkcanvas.pack(fill=tk.BOTH, expand=True)  # Re-pack the canvas to include the toolbar space  
    toolbar.pack(fill=tk.X)  # Place the toolbar at the top or appropriate location  
  
    cols = ["-不选择-"]+df.columns.to_list()  # Get a list of column names  
  
    # Create labels and drop-down menus (using the create_dropdown function)  
    x_label_vars = tk.StringVar(root)  
    #x_label_vars.set("单位名称")  # Default value for "X-axis column" (should be a valid column name from cols)  
    create_dropdown(main_frame, "X标签:", x_label_vars,cols)  # Use column names list as selectable values  
    
    x_value_vars = tk.StringVar(root)  
    x_value_vars.set("-自动生成-")  # Default value for "X-axis column" (should be a valid column name from cols)  
    create_dropdown(main_frame, "X数值:", x_value_vars, ["-自动生成-"]+df.columns.to_list())  # Use column names list as selectable values  

    y_label_vars = tk.StringVar(root)  
    y_label_vars.set("-不显示-")  # Default value for "X-axis column" (should be a valid column name from cols)  
    create_dropdown(main_frame, "Y标签:", y_label_vars, ["-不显示-"]+df.columns.to_list())  # Use column names list as selectable values  

    y_value_vars = tk.StringVar(root)  
    #y_value_vars.set("报告编码其他")  # Default value for "Y-axis column" (should be a valid column name from cols)  
    create_dropdown(main_frame, "Y数值:", y_value_vars, cols)  # Use column names list as selectable values  
     
    z_column_var = tk.StringVar(root)    
    z_column_var.set("-不选择-")  # 假设数据框中有一个名为"风险级别"的列  
    create_dropdown(main_frame, "Z分层:", z_column_var, cols)     
    
    m_column_var = tk.StringVar(root)  
    m_column_var.set("-不选择-")  # 
    create_dropdown(main_frame, "M补充:", m_column_var, cols)  
  

    plot_type_var = tk.StringVar(root)  
    plot_type_var.set("散点图(XY-Z分层-M大小)")  # Default plot type is scatter plot  
    plot_type_dropdown = create_dropdown(main_frame, "绘图类型:", plot_type_var, ["散点图(XY-Z分层-M大小)","---------", "柱状图(XY-Z分层)","堆叠柱状图(XYM-小到大)","横向条形图(XY-Z分层)","左右比对条形图(XYM)","横向堆叠条形图(XYM-小到大)","---------","折线图(XY-Z分层)","帕累托图(XY)","---------","饼图(XY)","---------", "热力图","添加网格"])  # Allow selecting plot type from a list of options  

    
    # 创建一个清除按钮   
    #clear_button = ttk.Button(main_frame, text="清除", command=clear_plot)  
    #clear_button.pack(side=tk.LEFT, pady=5) 

    # 创建一个清除按钮   
    Mclear_button = ttk.Button(main_frame, text="图例", command=clear_legend)  
    Mclear_button.pack(side=tk.LEFT, pady=5)  

    # Create a button to trigger plotting and lambda function to call draw_plot with selected options and pass Y-axis position value correctly  
    plot_button = ttk.Button(main_frame, text="绘图", command=lambda: draw_plot(x_label_vars.get(), x_value_vars.get(),y_label_vars.get(),y_value_vars.get(), z_column_var.get(), m_column_var.get(), plot_type_var.get(),my_colors))  
    plot_button.pack(pady=5) 
     

    root.lift()
    root.attributes("-topmost", True)
    root.attributes("-topmost", False)
      
    # 定义一个trace方法，当z_column_var的值发生变化时更新Label的内容  
    def update_risk_label(*args):  
        text=z_column_var.get()
        print(text)  
      
    # 设置trace方法，使得当z_column_var的值发生变化时调用update_risk_label函数  
    z_column_var.trace("w", update_risk_label) 
    root.mainloop()  # Start the main event loop to run the GUI application

######################################
#######################################
#主界面函数
def AAAA_PROGRAM():
    pass
######################################


def PROGRAM_update_progress_bar(win,win_progressbar,now_schedule, all_schedule):  
    """更新进度条"""  
    win_progressbar['value']  = min(now_schedule / all_schedule * 100, 100)  # 限制在0-100%范围内  
    win.update()  # 更新界面  
    
def PROGRAM_thread_it(func, *args):
    """将函数打包进线程"""
    # 创建
    t = threading.Thread(target=func, args=args)
    # 守护 !!!
    t.setDaemon(True)
    # 启动
    t.start()    
 




def PROGRAM_display_content_in_textbox(content):    
    """  
    在弹出窗口中显示给定内容。  
      
    参数:  
        content (any): 要显示的内容，可以是字典、整数、浮点数或字符串。  
    """  
    def create_popup_menu(event):  
        """  
        为文本框创建右键菜单，并添加复制功能。  
          
        参数:  
            event (tkinter.Event): 触发此函数的事件对象。  
        """  
        popup_menu = tk.Menu(text_widget, tearoff=0)  # 创建右键菜单  
        popup_menu.add_command(label="复制", command=lambda: text_widget.event_generate("<<Copy>>"))  # 添加复制功能  
        popup_menu.tk_popup(event.x_root, event.y_root)  # 在指定位置显示菜单  
      
    root = tk.Tk()  # 创建主窗口  
    root.withdraw()  # 隐藏主窗口  
    
    popup = tk.Toplevel(root)  # 创建弹出窗口  
    popup.title("查看内容")  # 设置弹出窗口的标题  

    # 得到屏幕宽度
    sw_treeQ = popup.winfo_screenwidth()
    # 得到屏幕高度
    sh_treeQ = popup.winfo_screenheight()
    ww_treeQ = 800
    wh_treeQ = 600
    # 窗口宽高为100
    x_treeQ = (sw_treeQ - ww_treeQ) / 2
    y_treeQ = (sh_treeQ - wh_treeQ) / 2
    popup.geometry("%dx%d+%d+%d" % (ww_treeQ, wh_treeQ, x_treeQ, y_treeQ))

    text_widget = ScrolledText(popup, height=400, width=400, bg="#FFFFFF", wrap=tk.WORD)

    text_widget.pack(fill=tk.BOTH, expand=True, padx=10, pady=10)  # 设置文本框的布局和填充方式  
      
    # 根据内容类型显示内容  
    if isinstance(content, dict):  # 如果内容是字典类型  
        for key, value in content.items():  # 遍历字典并显示其内容  
            text_widget.insert(tk.END, f"{key}：{value}\n")  # 在文本框末尾插入内容  
    elif isinstance(content, (int, float, str)):  # 如果内容是整数、浮点数或字符串类型  
        text_widget.insert(tk.END, str(content))  # 将内容转换为字符串并显示在文本框末尾  
    else:  # 如果内容是其他类型  
        messagebox.showerror("错误", "不支持的内容类型")  # 显示错误消息框  
        popup.destroy()  # 销毁弹出窗口  
        return    # 结束函数执行  
          
    text_widget.config(state=tk.DISABLED)  # 设置文本框为不可编辑状态  
    text_widget.bind("<Button-3>", create_popup_menu)  # 绑定鼠标右键单击事件到创建右键菜单的函数  
    popup.lift()
    popup.attributes("-topmost", True)
    popup.attributes("-topmost", False)
    popup.mainloop()  # 显示弹出窗口，并进入消息循环，等待用户操作



              
def PROGRAM_display_df_in_treeview(ori_owercount_easyread, methon, ori, page_size=100):    
    """    
    在Treeview中分页展示DataFrame的内容，并添加点击列标题进行排序的功能。       
    参数:    
        ori_owercount_easyread: 工作文件，一个Pandas DataFrame。    
        ori:源文件。
        page_size (int): 每页显示的行数，默认为10000。    
    """  
    #兼容PSUR
    if methon=="psur":
        ori_owercount_easyread=PSRU_ori_owercount_easyread(ori_owercount_easyread)
    
            
    if isinstance(ori, pd.DataFrame):  
        pass  
    else:  
        ori=ori_owercount_easyread
     
    
      
    def sort_column(col, tree):   
        nonlocal current_page  # 声明current_page为非局部变量 
        nonlocal current_sort_order
        nonlocal ori_owercount_easyread           
        current_page = 0  
        ori_owercount_easyread = ori_owercount_easyread.sort_values(by=col, ascending=current_sort_order)    
        update_treeview(tree, ori_owercount_easyread)    
        # 切换排序顺序  
        current_sort_order = not current_sort_order   
    def update_treeview(tree, df):    
        tree.delete(*tree.get_children())  # 清除Treeview中的所有数据    
        start = current_page * page_size    
        end = start + page_size    
        for index, row in df[start:end].iterrows():    
            tree.insert("", "end", text=str(index), values=list(row))    
      
    def go_to_page(BF,page):    
        nonlocal current_page  # 声明current_page为非局部变量 
        
        total_pages = len(ori_owercount_easyread) // page_size + (len(ori_owercount_easyread) % page_size > 0) 
        current_page = page  
        update_treeview(tree, ori_owercount_easyread)    
        update_textbox(entry_text, str(len(ori_owercount_easyread))+"，"+str(current_page+1)+"/"+str(total_pages))  
        print("当前页：",current_page+1,"/",total_pages)

    def copy_to_clipboard():    
        try:    
            selected_item = tree.selection()[0]  # 获取选定的行    
            selected_values = tree.item(selected_item, "values")  # 获取选定的行的值    
            selected_dict = dict(zip(ori_owercount_easyread.columns, selected_values))  # 转换为字典格式    
            root.clipboard_clear()  # 清除Tk的剪贴板内容    
            root.clipboard_append(str(selected_dict))  # 将字典字符串复制到Tk的剪贴板中    
        except IndexError:  # 如果没有选定任何行，则不执行复制操作    
            pass    

    def view_to_clipboard():    
        try:    
            selected_item = tree.selection()[0]  # 获取选定的行    
            selected_values = tree.item(selected_item, "values")  # 获取选定的行的值    
            selected_dict = dict(zip(ori_owercount_easyread.columns, selected_values))  # 转换为字典格式    
            PROGRAM_display_content_in_textbox(selected_dict) 
        except IndexError:  # 如果没有选定任何行，则不执行复制操作    
            pass  
    
    def show_popup_menu(event):    
        try:    
            tree.selection_set(tree.identify_row(event.y))  # 选定点击的行    
            popup_menu.tk_popup(event.x_root, event.y_root)  # 弹出右键菜单    
        except AttributeError:  # 如果没有点击在行上，则不弹出菜单    
            pass    
    
    def handle_double_click(event):    
        df = ori_owercount_easyread  # 在这里明确df的来源是ori_owercount_easyread  
        if "报表类型" in df.columns:  # 如果双击的是“报表类型”列，显示警告框    
            for item in tree.selection():    
                selection = tree.item(item, "values")    
                content_dict = dict(zip(df.columns, selection))    
                SETTING_exe(content_dict.copy(),methon,ori) 
        else:  # 如果双击的不是“报表类型”列，弹出名为VIEWDICT的文本窗口展示“xxxx”    
            for item in tree.selection():    
                selection = tree.item(item, "values")    
                content_dict = dict(zip(df.columns, selection))    
                PROGRAM_display_content_in_textbox(content_dict)  # 同上，确保这个函数在其他地方已经定义  

    def PROGRAM_create_button(frame, text, command):  
        """创建一个按钮并添加到frame"""  
        tk.Button(  
            frame,  
            text=text,  
            bg="white",  
            font=("微软雅黑", 10),  
            relief=tk.GROOVE,  
            activebackground="green",  
            command=command,  
        ).pack(side=tk.LEFT) 
        
           
    def update_textbox(textbox, text_to_display):  
        textbox.config(text=text_to_display)  # 更新标签的文本内容   
           
    # 初始化Tk窗口    
    root = tk.Tk()    
    root.title(mytitle)    
    
    # 得到屏幕宽度和高度，并设置窗口位置和大小（这里可以自定义）    
    sw = root.winfo_screenwidth()    
    sh = root.winfo_screenheight()    
    ww = 1310  # 窗口宽度    
    wh = 600  # 窗口高度    
    x = (sw - ww) // 2    
    y = (sh - wh) // 2    
    root.geometry(f"{ww}x{wh}+{x}+{y}")    
    
    # 创建Treeview并设置列    
    tree = ttk.Treeview(root)    
    tree["columns"] = list(ori_owercount_easyread.columns)    
    current_sort_column = None    
    current_sort_order = True  # True表示升序，False表示降序    
    current_page = 0  # 当前页数，从0开始计数。这里初始化了current_page变量。  
    total_pages = len(ori_owercount_easyread) // page_size + (len(ori_owercount_easyread) % page_size > 0)  # 总页数    



    
    # 配置列宽度和标题，并绑定点击事件和双击事件    
    for i, col in enumerate(ori_owercount_easyread.columns):    
        tree.column(col, width=200, stretch=tk.NO)    
        tree.heading(col, text=col, command=lambda c=col: sort_column(c, tree))    
    tree.bind("<Double-1>", handle_double_click)  # 绑定双击事件    


    # 个性化设置列宽 
     
    column_widths=SETTING_get_width()
    for column, width in column_widths.items():  
        try:  
            tree.column(column, minwidth=0, width=width, stretch=tk.NO)  
        except:   
            pass 

    
    update_treeview(tree, ori_owercount_easyread)  # 初次加载DataFrame数据到Treeview中（第一页）    
    

    # 创建分页控件（使用ttk.Label和ttk.Button）    
    frame0 = ttk.Frame(root)    
    ttk.Button(frame0, text="<<", command=lambda: go_to_page("B",max(0, current_page - 1))).pack(side="left") 
    entry_text = tk.Label(frame0, text=str(len(ori_owercount_easyread))+"，"+str(current_page+1)+"/"+str(total_pages),width=20)  
    entry_text.pack(side=tk.LEFT)  # 确保文本框可见    
    ttk.Button(frame0, text=">>", command=lambda: go_to_page("F",min(total_pages - 1, current_page + 1))).pack(side="left")        
    frame0.pack(side="bottom", fill="x")  # 将分页控件放置在底部  
 
    framecanvas = ttk.Frame(root, width=1310, height=20)
    framecanvas.pack(side="bottom")
            
    # 右键菜单的实现      
    popup_menu = tk.Menu(root, tearoff=0)  # 创建弹出菜单实例      
    popup_menu.add_command(label="复制", command=copy_to_clipboard)  # 添加复制选项到菜单中      
    popup_menu.add_command(label="查看", command=view_to_clipboard)  # 添加复制选项到菜单中      
    tree.bind("<Button-3>", show_popup_menu)  # 绑定右键点击事件到弹出菜单函数上    
    tree.bind("<Double-1>", handle_double_click)  # 绑定双击事件     
    # 创建垂直滚动条  
    vertical_scrollbar = ttk.Scrollbar(root, orient="vertical", command=tree.yview)  
    tree.configure(yscrollcommand=vertical_scrollbar.set)  
  
    # 创建水平滚动条  
    horizontal_scrollbar = ttk.Scrollbar(root, orient="horizontal", command=tree.xview)  
    tree.configure(xscrollcommand=horizontal_scrollbar.set)    
          
    vertical_scrollbar.pack(side="right", fill="y")  # 垂直滚动条放置在右侧  
    horizontal_scrollbar.pack(side="bottom", fill="x")  # 水平滚动条放置在底部  
    
    #创建进度条
    win_progress_frame = ttk.Frame(frame0)  
    win_progress_frame.pack(side="left",pady=1)  
    win_progressbar = ttk.Progressbar(win_progress_frame, orient='horizontal', mode='determinate', length=100)  
    win_progressbar.pack(fill=tk.X, expand=True, pady=5) 
    #创建选择组件
    xt11 = tk.StringVar()  
    xt11.set("列名")  
     
    import_se1 = tk.Label(frame0, text="位置：")  
    import_se1.pack(side=tk.LEFT)      
    comvalue = tk.StringVar()  # 窗体自带的文本，新建一个值  
    comboxlist = ttk.Combobox(  
        frame0, width=12, height=30, state="readonly", textvariable=comvalue  
    )   
    comboxlist["values"] = ori_owercount_easyread.columns.tolist()  
    #comboxlist.current(0)  # 选择第一个  
    comboxlist.bind("<<ComboboxSelected>>", lambda *arg: xt11.set(comboxlist.get()))  # 绑定事件,(下拉列表框被选中时，绑定XT11SET函数)  
    comboxlist.pack(side=tk.LEFT)    
    import_se3 = tk.Label(frame0, text="文本：")  
    import_se3.pack(side=tk.LEFT)  
    
	# 使用 Text widget 替换 Entry widget  
    xentry_t22 = tk.Text(frame0, width=20, height=1)  # 设置宽度和高度以适应多行文本  
    xentry_t22.pack(side=tk.LEFT)  
	# 插入原始内容到 Text widget  
    xentry_t22.insert(tk.END,"关键词1|关键词2")  
	
    import_se4 = tk.Label(frame0, text="数据清洗：") 
    import_se4.pack(side=tk.LEFT) 
    #创建通用按钮  
    #PROGRAM_create_button(framecanvas, "测试专用", lambda:PROGRAM_display_df_in_treeview(TOOLS_ROR_STAT_0(ori_owercount_easyread),0,0)) 
	

		
		      
    PROGRAM_create_button(frame0, "含", lambda: PROGRAM_display_df_in_treeview(ori_owercount_easyread.loc[ori_owercount_easyread[xt11.get()].astype(str).str.contains(str(xentry_t22.get("1.0", tk.END).strip()), na=False)], methon, ori_owercount_easyread)) 
    PROGRAM_create_button(frame0, "无", lambda: PROGRAM_display_df_in_treeview(ori_owercount_easyread.loc[~ori_owercount_easyread[xt11.get()].astype(str).str.contains(str(xentry_t22.get("1.0", tk.END).strip()), na=False)], methon, ori_owercount_easyread))  
    PROGRAM_create_button(frame0, "是", lambda: PROGRAM_display_df_in_treeview(ori_owercount_easyread.loc[ori_owercount_easyread[xt11.get()].astype(str)==str(xentry_t22.get("1.0", tk.END).strip())], methon, ori_owercount_easyread))
    PROGRAM_create_button(frame0, "大", lambda: PROGRAM_display_df_in_treeview(ori_owercount_easyread.loc[ori_owercount_easyread[xt11.get()].astype(float)>float(xentry_t22.get("1.0", tk.END).strip())], methon, ori_owercount_easyread))  
    PROGRAM_create_button(frame0, "小", lambda: PROGRAM_display_df_in_treeview(ori_owercount_easyread.loc[ori_owercount_easyread[xt11.get()].astype(float)<float(xentry_t22.get("1.0", tk.END).strip())],methon, ori_owercount_easyread))  
    PROGRAM_create_button(frame0, "等", lambda: PROGRAM_display_df_in_treeview(ori_owercount_easyread.loc[ori_owercount_easyread[xt11.get()].astype(str)==str(xentry_t22.get("1.0", tk.END).strip())], methon, ori_owercount_easyread))
    PROGRAM_create_button(frame0, "升", lambda: PROGRAM_display_df_in_treeview(ori_owercount_easyread.sort_values(by=(xt11.get()),ascending=[True],na_position="last") , methon, ori_owercount_easyread))
    PROGRAM_create_button(frame0, "降", lambda: PROGRAM_display_df_in_treeview(ori_owercount_easyread.sort_values(by=(xt11.get()),ascending=[False],na_position="last") , methon, ori_owercount_easyread))
    PROGRAM_create_button(frame0, "扩", lambda: PROGRAM_display_df_in_treeview(CLEAN_expand_rows(ori_owercount_easyread.copy(),str(xentry_t22.get("1.0", tk.END).strip()), [xt11.get()]),0,0)) 




    #增加PSUR兼容性
    if psur==0:

        PROGRAM_create_button(frame0, "找", lambda: PROGRAM_display_df_in_treeview(SMALL_find_based_on_expression(ori_owercount_easyread,str(xentry_t22.get("1.0", tk.END).strip())), methon, ori_owercount_easyread))
        PROGRAM_create_button(frame0, "透", lambda: PROGRAM_display_df_in_treeview(SMALL_add_count_and_ratio(ori_owercount_easyread,xt11.get()),1,ori_owercount_easyread) )
        PROGRAM_create_button(frame0, "拆", lambda: PROGRAM_display_df_in_treeview(SMALL_add_count_and_ratio_exp(ori_owercount_easyread,str(xt11.get()),xentry_t22.get("1.0", tk.END).strip()),methon,ori_owercount_easyread) )
        PROGRAM_create_button(frame0, "存", lambda: TOOLS_temp_save_df(ori_owercount_easyread.copy()) )


        PROGRAM_create_button(framecanvas, "数据清洗", lambda: CLEAN_table(ori_owercount_easyread.copy()))
        PROGRAM_create_button(framecanvas, "分组透视", lambda:TOOLS_create_pivot_tool_gui(ori_owercount_easyread.copy(),ori_owercount_easyread)) 	
        PROGRAM_create_button(framecanvas, "暂存合并", lambda: TOOLS_temp_save_df(ori_owercount_easyread.copy()) )
        PROGRAM_create_button(framecanvas, "图形绘制", lambda:DRAW_plot_df(ori_owercount_easyread))     
            
        #创建菜单栏
        SETTING_create_menu({"windows":root,"win_progressbar":win_progressbar,"ori_owercount_easyread":ori_owercount_easyread,"ori":ori})  
    
    #创建PSUR专属按钮
    if methon=="psur" or psur==1:
        PROGRAM_create_button(framecanvas, "导入数据",lambda:PSUR_open(0))
        PROGRAM_create_button(framecanvas, "导出数据", command=lambda:SMALL_save_dict(ori_owercount_easyread))
    if methon=="psur" and '报告表编码' in ori_owercount_easyread.columns:
        PROGRAM_create_button(framecanvas, "导入说明书", lambda:PSUR_check_adr_in_word(ori_owercount_easyread)) 	
        PROGRAM_create_button(framecanvas, "报告类型", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['-伤害'], ['怀疑/并用'], ['报告表编码'], ['nunique'], {'不良反应名称（规整）': 'count_ALL', '疑似新的ADR': 'count_ALL', '重点关注ADR': 'count_ALL'}, ['报告表编码合计']]),'psur',ori))       
        PROGRAM_create_button(framecanvas, "不良反应汇总表（SOC）",  lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['SOC'], ['-伤害'], ['报告表编码'], ['count'], {'不良反应名称（规整）': 'count_ALL', '疑似新的ADR': 'count_ALL', '重点关注ADR': 'count_ALL'},  ['报告表编码合计']]),'psur',ori))      
        PROGRAM_create_button(framecanvas, "不良反应汇总表（PT）", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['SOC', '不良反应名称（规整）'], ['-伤害'], ['报告表编码'], ['count'], {'疑似新的ADR': 'count_ALL', '重点关注ADR': 'count_ALL'},  ['报告表编码合计']]),'psur',ori))      
        PROGRAM_create_button(framecanvas, "特殊人群用药（年龄段）", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['年龄段'], ['-伤害'], ['报告表编码'], ['nunique'], {'不良反应名称（规整）': 'count_ALL', '疑似新的ADR': 'count_ALL', '重点关注ADR': 'count_ALL'}, ['报告表编码合计']]),'psur',ori))      
        PROGRAM_create_button(framecanvas, "特殊人群用药（肝肾孕哺）", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['特殊人群'], ['-伤害'], ['报告表编码'], ['nunique'], {'不良反应名称（规整）': 'count_ALL', '疑似新的ADR': 'count_ALL', '重点关注ADR': 'count_ALL'}, ['报告表编码合计']]),'psur',ori))       
        PROGRAM_create_button(framecanvas, "性别", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['性别'], ['-伤害'], ['报告表编码'], ['nunique'], {'不良反应名称（规整）': 'count_ALL', '疑似新的ADR': 'count_ALL', '重点关注ADR': 'count_ALL'}, ['报告表编码合计']]),'psur',ori))       
        PROGRAM_create_button(framecanvas, "时隔", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['不良反应发生时间减用药开始时间'], ['-伤害'], ['报告表编码'], ['nunique'], {'不良反应名称（规整）': 'count_ALL', '疑似新的ADR': 'count_ALL', '重点关注ADR': 'count_ALL'}, ['报告表编码合计']]),'psur',ori))       
        PROGRAM_create_button(framecanvas, "报告年份", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['报告年份'], ['-伤害'], ['报告表编码'], ['nunique'], {'不良反应名称（规整）': 'count_ALL', '疑似新的ADR': 'count_ALL', '重点关注ADR': 'count_ALL'}, ['报告表编码合计']]),'psur',ori))       
        PROGRAM_create_button(framecanvas, "用法用量", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['用法用量'], ['-伤害'], ['报告表编码'], ['nunique'], {'不良反应名称（规整）': 'count_ALL', '疑似新的ADR': 'count_ALL', '重点关注ADR': 'count_ALL'}, ['报告表编码合计']]),'psur',ori))       
        PROGRAM_create_button(framecanvas, "批号", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['生产批号'], ['-伤害'], ['报告表编码'], ['nunique'], {'不良反应名称（规整）': 'count_ALL', '疑似新的ADR': 'count_ALL', '重点关注ADR': 'count_ALL'}, ['报告表编码合计']]),'psur',ori))       
        PROGRAM_create_button(framecanvas, "批准文号", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['怀疑/并用','批准文号','通用名称'], ['-伤害'], ['报告表编码'], ['nunique'], {'不良反应名称（规整）': 'count_ALL','重点关注ADR': 'count_ALL'},  ['报告表编码合计']]),'psur',ori) )
        PROGRAM_create_button(framecanvas, "通用名称", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[['怀疑/并用','通用名称'], ['-伤害'], ['报告表编码'], ['nunique'], {'不良反应名称（规整）': 'count_ALL','重点关注ADR': 'count_ALL'},  ['报告表编码合计']]),'psur',ori) )

        PROGRAM_create_button(frame0, "PT计数", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[[xt11.get()], ['-伤害'], ['报告表编码'], ['count'], {'不良反应名称（规整）': 'count_ALL', '疑似新的ADR': 'count_ALL', '重点关注ADR': 'count_ALL'}, ['报告表编码合计']]),'psur',ori))       
        PROGRAM_create_button(frame0, "编码唯一值计数", lambda:PROGRAM_display_df_in_treeview(TOOLS_create_pivot_tool(ori_owercount_easyread,[[xt11.get()], ['-伤害'], ['报告表编码'], ['nunique'], {'不良反应名称（规整）': 'count_ALL', '疑似新的ADR': 'count_ALL', '重点关注ADR': 'count_ALL'}, ['报告表编码合计']]),'psur',ori))       


           
    # 在Tk窗口中放置Treeview并运行主循环      
    tree.pack(side="top", fill="both", expand=True) # Treeview填充窗口的剩余空间      

    root.lift()
    #root.attributes("-topmost", True)
    #root.attributes("-topmost", False)  
    root.mainloop() # 启动Tk窗口的主事件循环```python
        
#





default_data = pd.DataFrame({  
    'PT': ['A;B;C', 'C;D', 'E;F', 'G;H', 'I'],  
    'SOC': ['K;L;M', 'M;N', 'O;P', 'Q;R', 'S'],  
    'col3': [1, 2, 3, 4, 5],  
    'col4': [11, 22, 33, 44, 55],  
    'col5': [111, 222, 333, 444, 555]  
})  
  
#default_data=pd.read_excel(r"0.xlsx") 

def mainx(df,title, textbox_text, method_list, method_text):  
    root = tk.Tk()  
    root.title(title)  
    sw = root.winfo_screenwidth()    
    sh = root.winfo_screenheight()    
    ww = 600  # 窗口宽度    
    wh = 400  # 窗口高度    
    x = (sw - ww) // 2    
    y = (sh - wh) // 2    
    root.geometry(f"{ww}x{wh}+{x}+{y}")  
  
    # 创建左侧树状视图，显示df的列名  
    tree_frame = ttk.Frame(root)  
    tree_frame.pack(side="left", fill="both", expand=True, padx=10, pady=10)  
    tree_label = ttk.Label(tree_frame, text="多选列名:")  
    tree_label.pack(anchor='w')  
  
    # 添加滚动条  
    scrollbar = ttk.Scrollbar(tree_frame)  
    scrollbar.pack(side="right", fill="y")  
  
    tree = ttk.Treeview(tree_frame, yscrollcommand=scrollbar.set)  
    scrollbar.config(command=tree.yview)  
    tree['columns'] = ('Column',)  
    tree.column('#0', width=200, stretch=tk.NO)  
    tree.heading('#0', text='列名', anchor=tk.W)  
  
    for col in df.columns:  
        tree.insert("", tk.END, text=str(col), values=())
 
  
    tree.pack(fill="both", expand=True, padx=5, pady=5)  
  
    # 创建右侧内容区域  
    right_frame = ttk.Frame(root)  
    right_frame.pack(side="right", fill="both", expand=True, padx=10, pady=10)  
  
    # 创建文本框  
    text_label = ttk.Label(right_frame, text=textbox_text)  
    text_label.pack(anchor='w')  
  
    text_box = tk.Entry(right_frame, width=20)  # 使用Entry作为文本框，只需要一行  
    text_box.pack(padx=5, pady=[10, 0])  # 增加上方间距，使文本框离标签有一定距离  
  
    # 创建下拉菜单  
    dropdown_label = ttk.Label(right_frame, text=method_text)  
    dropdown_label.pack(anchor='w')  
  
    var = tk.StringVar(right_frame)  
    var.set(method_list[0])  # 设置默认选项为第一列  
    dropdown = ttk.Combobox(right_frame, textvariable=var, values=method_list)  
    dropdown.pack(padx=5, pady=[10, 0])  # 增加上方间距，使下拉菜单离文本框有一定距离  
  
    # 创建确定按钮  
    def on_confirm():  
        selected_columns = [tree.item(item_id, "text") for item_id in tree.selection()]   # 获取多选的列名  
        selected_column = var.get()  # 获取单选的列名 
        entry_content = text_box.get()  # 获取Entry的内容  
        print("Entry的内容:", entry_content) 
        print("多选的列名:", selected_columns)  
        print("单选的列名:", selected_column)  
        # 在这里添加确认按钮的逻辑，可以使用selected_columns和selected_column进行后续操作  
  
    confirm_button = tk.Button(right_frame, text="确定", command=on_confirm)  
    confirm_button.pack(padx=10, pady=[20, 10])  # 增加上下间距，使按钮离下拉菜单和底部都有一定距离  
  
    root.mainloop()




#main(default_data,"行列扩展","分隔符：",["扩行","扩列"],"需要的操作") 
  
#default_data = expand_rows(default_data, ";", 'PT', 'SOC')  




  





    

default_data = pd.DataFrame({  
        '条目': ['名称', '版本', '用途说明'],  
        '信息': ['易析数据统计分析工具', version_now, '供数据统计使用。']})    
     

if psur==1:
    default_data = pd.DataFrame({  
        '条目': ['名称', '版本', '用途说明', "❤注意事项"],  
        '信息': ['药品品种监测数据辅助分析工具', version_now,'供监管机构使用。', "❤不良反应汇总（PT或者SOC）报表中，一般、严重、新的一般、新的严重以及合计统计的是PT的数量小计；❤其他报表上述项目统计的是报告编码唯一值的数量。"]})    
     
    #setting=pd.read_excel(r"1.xls") 
    #default_data = CLEAN_expand_cols(default_data, ";", ['PT', 'SOC']) 



def AAA_my_main():

    global version_now
    global usergroup
    global setting_cfg
    global csdir
    global peizhidir
    global biaozhun
    global global_dfs
    global psur
    global mytitle 
    #序列好验证、配置表生成与自动更新。
    setting_cfg=INI_read_setting_cfg()
    INI_generate_random_file()
    setting_cfg=INI_open_setting_cfg()
    if setting_cfg["settingdir"]==0:
        showinfo(title="提示", message="未发现默认配置文件夹，请选择一个。如该配置文件夹中并无配置文件，将生成默认配置文件。")
        filepathu=filedialog.askdirectory()
        path=INI_get_directory_path(filepathu)
        INI_update_setting_cfg("settingdir",path)        
    setting_cfg=INI_open_setting_cfg()
    random_number=int(setting_cfg["sidori"])
    input_number=int(str(setting_cfg["sidfinal"])[0:6])
    day_end=INI_convert_and_compare_dates(str(setting_cfg["sidfinal"])[6:14])
    sid=random_number*2+183576
    if input_number == sid  and day_end=="未过期":
        usergroup="用户组=1" 
        print(usergroup+"   有效期至：")
        print(datetime.strptime(str(int(int(str(setting_cfg["sidfinal"])[6:14])/4)), "%Y%m%d") )
    else:
        print(usergroup)    
    print("\n配置文件路径："+setting_cfg["settingdir"]+"\n")
    peizhidir=str(setting_cfg["settingdir"])+csdir.split("pinggutools")[0][-1]

    try:
        biaozhun["药品清洗"]=pd.read_excel(peizhidir+'easy_药品规整-基础清洗.xlsx').reset_index(drop=True) 
        biaozhun["药品PT清洗"]=pd.read_excel(peizhidir+'easy_药品规整-PT标准化.xlsx').reset_index(drop=True)
        biaozhun["药品关键词"]=pd.read_excel(peizhidir+'easy_药品规整-SOC-关键词.xlsx').reset_index(drop=True) 
        
        biaozhun["药品不良反应库"]=pd.read_excel(peizhidir+'easy_药品规整-不良反应库.xlsx',sheet_name="已知不良反应库").reset_index(drop=True)  

        biaozhun["药品重点关注库"]=pd.read_excel(peizhidir+'easy_药品规整-不良反应库.xlsx',sheet_name="重点关注").reset_index(drop=True)       

        biaozhun["器械清洗"]=pd.read_excel(peizhidir+'easy_器械规整-基础清洗.xlsx').reset_index(drop=True) 
        biaozhun["器械关键词"]=pd.read_excel(peizhidir+'easy_器械规整-SOC-关键词.xlsx').reset_index(drop=True) 
        biaozhun["器械关键词（仅故障表现）"]= biaozhun["器械关键词"].copy()
        biaozhun["器械关键词（仅故障表现）"]['查找位置']="器械故障表现"

        biaozhun["meddra"]=0   
       
        
        print("已载入标准库。")
    except:
        print(csdir)
        if psur!=1:
            print("未载入标准库,部分功能可能无法使用。如需标准库，请联系开发者索取。")

    try:
        biaozhun["药品不良反应库-AI"]=pd.read_excel(peizhidir+'easy_药品规整-不良反应库.xlsx',sheet_name="已知不良反应库-AI").reset_index(drop=True)  
        biaozhun["药品分类库"]=pd.read_excel(peizhidir+'easy_药品规整-药品分类-关键词.xlsx').reset_index(drop=True)  
    except:
        pass      

    if psur==1:
        PSUR_ini(0)  
           
    app = PROGRAM_display_df_in_treeview(default_data,0,default_data)  # 创建应用实例  
    #app=DRAW_plot_df(default_data)




       
if __name__ == "__main__":  
	
    AAA_my_main()
    
    
    
    print("done.")
